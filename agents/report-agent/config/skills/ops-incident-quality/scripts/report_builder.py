"""
Report Builder for Ops Deep Dive Report.
Generates HTML (flowing document) and DOCX (Word) reports.
Modern Light Theme with AI Insights
"""
import os
from datetime import datetime
from jinja2 import Template
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from utils import OUTPUT_DIR, IMG_DIR, COLORS
from i18n import get_text, format_date, get_report_filename

class ReportBuilder:
    def __init__(self, results, charts, insights=None, language='en'):
        self.r = results
        self.charts = charts
        self.insights = insights or {}
        self.language = language
        
    def build_html(self):
        """Generate HTML flowing document report."""
        template = Template(HTML_TEMPLATE)
        html = template.render(
            date=format_date(datetime.now().strftime('%Y-%m-%d'), self.language),
            r=self.r,
            charts=self.charts,
            insights=self.insights,
            lang=self.language,
            get_text=get_text
        )

        # Get start and end dates from results
        date_range = self.r['summary']['date_range']
        if ' to ' in date_range:
            start_date, end_date = date_range.split(' to ')
        elif ' 至 ' in date_range:
            start_date, end_date = date_range.split(' 至 ')
        else:
            start_date = end_date = datetime.now().strftime('%Y-%m-%d')

        filename = get_report_filename(start_date, end_date, self.language, 'html')
        path = os.path.join(OUTPUT_DIR, filename)
        with open(path, 'w', encoding='utf-8') as f:
            f.write(html)
        return path
    
    def build_docx(self):
        """Generate Word document report with AI insights and professional styling."""
        doc = Document()
        r = self.r
        insights = self.insights
        
        # --- Helper for Styling ---
        def style_table(table):
            """Apply professional styling to table (three-line style)."""
            # Header Row
            for cell in table.rows[0].cells:
                # Add light gray shading
                from docx.oxml.ns import nsdecls
                from docx.oxml import parse_xml
                shading_elm = parse_xml(r'<w:shd {} w:fill="F3F4F6"/>'.format(nsdecls('w')))
                cell._tc.get_or_add_tcPr().append(shading_elm)
                # Bold Text & Borders (This relies on Table Grid + manual font tweaks)
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs:
                        run.font.bold = True
                        run.font.size = Pt(11)
            
            # Content Rows
            for row in table.rows[1:]:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in p.runs:
                            run.font.size = Pt(11)

        # --- Cover Page ---
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph()
        title = doc.add_heading(get_text('report_title', self.language), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Company/Team Name (Placeholder)
        team_subtitle = doc.add_paragraph()
        team_subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = team_subtitle.add_run(get_text('report_subtitle', self.language))
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(59, 130, 246)  # Primary Blue

        info_para = doc.add_paragraph()
        info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info_para.space_before = Pt(24)
        run = info_para.add_run(f"\n{get_text('data_period', self.language)}: {r['summary']['date_range']}")
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(107, 114, 128)

        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = date_para.add_run(f"{get_text('report_generated', self.language)}: {format_date(datetime.now().strftime('%Y-%m-%d'), self.language)}")
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(156, 163, 175) # Lighter gray
        
        doc.add_page_break()
        
        # --- Helper to add insight box ---
        def add_insight(insight_key):
            if insight_key in insights and insights[insight_key]:
                # Add a visually distinct paragraph for insights
                p = doc.add_paragraph()
                p.space_before = Pt(12)
                p.space_after = Pt(12)

                # Add border or background simulation (Word API is limited, using bold/color)
                run = p.add_run(f"💡 {get_text('ai_insight', self.language)}")
                run.bold = True
                run.font.color.rgb = RGBColor(59, 130, 246) # Primary blue
                run.font.size = Pt(12)

                p_content = doc.add_paragraph(insights[insight_key])
                p_content.paragraph_format.left_indent = Inches(0.25)
                # Italicize slightly for "voice" feel
                for run in p_content.runs:
                    run.font.italic = True
                    run.font.color.rgb = RGBColor(55, 65, 81) # Dark gray
                doc.add_paragraph() # Spacer

        # 1. Executive Summary
        doc.add_heading(f'1. {get_text("executive_summary", self.language)}', level=1)
        summary = r['summary']

        # Professional KPI Table
        table = doc.add_table(rows=2, cols=4)
        table.style = 'Table Grid'
        style_table(table)

        headers = [
            get_text('sla_compliance', self.language),
            get_text('avg_mttr', self.language),
            get_text('p1_incidents', self.language),
            get_text('total_tickets', self.language)
        ]
        values = [f"{summary['sla_rate']:.1f}%", f"{summary['avg_mttr']:.1f}{get_text('hours', self.language)}",
                  str(summary['p1_count']), str(summary['total'])]
        
        # Fill Header
        for i, h in enumerate(headers):
            table.cell(0, i).text = h
        
        # Fill Values with Colors for KPIs
        for i, v in enumerate(values):
            cell = table.cell(1, i)
            cell.text = v
            # Highlight P1 count in red if > 0
            if i == 2 and int(summary['p1_count']) > 0:
                 for p in cell.paragraphs:
                     for run in p.runs:
                         run.font.color.rgb = RGBColor(239, 68, 68) # Red
                         run.font.bold = True
        
        # Re-apply styling to ensure bold headers persist after text assignment
        style_table(table)
        
        add_insight('summary')

        if summary['anomalies']:
            h2 = doc.add_heading(get_text('key_alerts', self.language), level=2)
            h2.paragraph_format.space_before = Pt(12)
            for a in summary['anomalies']:
                p = doc.add_paragraph(style='List Bullet')
                run = p.add_run(f"⚠️ {a}")
                run.font.color.rgb = RGBColor(245, 158, 11) # Amber

        # 2. SLA Analysis
        doc.add_heading(f'2. {get_text("sla_analysis", self.language)}', level=1)
        if 'sla_rate_chart' in self.charts:
            doc.add_picture(self.charts['sla_rate_chart'], width=Inches(6))
        if 'sla_violations_by_person' in self.charts:
            doc.add_paragraph()
            doc.add_picture(self.charts['sla_violations_by_person'], width=Inches(6))
        add_insight('sla')
        
        # 3. Priority Analysis
        doc.add_heading(f'3. {get_text("priority_analysis", self.language)}', level=1)
        if 'priority_dist' in self.charts:
            doc.add_picture(self.charts['priority_dist'], width=Inches(5))
        if 'mttr_by_priority' in self.charts:
            doc.add_paragraph()
            doc.add_picture(self.charts['mttr_by_priority'], width=Inches(6))
        add_insight('priority')

        # 4. Personnel Analysis
        doc.add_heading(f'4. {get_text("personnel_analysis", self.language)}', level=1)
        if 'personnel_matrix' in self.charts:
            doc.add_picture(self.charts['personnel_matrix'], width=Inches(6))
        add_insight('personnel')

        # 5. Category Analysis
        doc.add_heading(f'5. {get_text("category_analysis", self.language)}', level=1)
        if 'category_dist' in self.charts:
            doc.add_picture(self.charts['category_dist'], width=Inches(5))
        add_insight('category')

        # 6. Time Analysis
        doc.add_heading(f'6. {get_text("time_analysis", self.language)}', level=1)
        if 'heatmap' in self.charts:
            doc.add_picture(self.charts['heatmap'], width=Inches(6))
        add_insight('time')

        # 7. Recommendations
        doc.add_heading(f'7. {get_text("recommendations", self.language)}', level=1)
        if r['recommendations']:
            for rec in r['recommendations']:
                # Use Word-compatible text markers instead of emoji
                priority_marker = '[HIGH]' if rec['priority'] == 'HIGH' else '[MEDIUM]'
                
                p = doc.add_paragraph(style='List Bullet')
                p.paragraph_format.space_after = Pt(6)
                
                # Priority Tag
                run_p = p.add_run(priority_marker)
                run_p.bold = True
                if rec['priority'] == 'HIGH':
                    run_p.font.color.rgb = RGBColor(220, 38, 38)  # Red
                else:
                    run_p.font.color.rgb = RGBColor(217, 119, 6)  # Amber
                
                # Category Tag
                run_c = p.add_run(f" [{rec['category']}] ")
                run_c.bold = True
                run_c.font.color.rgb = RGBColor(75, 85, 99) # Gray
                
                # Text
                p.add_run(rec['text'])
        else:
            no_issues_text = "No urgent issues requiring attention." if self.language == 'en' else "没有需要关注的紧急问题。"
            doc.add_paragraph(no_issues_text)

        # Get start and end dates from results
        date_range = r['summary']['date_range']
        if ' to ' in date_range:
            start_date, end_date = date_range.split(' to ')
        elif ' 至 ' in date_range:
            start_date, end_date = date_range.split(' 至 ')
        else:
            start_date = end_date = datetime.now().strftime('%Y-%m-%d')

        filename = get_report_filename(start_date, end_date, self.language, 'docx')
        path = os.path.join(OUTPUT_DIR, filename)
        doc.save(path)
        return path


# ============================================
# HTML TEMPLATE - SINGLE COLUMN WITH AI INSIGHTS
# ============================================
HTML_TEMPLATE = '''
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{{ get_text('report_title', lang) }}</title>
    <link href="https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&display=swap" rel="stylesheet">
    <style>
        :root {
            /* Corporate Modern Palette - Lower Saturation Matches utils.py */
            --primary: #3b82f6;       /* Blue 500 */
            --primary-light: #93c5fd;
            --success: #10b981;       /* Emerald 500 */
            --warning: #f59e0b;       /* Amber 500 */
            --danger: #ef4444;        /* Red 500 */
            --text: #374151;          /* Gray 700 */
            --text-secondary: #6b7280;/* Gray 500 */
            --border: #e5e7eb;        /* Gray 200 */
            --bg: #f3f4f6;            /* Gray 100 */
            --card: #ffffff;
        }
        
        * { box-sizing: border-box; margin: 0; padding: 0; }
        
        body {
            font-family: 'Inter', -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, sans-serif;
            background: var(--bg);
            color: var(--text);
            line-height: 1.6;
            -webkit-font-smoothing: antialiased;
        }
        
        .container { max-width: 1000px; margin: 0 auto; padding: 40px 24px; }
        
        header {
            text-align: center;
            padding: 40px 0 40px;
            margin-bottom: 32px;
        }
        
        header h1 { 
            font-size: 32px; 
            font-weight: 700; 
            margin-bottom: 8px; 
            color: #111827; /* Gray 900 */
            letter-spacing: -0.025em;
        }
        header .subtitle { 
            font-size: 14px; 
            color: var(--text-secondary); 
            background: #fff;
            display: inline-block;
            padding: 4px 12px;
            border-radius: 999px;
            border: 1px solid var(--border);
        }
        
        .section { margin-bottom: 48px; }
        
        .section-title {
            font-size: 20px;
            font-weight: 700;
            margin-bottom: 20px;
            display: flex;
            align-items: center;
            gap: 10px;
            color: #111827;
        }
        
        .kpi-grid {
            display: grid;
            grid-template-columns: repeat(4, 1fr);
            gap: 20px;
            margin-bottom: 24px;
        }
        
        /* New "Clean & Professional" KPI Card Design */
        .kpi-box {
            background: var(--card);
            border: 1px solid var(--border);
            border-radius: 12px;
            padding: 24px 20px;
            box-shadow: 0 1px 3px rgba(0,0,0,0.05);
            position: relative;
            overflow: hidden;
            transition: transform 0.2s;
        }
        
        .kpi-box::before {
            content: '';
            position: absolute;
            top: 0; left: 0; width: 100%; height: 4px;
            background: var(--primary);
        }
        
        .kpi-box.success::before { background: var(--success); }
        .kpi-box.warning::before { background: var(--warning); }
        .kpi-box.danger::before { background: var(--danger); }
        
        .kpi-box .label { 
            font-size: 12px; 
            font-weight: 600; 
            text-transform: uppercase; 
            color: var(--text-secondary); 
            letter-spacing: 0.05em;
            margin-bottom: 8px;
        }
        
        .kpi-box .value { 
            font-size: 36px; 
            font-weight: 700; 
            color: #111827;
            line-height: 1.1;
            margin-bottom: 8px;
        }
        
        .kpi-box .trend { 
            font-size: 13px; 
            font-weight: 500;
            display: inline-flex;
            align-items: center;
            gap: 4px;
        }
        .trend.neutral { color: var(--text-secondary); }
        
        /* Structured Insight Box */
        .insight-box {
            background: #f8fafc; /* Slate 50 */
            border: 1px solid #e2e8f0;
            border-radius: 8px;
            padding: 20px;
            margin: 24px 0;
            position: relative;
        }
        
        .insight-box h4 { 
            font-size: 13px; 
            font-weight: 700; 
            color: var(--primary); 
            text-transform: uppercase;
            letter-spacing: 0.05em;
            margin-bottom: 10px;
            display: flex;
            align-items: center; 
            gap: 6px; 
        }
        
        .insight-box p { 
            font-size: 15px; 
            color: #334155; /* Slate 700 */
            line-height: 1.6;
        }
        
        .chart-card {
            background: var(--card);
            border-radius: 12px;
            padding: 24px;
            border: 1px solid var(--border);
            margin-bottom: 20px;
            box-shadow: 0 1px 2px rgba(0,0,0,0.05);
        }
        
        .chart-card h3 {
            font-size: 14px;
            font-weight: 600;
            color: #4b5563;
            margin-bottom: 16px;
        }
        
        .chart-card img { width: 100%; height: auto; border-radius: 4px; }
        
        .alert-box {
            background: #fffbeb;
            border: 1px solid #fcd34d;
            border-radius: 8px;
            padding: 16px 20px;
            margin: 24px 0;
        }
        
        .alert-box h3 { color: #b45309; font-size: 14px; font-weight: 700; margin-bottom: 8px; text-transform: uppercase; }
        .alert-box ul { list-style: none; }
        .alert-box li { 
            margin-bottom: 6px; 
            color: #92400e; 
            font-size: 14px;
            display: flex;
            gap: 8px;
        }
        .alert-box li::before { content: "⚠️"; font-size: 12px; margin-top: 3px; }
        
        /* Table Styles */
        table { width: 100%; border-collapse: collapse; font-size: 14px; margin-top: 16px; }
        th { 
            background: #f9fafb; 
            color: #4b5563; 
            padding: 12px 16px; 
            text-align: left; 
            font-size: 12px; 
            font-weight: 600;
            text-transform: uppercase; 
            border-bottom: 1px solid var(--border);
        }
        td { padding: 12px 16px; border-bottom: 1px solid var(--border); color: #374151; }
        tr:last-child td { border-bottom: none; }
        
        .priority { display: inline-block; padding: 2px 8px; border-radius: 4px; font-size: 11px; font-weight: 600; color: white; }
        .priority-p1 { background: var(--danger); }
        .priority-p2 { background: var(--warning); }
        .priority-p3 { background: var(--primary); }
        .priority-p4 { background: var(--success); }
        
        .case-card {
            background: var(--card);
            border: 1px solid var(--border);
            border-radius: 8px;
            padding: 16px;
            margin-bottom: 12px;
            font-size: 14px;
        }
        
        .rec-item {
            display: flex;
            align-items: flex-start;
            gap: 12px;
            padding: 16px;
            background: var(--card);
            border: 1px solid var(--border);
            border-radius: 8px;
            margin-bottom: 12px;
        }
        
        .rec-priority { 
            font-size: 10px; 
            font-weight: 700; 
            padding: 2px 6px; 
            border-radius: 4px;
            text-transform: uppercase;
            white-space: nowrap;
            margin-top: 2px;
        }
        .rec-priority.high { background: #fee2e2; color: #b91c1c; }
        .rec-priority.medium { background: #fef3c7; color: #b45309; }
        
        .rec-content { flex: 1; }
        .rec-category { 
            font-size: 11px; 
            font-weight: 600; 
            color: var(--text-secondary); 
            text-transform: uppercase; 
            margin-bottom: 4px;
            display: block;
        }
        .rec-text { color: var(--text); font-size: 14px; }
    </style>
</head>
<body>
<div class="container">

<header>
    <h1>{{ get_text('report_title', lang) }}</h1>
    <span class="subtitle">{{ r.summary.date_range }} | {{ get_text('report_generated', lang) }}: {{ date }}</span>
</header>

<!-- Executive Summary -->
<section class="section">
    <h2 class="section-title">{{ get_text('executive_summary', lang) }}</h2>
    <div class="kpi-grid">
        <div class="kpi-box">
            <div class="label">{{ get_text('sla_compliance', lang) }}</div>
            <div class="value" style="color: var(--primary);">{{ "%.1f"|format(r.summary.sla_rate) }}%</div>
        </div>
        <div class="kpi-box success">
            <div class="label">{{ get_text('avg_mttr', lang) }}</div>
            <div class="value" style="color: var(--success);">{{ "%.1f"|format(r.summary.avg_mttr) }}{{ get_text('hours', lang) }}</div>
            <div class="trend neutral">
                <span style="color: {{ '#10b981' if r.summary.mttr_trend < 0 else '#ef4444' }}; font-weight: 700;">
                    {{ "↓" if r.summary.mttr_trend < 0 else "↑" }} {{ "%.1f"|format(r.summary.mttr_trend|abs) }}%
                </span> vs prev
            </div>
        </div>
        <div class="kpi-box danger">
            <div class="label">{{ get_text('p1_incidents', lang) }}</div>
            <div class="value" style="color: var(--danger);">{{ r.summary.p1_count }}</div>
        </div>
        <div class="kpi-box warning">
            <div class="label">{{ get_text('total_tickets', lang) }}</div>
            <div class="value" style="color: var(--warning);">{{ r.summary.total }}</div>
            <div class="trend neutral">
                <span style="color: var(--text-secondary);">
                    {{ "↑" if r.summary.vol_trend > 0 else "↓" }} {{ "%.1f"|format(r.summary.vol_trend|abs) }}%
                </span> vs prev
            </div>
        </div>
    </div>
    {% if insights.summary %}
    <div class="insight-box">
        <h4>💡 {{ get_text('ai_insight', lang) }}</h4>
        <p>{{ insights.summary }}</p>
    </div>
    {% endif %}
    {% if r.summary.anomalies %}
    <div class="alert-box">
        <h3>{{ get_text('key_alerts', lang) }}</h3>
        <ul>
        {% for a in r.summary.anomalies %}
            <li>{{ a }}</li>
        {% endfor %}
        </ul>
    {% endif %}
</section>

<!-- SLA Analysis -->
<section class="section">
    <h2 class="section-title">📊 {{ get_text('sla_analysis', lang) }}</h2>
    <div class="chart-card"><h3>Compliance by Priority</h3><img src="images/sla_rate_chart.png"></div>
    <div class="chart-card"><h3>Top Violators (Personnel)</h3><img src="images/sla_violations_by_person.png"></div>
    <div class="chart-card"><h3>Top Violators (Category)</h3><img src="images/sla_violations_by_category.png"></div>
    <div class="chart-card"><h3>Violation Severity</h3><img src="images/sla_severity.png"></div>
    {% if insights.sla %}
    <div class="insight-box">
        <h4>💡 {{ get_text('ai_insight', lang) }}</h4>
        <p>{{ insights.sla }}</p>
    </div>
    {% endif %}
</section>

<!-- Priority Analysis -->
<section class="section">
    <h2 class="section-title">🎯 {{ get_text('priority_analysis', lang) }}</h2>
    <div class="chart-card"><h3>Priority Distribution</h3><img src="images/priority_dist.png"></div>
    <div class="chart-card"><h3>MTTR by Priority</h3><img src="images/mttr_by_priority.png"></div>
    {% if insights.priority %}
    <div class="insight-box">
        <h4>💡 {{ get_text('ai_insight', lang) }}</h4>
        <p>{{ insights.priority }}</p>
    </div>
    {% endif %}
    <h3 style="margin: 24px 0 16px; font-size: 16px;">P1/P2 Case Review</h3>
    {% for case in r.priority.p1_cases[:3] %}
    <div class="case-card">
        <span class="priority priority-p1">P1</span>
        <strong style="margin-left: 10px;">{{ case['Order Number'] }}</strong>: {{ case['Order Name'][:55] }}...
        <br><small style="color: var(--text-secondary);">{{ get_text('resolver', lang) }}: {{ case['Resolver'] }} | MTTR: {{ "%.1f"|format(case['Resolution_Hours']) }}{{ get_text('hours', lang) }}</small>
    </div>
    {% endfor %}
    {% for case in r.priority.p2_cases[:2] %}
    <div class="case-card">
        <span class="priority priority-p2">P2</span>
        <strong style="margin-left: 10px;">{{ case['Order Number'] }}</strong>: {{ case['Order Name'][:55] }}...
        <br><small style="color: var(--text-secondary);">{{ get_text('resolver', lang) }}: {{ case['Resolver'] }} | MTTR: {{ "%.1f"|format(case['Resolution_Hours']) }}{{ get_text('hours', lang) }}</small>
    </div>
    {% endfor %}
</section>

<!-- Personnel Analysis -->
<section class="section">
    <h2 class="section-title">👥 {{ get_text('personnel_analysis', lang) }}</h2>
    <div class="chart-card"><h3>Performance Matrix</h3><img src="images/personnel_matrix.png"></div>
    <div class="chart-card"><h3>Volume Ranking</h3><img src="images/volume_rank.png"></div>
    <div class="chart-card"><h3>SLA Violation Ranking</h3><img src="images/sla_violation_rank.png"></div>
    <div class="stat-box" style="margin-top: 20px;">
        <div class="number">{{ r.personnel.unassigned.count }}</div>
        <div class="label">Unassigned Tickets ({{ "%.1f"|format(r.personnel.unassigned.percentage) }}%)</div>
    </div>
    {% if insights.personnel %}
    <div class="insight-box">
        <h4>💡 {{ get_text('ai_insight', lang) }}</h4>
        <p>{{ insights.personnel }}</p>
    </div>
    {% endif %}
</section>

<!-- Category Analysis -->
<section class="section">
    <h2 class="section-title">📁 {{ get_text('category_analysis', lang) }}</h2>
    <div class="chart-card"><h3>Category Distribution</h3><img src="images/category_dist.png"></div>
    <div class="chart-card"><h3>Keyword Mining</h3><img src="images/keywords.png"></div>
    <div class="chart-card">
        <h3>Fastest Growing Categories</h3>
        <table>
            <tr><th>{{ get_text('category', lang) }}</th><th>Change</th></tr>
            {% for name, change in r.category.fastest_growing[:5] %}
            <tr>
                <td>{{ name }}</td>
                <td style="color: {% if change > 0 %}var(--danger){% else %}var(--success){% endif %}; font-weight: 600;">{% if change > 0 %}+{% endif %}{{ "%.1f"|format(change) }}%</td>
            </tr>
            {% endfor %}
        </table>
    </div>
    {% if insights.category %}
    <div class="insight-box">
        <h4>💡 {{ get_text('ai_insight', lang) }}</h4>
        <p>{{ insights.category }}</p>
    </div>
    {% endif %}
</section>

<!-- Time Analysis -->
<section class="section">
    <h2 class="section-title">🕐 {{ get_text('time_analysis', lang) }}</h2>
    <div class="chart-card"><h3>Heatmap (Day × Hour)</h3><img src="images/heatmap.png"></div>
    <div class="chart-card"><h3>Monthly Trend</h3><img src="images/monthly_trend.png"></div>
    <div class="stat-box" style="margin-top: 20px;">
        <div class="number">{{ r.time.long_tail.count }}</div>
        <div class="label">Long-tail Tickets >48h ({{ "%.1f"|format(r.time.long_tail.percentage) }}%)</div>
    </div>
    {% if insights.time %}
    <div class="insight-box">
        <h4>💡 {{ get_text('ai_insight', lang) }}</h4>
        <p>{{ insights.time }}</p>
    </div>
    {% endif %}
</section>

<!-- SLA Violations Detail -->
<section class="section">
    <h2 class="section-title">⚠️ {{ get_text('sla_violations', lang) }}</h2>
    <div class="chart-card">
        <h3>Violation List ({{ r.violations.total }} total)</h3>
        <table>
            <tr><th>{{ get_text('order_number', lang) }}</th><th>{{ get_text('priority', lang) }}</th><th>{{ get_text('resolver', lang) }}</th><th>{{ get_text('category', lang) }}</th><th>Actual</th><th>{{ get_text('overage', lang) }}</th></tr>
            {% for v in r.violations.details[:15] %}
            <tr>
                <td><strong>{{ v['Order Number'] }}</strong></td>
                <td><span class="priority priority-{{ v['Priority']|lower }}">{{ v['Priority'] }}</span></td>
                <td>{{ v['Resolver'] }}</td>
                <td>{{ v['Category'][:20] }}...</td>
                <td>{{ "%.1f"|format(v['Resolution_Hours']) }}{{ get_text('hours', lang) }}</td>
                <td style="color: var(--danger); font-weight: 600;">+{{ "%.1f"|format(v['SLA_Res_Overage_Hours']) }}{{ get_text('hours', lang) }}</td>
            </tr>
            {% endfor %}
        </table>
        {% if r.violations.total > 15 %}
        <p style="text-align: center; margin-top: 16px; color: var(--text-secondary);">Showing top 15 of {{ r.violations.total }} violations</p>
        {% endif %}
    </div>
</section>

<!-- Recommendations -->
<section class="section">
    <h2 class="section-title">💡 {{ get_text('recommendations', lang) }}</h2>
    {% if r.recommendations %}
    {% for rec in r.recommendations %}
    <div class="rec-item">
        <div class="rec-priority {% if rec.priority == 'HIGH' %}high{% else %}medium{% endif %}"></div>
        <div><strong style="color: var(--primary);">[{{ rec.category }}]</strong> {{ rec.text }}</div>
    </div>
    {% endfor %}
    {% else %}
    <div class="chart-card" style="text-align: center; padding: 48px; color: var(--success);">
        <div style="font-size: 48px;">✅</div>
        <div style="font-size: 18px; margin-top: 16px;">{% if lang == 'zh' %}没有需要关注的紧急问题{% else %}No urgent issues requiring attention{% endif %}</div>
    </div>
    {% endif %}
</section>

<footer>{{ get_text('report_title', lang) }} | {{ get_text('report_generated', lang) }} {{ date }}</footer>

</div>
</body>
</html>
'''
