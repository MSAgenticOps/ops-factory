"""
Report builder for Major Incident Analysis Report.

Generates HTML and DOCX reports from analysis results.
"""
from datetime import datetime
from pathlib import Path
from typing import Optional

from models import IncidentAnalysis, IssueSeverity, ActionItemStatus
import config


# =============================================================================
# Internationalization
# =============================================================================
I18N = {
    "en": {
        "report_title": "Major Incident Analysis Report",
        "incident_overview": "Incident Overview",
        "timeline": "Timeline",
        "time_analysis": "Time Analysis",
        "flow_analysis": "Flow Analysis",
        "issues_detected": "Issues Detected",
        "ai_insights": "AI Insights",
        "root_cause_analysis": "Root Cause & Control Point Analysis",
        "improvement_tracking": "Improvement Action Items",
        "order_number": "Order Number",
        "title": "Title",
        "priority": "Priority",
        "category": "Category",
        "status": "Status",
        "created_at": "Created At",
        "resolved_at": "Resolved At",
        "duration": "Duration",
        "affected_systems": "Affected Systems",
        "response_time": "Response Time",
        "resolution_time": "Resolution Time",
        "sla_status": "SLA Status",
        "met": "Met",
        "violated": "Violated",
        "ongoing": "Ongoing",
        "escalations": "Escalations",
        "reassignments": "Reassignments",
        "participants": "Participants",
        "no_issues": "No issues detected",
        "minutes": "minutes",
        "minute": "minute",
        "hours": "hours",
        "generated_at": "Generated at",
        "severity": "Severity",
        "issue": "Issue",
        "description": "Description",
        "actor": "Actor",
        "event": "Event",
        "from": "From",
        "to": "To",
        "detail": "Detail",
        "phase": "Phase",
        "phase_duration": "Phase Duration",
        "root_cause": "Root Cause",
        "five_whys": "5-Why Analysis",
        "control_failures": "Control Point Failures",
        "contributing_factors": "Contributing Factors",
        "action_item": "Action Item",
        "owner": "Owner",
        "due_date": "Due Date",
        "open": "Open",
        "in_progress": "In Progress",
        "completed": "Completed",
        "blocked": "Blocked",
        "time_breakdown": "Time Breakdown",
        "sla_compliance": "SLA Compliance",
        "escalation_path": "Escalation Path",
    },
    "zh": {
        "report_title": "重大事故分析报告",
        "incident_overview": "事故概览",
        "timeline": "时间线",
        "time_analysis": "时间分析",
        "flow_analysis": "流转分析",
        "issues_detected": "发现的问题",
        "ai_insights": "AI 洞察",
        "root_cause_analysis": "根因与控制点分析",
        "improvement_tracking": "改进项跟踪",
        "order_number": "工单号",
        "title": "标题",
        "priority": "优先级",
        "category": "类别",
        "status": "状态",
        "created_at": "创建时间",
        "resolved_at": "解决时间",
        "duration": "持续时长",
        "affected_systems": "受影响系统",
        "response_time": "响应时间",
        "resolution_time": "解决时间",
        "sla_status": "SLA 状态",
        "met": "达标",
        "violated": "违规",
        "ongoing": "进行中",
        "escalations": "升级次数",
        "reassignments": "重分配次数",
        "participants": "参与人员",
        "no_issues": "未发现问题",
        "minutes": "分钟",
        "minute": "分钟",
        "hours": "小时",
        "generated_at": "生成时间",
        "severity": "严重程度",
        "issue": "问题",
        "description": "描述",
        "actor": "操作人",
        "event": "事件",
        "from": "从",
        "to": "到",
        "detail": "详情",
        "phase": "阶段",
        "phase_duration": "阶段耗时",
        "root_cause": "根本原因",
        "five_whys": "5-Why 分析",
        "control_failures": "控制点失效",
        "contributing_factors": "贡献因素",
        "action_item": "改进项",
        "owner": "负责人",
        "due_date": "截止日期",
        "open": "待处理",
        "in_progress": "进行中",
        "completed": "已完成",
        "blocked": "已阻塞",
        "time_breakdown": "时间分解",
        "sla_compliance": "SLA 合规",
        "escalation_path": "升级路径",
    },
}


def get_text(key: str, language: str = "en") -> str:
    """Get localized text."""
    return I18N.get(language, I18N["en"]).get(key, key)


# =============================================================================
# HTML Report Builder
# =============================================================================
class HTMLReportBuilder:
    """Builds HTML report from analysis results."""

    def __init__(self, analysis: IncidentAnalysis, language: str = "en"):
        self.analysis = analysis
        self.language = language
        self.t = lambda key: get_text(key, language)

    def build(self) -> str:
        """Build complete HTML report."""
        return f"""<!DOCTYPE html>
<html lang="{self.language}">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{self.t('report_title')} - {self.analysis.incident.order_number}</title>
    <style>
{self._get_styles()}
    </style>
</head>
<body>
    <div class="container">
        {self._build_header()}
        {self._build_overview()}
        {self._build_time_visualizations()}
        {self._build_timeline()}
        {self._build_time_analysis()}
        {self._build_flow_analysis()}
        {self._build_root_cause_analysis()}
        {self._build_issues()}
        {self._build_insights()}
        {self._build_action_items()}
        {self._build_footer()}
    </div>
</body>
</html>"""

    def _get_styles(self) -> str:
        """Return CSS styles matching SLA analysis report styling."""
        return """
* { margin: 0; padding: 0; box-sizing: border-box; }
body { font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, 'PingFang SC', 'Microsoft YaHei', sans-serif; background: #f3f4f6; color: #374151; line-height: 1.6; }
.container { max-width: 1200px; margin: 0 auto; padding: 20px; }
.header { background: linear-gradient(135deg, #1e40af, #3b82f6); color: white; padding: 30px; border-radius: 12px; margin-bottom: 24px; }
.header h1 { font-size: 28px; margin-bottom: 8px; }
.header .subtitle { opacity: 0.9; font-size: 16px; }
.card { background: white; border-radius: 12px; padding: 24px; margin-bottom: 20px; box-shadow: 0 1px 3px rgba(0,0,0,0.1); }
.card h2 { font-size: 20px; color: #1e40af; margin-bottom: 16px; padding-bottom: 8px; border-bottom: 2px solid #e5e7eb; }
.card h3 { font-size: 16px; color: #374151; margin: 20px 0 12px; }
.metrics { display: grid; grid-template-columns: repeat(auto-fit, minmax(180px, 1fr)); gap: 16px; margin-bottom: 20px; }
.metric { background: #f9fafb; padding: 16px; border-radius: 8px; text-align: center; }
.metric .value { font-size: 28px; font-weight: bold; color: #1e40af; }
.metric .label { color: #6b7280; font-size: 13px; margin-top: 4px; }
.metric.danger .value { color: #ef4444; }
.metric.warning .value { color: #f59e0b; }
.metric.success .value { color: #10b981; }
table { width: 100%; border-collapse: collapse; margin: 16px 0; font-size: 14px; }
th, td { padding: 10px 12px; text-align: left; border-bottom: 1px solid #e5e7eb; }
th { background: #f9fafb; font-weight: 600; color: #374151; }
tr:hover { background: #f9fafb; }
.badge { display: inline-block; padding: 4px 10px; border-radius: 20px; font-size: 11px; font-weight: 500; }
.badge.high, .badge.critical { background: #fee2e2; color: #991b1b; }
.badge.medium, .badge.severe { background: #fef3c7; color: #92400e; }
.badge.low, .badge.minor { background: #d1fae5; color: #065f46; }
.badge.process { background: #dbeafe; color: #1e40af; }
.badge.resource { background: #ede9fe; color: #6b21a8; }
.badge.external { background: #fef3c7; color: #92400e; }
.recommendation { padding: 16px; border-left: 4px solid #3b82f6; background: #f9fafb; margin: 12px 0; border-radius: 0 8px 8px 0; }
.recommendation.critical { border-left-color: #ef4444; background: #fef2f2; }
.recommendation.high { border-left-color: #f59e0b; background: #fffbeb; }
.recommendation strong { color: #374151; }
.recommendation ul { margin: 8px 0 0 20px; }
.recommendation li { margin: 4px 0; }
.grid-2 { display: grid; grid-template-columns: 1fr 1fr; gap: 20px; }
@media (max-width: 768px) { .grid-2 { grid-template-columns: 1fr; } }
.footnote { font-size: 12px; color: #6b7280; margin-top: 12px; padding-top: 12px; border-top: 1px dashed #e5e7eb; }
.alert { padding: 12px 16px; border-radius: 8px; margin: 16px 0; }
.alert.danger { background: #fef2f2; border: 1px solid #fecaca; color: #991b1b; }
.alert.warning { background: #fffbeb; border: 1px solid #fcd34d; color: #92400e; }
.alert.success { background: #f0fdf4; border: 1px solid #86efac; color: #166534; }
.footer { text-align: center; padding: 20px; color: #6b7280; font-size: 14px; }

/* Timeline Styles */
.timeline-container { position: relative; padding-left: 20px; }
.timeline-container::before { content: ''; position: absolute; left: 6px; top: 0; bottom: 0; width: 2px; background: linear-gradient(to bottom, #3b82f6, #10b981); }
.timeline-event { position: relative; display: flex; gap: 1rem; padding: 12px 0; padding-left: 20px; }
.timeline-event::before { content: ''; position: absolute; left: -17px; top: 16px; width: 10px; height: 10px; border-radius: 50%; background: white; border: 2px solid #3b82f6; }
.timeline-event:last-child::after { content: ''; position: absolute; left: -14px; top: 24px; bottom: 0; width: 2px; background: white; }
.timeline-event.event-resolved::before, .timeline-event.event-closed::before { background: #10b981; border-color: #10b981; }
.timeline-event.event-escalated::before { background: #f59e0b; border-color: #f59e0b; }
.timeline-time { width: 150px; flex-shrink: 0; color: #6b7280; font-size: 13px; }
.timeline-content { flex-grow: 1; }
.timeline-event-type { display: inline-block; padding: 3px 8px; border-radius: 4px; font-size: 11px; font-weight: 600; text-transform: uppercase; margin-right: 8px; }
.event-created { background: #dbeafe; color: #1e40af; }
.event-assigned { background: #d1fae5; color: #065f46; }
.event-escalated { background: #ffedd5; color: #c2410c; }
.event-reassigned { background: #fef3c7; color: #b45309; }
.event-status_change { background: #ede9fe; color: #6d28d9; }
.event-resolved { background: #d1fae5; color: #065f46; }
.event-closed { background: #e5e7eb; color: #374151; }
.event-note { background: #f3f4f6; color: #6b7280; }

/* Issue Cards - Enhanced */
.issue-card { padding: 20px; margin-bottom: 16px; border-radius: 12px; border-left: 5px solid; box-shadow: 0 2px 4px rgba(0,0,0,0.05); transition: all 0.2s ease; position: relative; overflow: hidden; }
.issue-card:hover { box-shadow: 0 4px 12px rgba(0,0,0,0.1); transform: translateY(-2px); }
.issue-card::before { content: ''; position: absolute; top: 0; right: 0; width: 60px; height: 60px; opacity: 0.1; background-size: 32px; background-repeat: no-repeat; background-position: center; }
.issue-critical { background: linear-gradient(135deg, #fef2f2 0%, #fee2e2 100%); border-left-color: #ef4444; }
.issue-critical::before { background-image: url("data:image/svg+xml,%3Csvg xmlns='http://www.w3.org/2000/svg' fill='%23ef4444' viewBox='0 0 24 24'%3E%3Cpath d='M12 2C6.48 2 2 6.48 2 12s4.48 10 10 10 10-4.48 10-10S17.52 2 12 2zm-2 15l-5-5 1.41-1.41L10 14.17l7.59-7.59L19 8l-9 9z'/%3E%3C/svg%3E"); }
.issue-high { background: linear-gradient(135deg, #fff7ed 0%, #ffedd5 100%); border-left-color: #f97316; }
.issue-medium { background: linear-gradient(135deg, #fffbeb 0%, #fef3c7 100%); border-left-color: #eab308; }
.issue-low { background: linear-gradient(135deg, #f0fdf4 0%, #dcfce7 100%); border-left-color: #22c55e; }
.issue-header { display: flex; align-items: center; gap: 12px; margin-bottom: 12px; }
.issue-icon { font-size: 20px; line-height: 1; }
.issue-severity { padding: 4px 10px; border-radius: 6px; font-size: 11px; font-weight: 600; text-transform: uppercase; color: white; letter-spacing: 0.5px; }
.severity-critical { background: linear-gradient(135deg, #ef4444, #dc2626); box-shadow: 0 2px 4px rgba(239,68,68,0.3); }
.severity-high { background: linear-gradient(135deg, #f97316, #ea580c); box-shadow: 0 2px 4px rgba(249,115,22,0.3); }
.severity-medium { background: linear-gradient(135deg, #eab308, #ca8a04); box-shadow: 0 2px 4px rgba(234,179,8,0.3); }
.severity-low { background: linear-gradient(135deg, #22c55e, #16a34a); box-shadow: 0 2px 4px rgba(34,197,94,0.3); }
.issue-title { font-weight: 600; font-size: 15px; color: #1f2937; }
.issue-description { color: #4b5563; font-size: 14px; line-height: 1.5; margin-top: 4px; }
.issue-meta { margin-top: 12px; padding-top: 12px; border-top: 1px dashed rgba(0,0,0,0.1); font-size: 13px; color: #6b7280; display: flex; align-items: center; gap: 6px; }
.issue-meta-icon { opacity: 0.6; }

/* Insights Box */
.insights-box { background: linear-gradient(135deg, #eff6ff 0%, #f0fdf4 100%); padding: 20px; border-radius: 8px; border: 1px solid #e5e7eb; }
.insight-item { margin-bottom: 16px; }
.insight-item:last-child { margin-bottom: 0; }

/* Root Cause Analysis */
.root-cause-box { background: linear-gradient(135deg, #fef3c7 0%, #ffedd5 100%); padding: 20px; border-radius: 8px; border: 1px solid #fcd34d; margin-bottom: 16px; }
.five-whys { counter-reset: why-counter; padding-left: 56px; margin-left: 0; }
.five-whys li { list-style: none; padding: 12px 0 12px 16px; position: relative; border-left: 2px solid #f59e0b; }
.five-whys li::before { counter-increment: why-counter; content: "Why " counter(why-counter); position: absolute; left: -56px; top: 12px; background: #f59e0b; color: white; padding: 2px 8px; border-radius: 4px; font-size: 11px; font-weight: 600; white-space: nowrap; }
.control-failures { background: #fef2f2; padding: 16px; border-radius: 8px; border-left: 4px solid #ef4444; }
.control-failures ul { margin: 8px 0 0 20px; }

/* Action Items */
.action-item { display: grid; grid-template-columns: auto 1fr auto auto; gap: 16px; padding: 16px; background: #f9fafb; border-radius: 8px; margin-bottom: 8px; align-items: center; }
.action-item-id { font-family: monospace; background: #1e40af; color: white; padding: 4px 8px; border-radius: 4px; font-size: 12px; }
.action-item-content { flex-grow: 1; }
.action-item-title { font-weight: 600; color: #374151; }
.action-item-meta { font-size: 13px; color: #6b7280; margin-top: 4px; }
.status-badge { padding: 4px 10px; border-radius: 12px; font-size: 11px; font-weight: 600; text-transform: uppercase; }
.status-open { background: #dbeafe; color: #1e40af; }
.status-in_progress { background: #fef3c7; color: #b45309; }
.status-completed { background: #d1fae5; color: #065f46; }
.status-blocked { background: #fef2f2; color: #dc2626; }
.priority-badge { padding: 3px 8px; border-radius: 4px; font-size: 11px; font-weight: 600; text-transform: uppercase; }
.priority-high { background: #fef2f2; color: #dc2626; }
.priority-medium { background: #fef3c7; color: #b45309; }
.priority-low { background: #d1fae5; color: #065f46; }

/* Visualization Cards */
.viz-container { display: grid; grid-template-columns: 1fr 1fr; gap: 20px; margin-bottom: 16px; }
@media (max-width: 768px) { .viz-container { grid-template-columns: 1fr; } }
.viz-card { background: #f9fafb; border-radius: 8px; padding: 16px; border: 1px solid #e5e7eb; }
.viz-title { font-size: 14px; font-weight: 600; color: #374151; margin-bottom: 12px; text-align: center; }
.gantt-chart { width: 100%; overflow-x: auto; }
.gantt-row { display: flex; align-items: center; margin: 6px 0; gap: 8px; }
.gantt-bar { height: 28px; border-radius: 4px; display: flex; align-items: center; padding: 0 10px; font-size: 12px; color: white; font-weight: 500; min-width: fit-content; white-space: nowrap; }
.gantt-label { font-size: 12px; color: #374151; font-weight: 500; white-space: nowrap; }
.sla-gauge { display: flex; flex-direction: column; align-items: center; padding: 16px; }
.gauge-circle { position: relative; width: 120px; height: 120px; }
.gauge-circle svg { transform: rotate(-90deg); }
.gauge-text { position: absolute; top: 50%; left: 50%; transform: translate(-50%, -50%); text-align: center; }
.gauge-value { font-size: 24px; font-weight: 700; }
.gauge-label { font-size: 12px; color: #6b7280; }

/* Escalation Path */
.escalation-path { display: flex; align-items: center; justify-content: center; flex-wrap: wrap; gap: 8px; padding: 16px; }
.escalation-node { background: #f9fafb; padding: 8px 16px; border-radius: 8px; border: 2px solid #e5e7eb; font-size: 13px; text-align: center; }
.escalation-node.active { border-color: #3b82f6; background: #dbeafe; }
.escalation-arrow { color: #6b7280; font-size: 20px; }

/* Tags */
.tag { display: inline-block; padding: 3px 8px; background: #f3f4f6; border-radius: 4px; font-size: 13px; margin-right: 4px; }

@media print {
    body { background: white; }
    .container { max-width: none; padding: 16px; }
    .card { box-shadow: none; border: 1px solid #e5e7eb; page-break-inside: avoid; }
    .header { background: #1e40af !important; -webkit-print-color-adjust: exact; print-color-adjust: exact; }
}
"""

    def _build_header(self) -> str:
        """Build report header."""
        incident = self.analysis.incident
        title = incident.get_title(self.language)
        return f"""
<div class="header">
    <h1>{self.t('report_title')}</h1>
    <div class="subtitle">{incident.order_number} - {title}</div>
</div>
"""

    def _build_overview(self) -> str:
        """Build incident overview section."""
        incident = self.analysis.incident

        # Format duration
        if self.analysis.resolution_time_minutes:
            hours = self.analysis.resolution_time_minutes / 60
            duration = f"{hours:.1f} {self.t('hours')}"
        else:
            duration = self.t('ongoing')

        # Format affected systems (with i18n)
        systems_list = incident.get_affected_systems(self.language)
        systems = ", ".join(systems_list) if systems_list else "-"

        # Get title with i18n
        title = incident.get_title(self.language)

        # SLA status
        if incident.sla_response_met is None:
            response_sla = "-"
            response_class = ""
        elif incident.sla_response_met:
            response_sla = self.t('met')
            response_class = "success"
        else:
            response_sla = self.t('violated')
            response_class = "danger"

        if incident.sla_resolution_met is None:
            resolution_sla = "-"
            resolution_class = ""
        elif incident.sla_resolution_met:
            resolution_sla = self.t('met')
            resolution_class = "success"
        else:
            resolution_sla = self.t('violated')
            resolution_class = "danger"

        return f"""
<div class="card">
    <h2>{self.t('incident_overview')}</h2>
    <div class="metrics">
        <div class="metric">
            <div class="value">{incident.priority}</div>
            <div class="label">{self.t('priority')}</div>
        </div>
        <div class="metric">
            <div class="value">{incident.category}</div>
            <div class="label">{self.t('category')}</div>
        </div>
        <div class="metric {'danger' if incident.is_ongoing else 'success'}">
            <div class="value">{incident.status}</div>
            <div class="label">{self.t('status')}</div>
        </div>
        <div class="metric">
            <div class="value">{duration}</div>
            <div class="label">{self.t('duration')}</div>
        </div>
        <div class="metric {response_class}">
            <div class="value">{response_sla}</div>
            <div class="label">{self.t('response_time')} SLA</div>
        </div>
    </div>
    <p style="margin-top: 16px;"><strong>{self.t('title')}:</strong> {title}</p>
    <p style="margin-top: 8px;"><strong>{self.t('affected_systems')}:</strong> {systems}</p>
</div>
"""

    def _build_time_visualizations(self) -> str:
        """Build time analysis visualizations (Gantt chart, SLA gauge, escalation path)."""
        return f"""
<div class="card">
    <h2>{self.t('time_breakdown')}</h2>
    <div class="viz-container">
        {self._build_gantt_chart()}
        {self._build_sla_gauge()}
    </div>
    {self._build_escalation_path()}
</div>
"""

    def _build_gantt_chart(self) -> str:
        """Build phase duration Gantt chart."""
        phases = self.analysis.phase_durations
        if not phases:
            return ""

        total_time = sum(phases.values()) or 1
        colors = {
            "Open": "#3b82f6",
            "In Progress": "#f59e0b",
            "resolved": "#10b981",
            "Resolved": "#10b981",
            "closed": "#6b7280",
            "Closed": "#6b7280",
        }

        bars_html = ""
        for phase, duration in phases.items():
            width_pct = (duration / total_time) * 100
            color = colors.get(phase, "#8b5cf6")
            # Capitalize phase name consistently
            display_phase = phase.capitalize() if phase.islower() else phase
            minutes_text = self.t('minute') if duration == 1 else self.t('minutes')
            label = f"{display_phase}: {duration} {minutes_text}"

            # For bars that would be too narrow, show label outside
            if width_pct < 25:
                bars_html += f"""
        <div class="gantt-row">
            <div class="gantt-bar" style="width: {max(width_pct, 8)}%; background: {color};"></div>
            <span class="gantt-label">{label}</span>
        </div>"""
            else:
                bars_html += f"""
        <div class="gantt-row">
            <div class="gantt-bar" style="width: {width_pct}%; background: {color};">{label}</div>
        </div>"""

        return f"""
        <div class="viz-card">
            <div class="viz-title">{self.t('phase_duration')}</div>
            <div class="gantt-chart">
                {bars_html}
            </div>
        </div>"""

    def _build_sla_gauge(self) -> str:
        """Build SLA compliance gauge."""
        incident = self.analysis.incident
        sla = incident.sla

        if not sla or not self.analysis.resolution_time_minutes:
            return ""

        # Calculate percentage (100% = on target, <100% = under, >100% = over)
        target_minutes = sla.resolution_hours * 60
        actual_minutes = self.analysis.resolution_time_minutes
        percentage = (actual_minutes / target_minutes) * 100

        # Determine color based on compliance
        if percentage <= 100:
            color = "#10b981"  # Green
            status = self.t('met')
        elif percentage <= 150:
            color = "#f59e0b"  # Warning
            status = self.t('violated')
        else:
            color = "#ef4444"  # Danger
            status = self.t('violated')

        # SVG gauge - show percentage of SLA used
        display_pct = min(percentage, 200)  # Cap at 200% for display
        stroke_dasharray = (display_pct / 200) * 283  # 283 is circumference of r=45

        return f"""
        <div class="viz-card">
            <div class="viz-title">{self.t('sla_compliance')}</div>
            <div class="sla-gauge">
                <div class="gauge-circle">
                    <svg width="120" height="120" viewBox="0 0 120 120">
                        <circle cx="60" cy="60" r="45" fill="none" stroke="#e5e7eb" stroke-width="10"/>
                        <circle cx="60" cy="60" r="45" fill="none" stroke="{color}" stroke-width="10"
                                stroke-dasharray="{stroke_dasharray} 283" stroke-linecap="round"/>
                    </svg>
                    <div class="gauge-text">
                        <div class="gauge-value" style="color: {color};">{percentage:.0f}%</div>
                        <div class="gauge-label">{status}</div>
                    </div>
                </div>
                <div style="margin-top: 8px; font-size: 13px; color: #6b7280;">
                    {actual_minutes} / {target_minutes:.0f} {self.t('minutes')}
                </div>
            </div>
        </div>"""

    def _build_escalation_path(self) -> str:
        """Build escalation path visualization."""
        escalations = [e for e in self.analysis.timeline if e.event.value == "escalated"]

        if not escalations:
            return ""

        nodes_html = ""
        for i, event in enumerate(escalations):
            from_val = event.get_from_value(self.language) or "Initial"
            to_val = event.get_to_value(self.language) or "Unknown"

            if i == 0:
                nodes_html += f'<div class="escalation-node">{from_val}</div>'

            nodes_html += '<span class="escalation-arrow">→</span>'
            nodes_html += f'<div class="escalation-node active">{to_val}</div>'

        return f"""
    <h3 style="margin-top: 1rem;">{self.t('escalation_path')}</h3>
    <div class="escalation-path">
        {nodes_html}
    </div>"""

    def _build_timeline(self) -> str:
        """Build timeline section with i18n support and vertical connecting line."""
        events_html = ""
        for event in self.analysis.timeline:
            event_class = f"event-{event.event.value}"
            time_str = event.timestamp.strftime("%Y-%m-%d %H:%M")

            # Get i18n values
            actor = event.get_actor(self.language)
            from_val = event.get_from_value(self.language)
            to_val = event.get_to_value(self.language)
            detail = event.get_detail(self.language)

            # Build detail string
            detail_parts = []
            if from_val and to_val:
                detail_parts.append(f"{from_val} → {to_val}")
            elif to_val:
                detail_parts.append(f"→ {to_val}")
            if detail:
                detail_parts.append(detail)
            detail_str = " | ".join(detail_parts) if detail_parts else ""

            # Duration indicator
            duration_str = ""
            if event.duration_from_previous and event.duration_from_previous > 0:
                duration_str = f"<span class='tag'>+{event.duration_from_previous}m</span>"

            events_html += f"""
<div class="timeline-event {event_class}">
    <div class="timeline-time">{time_str} {duration_str}</div>
    <div class="timeline-content">
        <span class="timeline-event-type {event_class}">{event.event.value}</span>
        <strong>{actor}</strong>
        {f'<div style="color: #6b7280; margin-top: 4px;">{detail_str}</div>' if detail_str else ''}
    </div>
</div>
"""

        return f"""
<div class="card">
    <h2>{self.t('timeline')}</h2>
    <div class="timeline-container">
        {events_html}
    </div>
</div>
"""

    def _build_time_analysis(self) -> str:
        """Build time analysis section."""
        # Response time with correct singular/plural
        response = self.analysis.response_time_minutes
        if response:
            minutes_word = self.t('minute') if response == 1 else self.t('minutes')
            response_str = f"{response} {minutes_word}"
        else:
            response_str = "-"

        # Resolution time
        resolution = self.analysis.resolution_time_minutes
        if resolution:
            resolution_str = f"{resolution / 60:.1f} {self.t('hours')} ({resolution} {self.t('minutes')})"
        else:
            resolution_str = self.t('ongoing')

        # Phase durations with capitalization fix
        phases_html = ""
        for phase, duration in self.analysis.phase_durations.items():
            # Capitalize phase names consistently
            display_phase = phase.capitalize() if phase.islower() else phase
            minutes_word = self.t('minute') if duration == 1 else self.t('minutes')
            phases_html += f"<tr><td>{display_phase}</td><td>{duration} {minutes_word}</td></tr>"

        return f"""
<div class="card">
    <h2>{self.t('time_analysis')}</h2>
    <div class="metrics" style="margin-bottom: 16px;">
        <div class="metric">
            <div class="value">{response_str}</div>
            <div class="label">{self.t('response_time')}</div>
        </div>
        <div class="metric">
            <div class="value">{resolution_str}</div>
            <div class="label">{self.t('resolution_time')}</div>
        </div>
    </div>
    {f'''
    <h3>{self.t('phase_duration')}</h3>
    <table>
        <thead><tr><th>{self.t('phase')}</th><th>{self.t('duration')}</th></tr></thead>
        <tbody>{phases_html}</tbody>
    </table>
    ''' if phases_html else ''}
</div>
"""

    def _build_flow_analysis(self) -> str:
        """Build flow analysis section."""
        # Get actors with proper i18n support
        actors_set = set()
        for event in self.analysis.timeline:
            actor = event.get_actor(self.language)
            actors_set.add(actor)
            to_val = event.get_to_value(self.language)
            if to_val:
                actors_set.add(to_val)

        # Filter out non-person values
        exclude_patterns = {
            "In Progress", "Open", "Resolved", "Closed", "Pending",
            "进行中", "已解决", "已关闭", "待处理",
        }

        # Normalize and deduplicate names
        normalized_map = {}  # base_name -> full_name
        for actor in actors_set:
            if not actor or actor in exclude_patterns:
                continue

            # Extract base name (remove annotations in parentheses)
            base_name = actor
            if "(" in actor:
                base_name = actor.split("(")[0].strip()
            elif "（" in actor:
                base_name = actor.split("（")[0].strip()

            # Keep the longer (more detailed) version
            if base_name not in normalized_map:
                normalized_map[base_name] = actor
            elif len(actor) > len(normalized_map[base_name]):
                normalized_map[base_name] = actor

        actors_list = sorted(normalized_map.values())
        actors = ", ".join(actors_list) if actors_list else "-"

        return f"""
<div class="card">
    <h2>{self.t('flow_analysis')}</h2>
    <div class="metrics">
        <div class="metric {'warning' if self.analysis.escalation_count >= 2 else ''}">
            <div class="value">{self.analysis.escalation_count}</div>
            <div class="label">{self.t('escalations')}</div>
        </div>
        <div class="metric {'warning' if self.analysis.reassignment_count >= 3 else ''}">
            <div class="value">{self.analysis.reassignment_count}</div>
            <div class="label">{self.t('reassignments')}</div>
        </div>
        <div class="metric">
            <div class="value">{len(actors_list)}</div>
            <div class="label">{self.t('participants')}</div>
        </div>
    </div>
    <p style="margin-top: 16px;"><strong>{self.t('participants')}:</strong> {actors}</p>
</div>
"""

    def _build_root_cause_analysis(self) -> str:
        """Build root cause and control point analysis section."""
        rca = self.analysis.root_cause_analysis

        if not rca:
            return ""

        # Root cause description
        description = rca.description_en if self.language == "en" and rca.description_en else rca.description

        # 5-Whys analysis
        five_whys_html = ""
        if rca.five_whys:
            five_whys_items = "".join(f"<li>{why}</li>" for why in rca.five_whys)
            five_whys_html = f"""
        <h3>{self.t('five_whys')}</h3>
        <ol class="five-whys">
            {five_whys_items}
        </ol>
"""

        # Control point failures
        control_html = ""
        if rca.control_point_failures:
            control_items = "".join(f"<li>{item}</li>" for item in rca.control_point_failures)
            control_html = f"""
        <h3>{self.t('control_failures')}</h3>
        <div class="control-failures">
            <ul>
                {control_items}
            </ul>
        </div>
"""

        # Contributing factors
        factors_html = ""
        if rca.contributing_factors:
            factors_items = "".join(f"<li>{item}</li>" for item in rca.contributing_factors)
            factors_html = f"""
        <h3>{self.t('contributing_factors')}</h3>
        <ul>
            {factors_items}
        </ul>
"""

        return f"""
<div class="card">
    <h2>{self.t('root_cause_analysis')}</h2>
    <div class="root-cause-box">
        <strong>{self.t('root_cause')} ({rca.category}):</strong>
        <p style="margin-top: 8px;">{description}</p>
    </div>
    {five_whys_html}
    {control_html}
    {factors_html}
</div>
"""

    def _translate_actor_name(self, actor: str) -> str:
        """Translate actor name to the report language using timeline data."""
        if not actor:
            return actor

        # In Chinese mode, return as-is
        if self.language == "zh":
            return actor

        # Build a mapping from Chinese to English names from timeline
        for event in self.analysis.timeline:
            # Check if this actor matches any Chinese name
            if event.actor == actor and event.actor_en:
                return event.actor_en
            if event.to_value == actor:
                if event.to_value_en:
                    return event.to_value_en
            if event.from_value == actor:
                if event.from_value_en:
                    return event.from_value_en

        # If no translation found, return original
        return actor

    def _build_issues(self) -> str:
        """Build issues section with enhanced card styling."""
        if not self.analysis.issues:
            return f"""
<div class="card">
    <h2>{self.t('issues_detected')}</h2>
    <p style="color: #10b981; display: flex; align-items: center; gap: 8px;">✅ {self.t('no_issues')}</p>
</div>
"""

        # Severity icons mapping
        severity_icons = {
            "critical": "🚨",
            "high": "⚠️",
            "medium": "📋",
            "low": "💡",
        }
        
        # Actor label based on language
        actor_label = "操作人" if self.language == "zh" else "Actor"

        issues_html = ""
        for issue in self.analysis.issues:
            severity_class = f"severity-{issue.severity.value}"
            issue_class = f"issue-{issue.severity.value}"
            icon = severity_icons.get(issue.severity.value, "📌")

            # Translate actor name if in English mode
            actor_display = self._translate_actor_name(issue.actor) if issue.actor else None

            issues_html += f"""
<div class="issue-card {issue_class}">
    <div class="issue-header">
        <span class="issue-icon">{icon}</span>
        <span class="issue-severity {severity_class}">{issue.severity.value}</span>
        <span class="issue-title">{issue.title}</span>
    </div>
    <div class="issue-description">{issue.description}</div>
    {f'<div class="issue-meta"><span class="issue-meta-icon">👤</span>{actor_label}: {actor_display}</div>' if actor_display else ''}
</div>
"""

        return f"""
<div class="card">
    <h2>{self.t('issues_detected')} ({len(self.analysis.issues)})</h2>
    {issues_html}
</div>
"""

    def _build_insights(self) -> str:
        """Build AI insights section. Shows placeholder when no insights available."""
        if not self.analysis.insights:
            # Show a placeholder message when AI insights are not available
            no_insights_msg = "暂无 AI 分析。请使用 LLM 模式生成（移除 --no-llm 参数）" if self.language == "zh" else "No AI insights available. Run with LLM enabled (remove --no-llm flag) to generate AI analysis."
            return f"""
<div class="card">
    <h2>{self.t('ai_insights')}</h2>
    <div class="insights-box" style="text-align: center; padding: 32px;">
        <div style="font-size: 48px; margin-bottom: 16px;">🤖</div>
        <p style="color: #6b7280;">{no_insights_msg}</p>
    </div>
</div>
"""

        insights = self.analysis.insights

        # Build structured insights HTML
        sections_html = ""

        # Overall assessment
        if "overall_assessment" in insights:
            label = "整体评估" if self.language == "zh" else "Overall Assessment"
            sections_html += f"""
<div class="insight-item">
    <strong>{label}:</strong>
    <p>{insights['overall_assessment']}</p>
</div>
"""

        # Highlights
        if "highlights" in insights and insights["highlights"]:
            label = "亮点" if self.language == "zh" else "Highlights"
            items = "".join(f"<li>{h}</li>" for h in insights["highlights"])
            sections_html += f"""
<div class="insight-item">
    <strong>{label}:</strong>
    <ul>{items}</ul>
</div>
"""

        # Problem analysis
        if "problem_analysis" in insights:
            label = "问题分析" if self.language == "zh" else "Problem Analysis"
            sections_html += f"""
<div class="insight-item">
    <strong>{label}:</strong>
    <p>{insights['problem_analysis']}</p>
</div>
"""

        # Improvement suggestions
        if "improvement_suggestions" in insights and insights["improvement_suggestions"]:
            label = "改进建议" if self.language == "zh" else "Improvement Suggestions"
            items = "".join(f"<li>{s}</li>" for s in insights["improvement_suggestions"])
            sections_html += f"""
<div class="insight-item">
    <strong>{label}:</strong>
    <ul>{items}</ul>
</div>
"""

        # Prevention measures
        if "prevention_measures" in insights and insights["prevention_measures"]:
            label = "预防措施" if self.language == "zh" else "Prevention Measures"
            items = "".join(f"<li>{m}</li>" for m in insights["prevention_measures"])
            sections_html += f"""
<div class="insight-item">
    <strong>{label}:</strong>
    <ul>{items}</ul>
</div>
"""

        # Fallback for old format
        if not sections_html:
            for key, value in insights.items():
                if isinstance(value, list):
                    items = "".join(f"<li>{v}</li>" for v in value)
                    sections_html += f"""
<div class="insight-item">
    <strong>{key}:</strong>
    <ul>{items}</ul>
</div>
"""
                else:
                    sections_html += f"""
<div class="insight-item">
    <strong>{key}:</strong>
    <p>{value}</p>
</div>
"""

        return f"""
<div class="card">
    <h2>{self.t('ai_insights')}</h2>
    <div class="insights-box">
        {sections_html}
    </div>
</div>
"""

    def _build_action_items(self) -> str:
        """Build improvement action items tracking section."""
        action_items = self.analysis.action_items

        if not action_items:
            return ""

        items_html = ""
        for item in action_items:
            status_class = f"status-{item.status.value}"
            priority_class = f"priority-{item.priority}"

            # Get localized status label
            status_labels = {
                "open": self.t('open'),
                "in_progress": self.t('in_progress'),
                "completed": self.t('completed'),
                "blocked": self.t('blocked'),
            }
            status_label = status_labels.get(item.status.value, item.status.value)

            items_html += f"""
<div class="action-item">
    <span class="action-item-id">{item.id}</span>
    <div class="action-item-content">
        <div class="action-item-title">{item.title}</div>
        <div class="action-item-meta">{self.t('owner')}: {item.owner} | {item.category.capitalize()}</div>
    </div>
    <span class="priority-badge {priority_class}">{item.priority}</span>
    <span class="status-badge {status_class}">{status_label}</span>
</div>
"""

        return f"""
<div class="card">
    <h2>{self.t('improvement_tracking')}</h2>
    {items_html}
</div>
"""

    def _build_footer(self) -> str:
        """Build report footer."""
        now = datetime.now().strftime("%Y-%m-%d %H:%M")
        return f"""
<div class="footer">
    {self.t('generated_at')}: {now} | Major Incident Analysis Report
</div>
"""

    def save(self, filepath: Path) -> None:
        """Save HTML report to file."""
        html = self.build()
        with open(filepath, "w", encoding="utf-8") as f:
            f.write(html)


# =============================================================================
# DOCX Report Builder
# =============================================================================
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


class DOCXReportBuilder:
    """Builds DOCX report from analysis results."""

    def __init__(self, analysis: IncidentAnalysis, language: str = "en"):
        self.analysis = analysis
        self.language = language
        self.t = lambda key: get_text(key, language)
        self.doc = None

    def build(self) -> Optional["Document"]:
        """Build DOCX document."""
        if not HAS_DOCX:
            print("Warning: python-docx not installed, skipping DOCX generation")
            return None

        self.doc = Document()

        # Set document styles
        self._setup_styles()

        # Build sections
        self._build_header()
        self._build_overview()
        self._build_timeline()
        self._build_time_analysis()
        self._build_flow_analysis()
        self._build_issues()
        self._build_insights()
        self._build_footer()

        return self.doc

    def _setup_styles(self) -> None:
        """Setup document styles."""
        # Set default font
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(11)

        # Set Chinese font
        style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    def _add_heading(self, text: str, level: int = 1) -> None:
        """Add a heading with custom styling."""
        heading = self.doc.add_heading(text, level=level)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(30, 64, 175)  # Primary blue #1e40af

    def _add_table(self, data: list[list[str]], headers: bool = True) -> None:
        """Add a table with data."""
        if not data:
            return

        table = self.doc.add_table(rows=len(data), cols=len(data[0]))
        table.style = 'Table Grid'

        for i, row_data in enumerate(data):
            row = table.rows[i]
            for j, cell_text in enumerate(row_data):
                cell = row.cells[j]
                cell.text = str(cell_text)
                if headers and i == 0:
                    # Bold header row
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True

        self.doc.add_paragraph()

    def _build_header(self) -> None:
        """Build document header."""
        incident = self.analysis.incident

        # Title
        title = self.doc.add_heading(self.t('report_title'), level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Subtitle
        subtitle = self.doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run(f"{incident.order_number} - {incident.title}")
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(107, 114, 128)

        self.doc.add_paragraph()

    def _build_overview(self) -> None:
        """Build incident overview section."""
        incident = self.analysis.incident

        self._add_heading(self.t('incident_overview'), level=1)

        # Format duration
        if self.analysis.resolution_time_minutes:
            hours = self.analysis.resolution_time_minutes / 60
            duration = f"{hours:.1f} {self.t('hours')}"
        else:
            duration = self.t('ongoing')

        # SLA status
        response_sla = self.t('met') if incident.sla_response_met else (
            self.t('violated') if incident.sla_response_met is False else "-"
        )
        resolution_sla = self.t('met') if incident.sla_resolution_met else (
            self.t('violated') if incident.sla_resolution_met is False else "-"
        )

        # Build overview table (Order Number removed - already shown in header)
        data = [
            [self.t('title'), incident.title],
            [self.t('priority'), incident.priority],
            [self.t('category'), incident.category],
            [self.t('status'), incident.status],
            [self.t('created_at'), incident.created_at.strftime("%Y-%m-%d %H:%M")],
            [self.t('resolved_at'), incident.resolved_at.strftime("%Y-%m-%d %H:%M") if incident.resolved_at else "-"],
            [self.t('duration'), duration],
            [f"{self.t('response_time')} SLA", response_sla],
            [f"{self.t('resolution_time')} SLA", resolution_sla],
            [self.t('affected_systems'), ", ".join(incident.affected_systems) if incident.affected_systems else "-"],
        ]

        for label, value in data:
            p = self.doc.add_paragraph()
            p.add_run(f"{label}: ").bold = True
            p.add_run(str(value))

        self.doc.add_paragraph()

    def _build_timeline(self) -> None:
        """Build timeline section."""
        self._add_heading(self.t('timeline'), level=1)

        # Timeline table
        headers = [self.t('actor'), self.t('event'), self.t('detail')]
        data = [headers]

        for event in self.analysis.timeline:
            time_str = event.timestamp.strftime("%m-%d %H:%M")

            # Build detail string
            detail_parts = []
            if event.from_value and event.to_value:
                detail_parts.append(f"{event.from_value} → {event.to_value}")
            elif event.to_value:
                detail_parts.append(f"→ {event.to_value}")
            if event.detail:
                detail_parts.append(event.detail)
            detail_str = " | ".join(detail_parts) if detail_parts else ""

            # Duration
            duration_str = f" (+{event.duration_from_previous}m)" if event.duration_from_previous else ""

            data.append([
                f"[{time_str}]{duration_str}",
                f"{event.event.value} - {event.actor}",
                detail_str[:100] + "..." if len(detail_str) > 100 else detail_str
            ])

        self._add_table(data)

    def _build_time_analysis(self) -> None:
        """Build time analysis section."""
        self._add_heading(self.t('time_analysis'), level=1)

        # Response time
        response = self.analysis.response_time_minutes
        response_str = f"{response} {self.t('minutes')}" if response else "-"

        # Resolution time
        resolution = self.analysis.resolution_time_minutes
        if resolution:
            resolution_str = f"{resolution / 60:.1f} {self.t('hours')} ({resolution} {self.t('minutes')})"
        else:
            resolution_str = self.t('ongoing')

        p = self.doc.add_paragraph()
        p.add_run(f"{self.t('response_time')}: ").bold = True
        p.add_run(response_str)

        p = self.doc.add_paragraph()
        p.add_run(f"{self.t('resolution_time')}: ").bold = True
        p.add_run(resolution_str)

        # Phase durations
        if self.analysis.phase_durations:
            self.doc.add_paragraph()
            p = self.doc.add_paragraph()
            p.add_run(f"{self.t('phase_duration')}:").bold = True

            data = [[self.t('phase'), self.t('duration')]]
            for phase, duration in self.analysis.phase_durations.items():
                data.append([phase, f"{duration} {self.t('minutes')}"])
            self._add_table(data)

    def _build_flow_analysis(self) -> None:
        """Build flow analysis section."""
        self._add_heading(self.t('flow_analysis'), level=1)

        p = self.doc.add_paragraph()
        p.add_run(f"{self.t('escalations')}: ").bold = True
        p.add_run(str(self.analysis.escalation_count))

        p = self.doc.add_paragraph()
        p.add_run(f"{self.t('reassignments')}: ").bold = True
        p.add_run(str(self.analysis.reassignment_count))

        p = self.doc.add_paragraph()
        p.add_run(f"{self.t('participants')}: ").bold = True
        p.add_run(", ".join(self.analysis.unique_actors) if self.analysis.unique_actors else "-")

        self.doc.add_paragraph()

    def _build_issues(self) -> None:
        """Build issues section."""
        self._add_heading(self.t('issues_detected'), level=1)

        if not self.analysis.issues:
            p = self.doc.add_paragraph()
            run = p.add_run(self.t('no_issues'))
            run.font.color.rgb = RGBColor(34, 197, 94)  # Green
            return

        for issue in self.analysis.issues:
            p = self.doc.add_paragraph()

            # Severity badge
            severity_colors = {
                "critical": RGBColor(239, 68, 68),
                "high": RGBColor(249, 115, 22),
                "medium": RGBColor(234, 179, 8),
                "low": RGBColor(34, 197, 94),
            }

            run = p.add_run(f"[{issue.severity.value.upper()}] ")
            run.bold = True
            run.font.color.rgb = severity_colors.get(issue.severity.value, RGBColor(0, 0, 0))

            p.add_run(f"{issue.title}: ").bold = True
            p.add_run(issue.description)

        self.doc.add_paragraph()

    def _build_insights(self) -> None:
        """Build AI insights section. Shows placeholder when no insights available."""
        self._add_heading(self.t('ai_insights'), level=1)

        if not self.analysis.insights:
            # Show a placeholder message when AI insights are not available
            no_insights_msg = "暂无 AI 分析。请使用 LLM 模式生成（移除 --no-llm 参数）" if self.language == "zh" else "No AI insights available. Run with LLM enabled (remove --no-llm flag) to generate AI analysis."
            p = self.doc.add_paragraph()
            run = p.add_run(no_insights_msg)
            run.font.color.rgb = RGBColor(107, 114, 128)
            run.font.italic = True
            self.doc.add_paragraph()
            return

        insights = self.analysis.insights

        # Overall assessment
        if "overall_assessment" in insights:
            label = "整体评估" if self.language == "zh" else "Overall Assessment"
            p = self.doc.add_paragraph()
            p.add_run(f"{label}: ").bold = True
            p.add_run(insights['overall_assessment'])
            self.doc.add_paragraph()

        # Highlights
        if "highlights" in insights and insights["highlights"]:
            label = "亮点" if self.language == "zh" else "Highlights"
            p = self.doc.add_paragraph()
            p.add_run(f"{label}:").bold = True
            for h in insights["highlights"]:
                self.doc.add_paragraph(f"• {h}", style='List Bullet')
            self.doc.add_paragraph()

        # Problem analysis
        if "problem_analysis" in insights:
            label = "问题分析" if self.language == "zh" else "Problem Analysis"
            p = self.doc.add_paragraph()
            p.add_run(f"{label}: ").bold = True
            p.add_run(insights['problem_analysis'])
            self.doc.add_paragraph()

        # Improvement suggestions
        if "improvement_suggestions" in insights and insights["improvement_suggestions"]:
            label = "改进建议" if self.language == "zh" else "Improvement Suggestions"
            p = self.doc.add_paragraph()
            p.add_run(f"{label}:").bold = True
            for s in insights["improvement_suggestions"]:
                self.doc.add_paragraph(f"• {s}", style='List Bullet')
            self.doc.add_paragraph()

        # Prevention measures
        if "prevention_measures" in insights and insights["prevention_measures"]:
            label = "预防措施" if self.language == "zh" else "Prevention Measures"
            p = self.doc.add_paragraph()
            p.add_run(f"{label}:").bold = True
            for m in insights["prevention_measures"]:
                self.doc.add_paragraph(f"• {m}", style='List Bullet')

    def _build_footer(self) -> None:
        """Build document footer."""
        self.doc.add_paragraph()
        now = datetime.now().strftime("%Y-%m-%d %H:%M")
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{self.t('generated_at')}: {now} | Major Incident Analysis Report")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(107, 114, 128)

    def save(self, filepath: Path) -> None:
        """Save DOCX report to file."""
        if self.doc:
            self.doc.save(filepath)


# =============================================================================
# Report Generation Functions
# =============================================================================
def generate_html_report(
    analysis: IncidentAnalysis,
    output_dir: Path,
    language: str = "en"
) -> Path:
    """Generate HTML report and return file path."""
    builder = HTMLReportBuilder(analysis, language)

    filename = f"Major_Incident_Analysis_{analysis.incident.order_number}_{language.upper()}.html"
    filepath = output_dir / filename

    builder.save(filepath)
    return filepath


def generate_docx_report(
    analysis: IncidentAnalysis,
    output_dir: Path,
    language: str = "en"
) -> Optional[Path]:
    """Generate DOCX report and return file path."""
    if not HAS_DOCX:
        print("Warning: python-docx not installed, skipping DOCX generation")
        return None

    builder = DOCXReportBuilder(analysis, language)
    doc = builder.build()

    if not doc:
        return None

    filename = f"Major_Incident_Analysis_{analysis.incident.order_number}_{language.upper()}.docx"
    filepath = output_dir / filename

    builder.save(filepath)
    return filepath
