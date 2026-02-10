"""
Report Builder for BO Workload Performance Report.

Generates DOCX and HTML reports from analysis results.
"""

import os
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Any

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import base64
from io import BytesIO

from jinja2 import Template
import markdown

from config import OUTPUT_DIR, COLORS
from analyzer import (
    AnalysisResult,
    ResolverProfile,
    CategoryRisk,
    WorkloadBalanceResult,
    KnowledgeSilo,
    TrendAnalysis,
    TrendPeriod
)
from i18n import get_text, get_all_texts
from utils import format_duration, format_percentage


# =============================================================================
# HTML Template
# =============================================================================

HTML_TEMPLATE = """
<!DOCTYPE html>
<html lang="{{ language }}">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{{ texts.report_title }}</title>
    <style>
        :root {
            --primary: {{ colors.primary }};
            --secondary: {{ colors.secondary }};
            --success: {{ colors.success }};
            --warning: {{ colors.warning }};
            --danger: {{ colors.danger }};
            --info: {{ colors.info }};
            --light: {{ colors.light }};
            --dark: {{ colors.dark }};
        }

        * {
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }

        body {
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, 'Helvetica Neue', Arial, sans-serif;
            line-height: 1.6;
            color: var(--dark);
            background: var(--light);
            padding: 2rem;
        }

        .container {
            max-width: 1200px;
            margin: 0 auto;
            background: white;
            padding: 2rem;
            border-radius: 8px;
            box-shadow: 0 2px 10px rgba(0,0,0,0.1);
        }

        /* Sticky TOC Navigation */
        .toc-nav {
            position: fixed;
            top: 20px;
            right: 20px;
            width: 200px;
            background: white;
            border-radius: 8px;
            box-shadow: 0 2px 10px rgba(0,0,0,0.15);
            padding: 1rem;
            font-size: 0.8rem;
            max-height: calc(100vh - 40px);
            overflow-y: auto;
            z-index: 1000;
        }

        .toc-nav h4 {
            color: var(--primary);
            margin-bottom: 0.5rem;
            font-size: 0.9rem;
            border-bottom: 2px solid var(--secondary);
            padding-bottom: 0.5rem;
        }

        .toc-nav ul {
            list-style: none;
        }

        .toc-nav li {
            margin: 0.4rem 0;
        }

        .toc-nav a {
            color: var(--dark);
            text-decoration: none;
            transition: color 0.2s;
        }

        .toc-nav a:hover {
            color: var(--primary);
        }

        @media (max-width: 1400px) {
            .toc-nav {
                display: none;
            }
        }

        header {
            text-align: center;
            margin-bottom: 2rem;
            padding-bottom: 1.5rem;
            border-bottom: 3px solid var(--primary);
            background: linear-gradient(135deg, rgba(46, 80, 144, 0.05), rgba(74, 144, 217, 0.05));
            margin: -2rem -2rem 2rem -2rem;
            padding: 2rem;
            border-radius: 8px 8px 0 0;
        }

        h1 {
            color: var(--primary);
            font-size: 2.2rem;
            margin-bottom: 0.5rem;
        }

        .subtitle {
            color: var(--secondary);
            font-size: 1.2rem;
            font-weight: 500;
        }

        .meta {
            color: #666;
            font-size: 0.9rem;
            margin-top: 0.5rem;
        }

        section {
            margin-bottom: 2.5rem;
            scroll-margin-top: 1rem;
        }

        h2 {
            color: var(--primary);
            font-size: 1.5rem;
            margin-bottom: 1rem;
            padding-bottom: 0.5rem;
            border-bottom: 2px solid var(--secondary);
            display: flex;
            align-items: center;
            gap: 0.5rem;
        }

        h2 .section-icon {
            font-size: 1.2rem;
        }

        h3 {
            color: var(--secondary);
            font-size: 1.2rem;
            margin: 1rem 0;
        }

        /* Enhanced KPI Cards */
        .kpi-grid {
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(200px, 1fr));
            gap: 1rem;
            margin-bottom: 1.5rem;
        }

        .kpi-card {
            background: linear-gradient(135deg, var(--primary), var(--secondary));
            color: white;
            padding: 1.5rem;
            border-radius: 12px;
            text-align: center;
            position: relative;
            overflow: hidden;
        }

        .kpi-card::before {
            content: '';
            position: absolute;
            top: -50%;
            right: -50%;
            width: 100%;
            height: 100%;
            background: rgba(255,255,255,0.1);
            border-radius: 50%;
        }

        .kpi-card.warning {
            background: linear-gradient(135deg, var(--warning), #ffdb4d);
            color: var(--dark);
        }

        .kpi-card.danger {
            background: linear-gradient(135deg, var(--danger), #ff6b6b);
        }

        .kpi-card.success {
            background: linear-gradient(135deg, var(--success), #51cf66);
        }

        .kpi-value {
            font-size: 2.2rem;
            font-weight: bold;
            margin-bottom: 0.25rem;
            position: relative;
        }

        .kpi-trend {
            font-size: 0.9rem;
            opacity: 0.9;
            margin-left: 0.5rem;
        }

        .kpi-trend.up { color: #90EE90; }
        .kpi-trend.down { color: #FFB6C1; }
        .kpi-trend.stable { opacity: 0.7; }

        .kpi-card.warning .kpi-trend.up,
        .kpi-card.warning .kpi-trend.down,
        .kpi-card.warning .kpi-trend.stable { color: var(--dark); }

        .kpi-label {
            font-size: 0.85rem;
            opacity: 0.9;
        }

        .kpi-sublabel {
            font-size: 0.75rem;
            opacity: 0.7;
            margin-top: 0.25rem;
        }

        /* Tables with responsive wrapper */
        .table-wrapper {
            overflow-x: auto;
            margin: 1rem 0;
            border-radius: 8px;
            box-shadow: 0 1px 3px rgba(0,0,0,0.1);
        }

        table {
            width: 100%;
            border-collapse: collapse;
            font-size: 0.9rem;
            min-width: 600px;
        }

        th, td {
            padding: 0.75rem;
            text-align: left;
            border-bottom: 1px solid #ddd;
        }

        th {
            background: var(--primary);
            color: white;
            font-weight: 600;
            position: sticky;
            top: 0;
        }

        tr:hover {
            background: var(--light);
        }

        .risk-high {
            color: var(--danger);
            font-weight: bold;
        }

        .risk-medium {
            color: var(--warning);
            font-weight: bold;
        }

        .risk-low {
            color: var(--success);
        }

        /* Charts */
        .chart-container {
            margin: 1.5rem 0;
            text-align: center;
        }

        .chart-container img {
            max-width: 100%;
            height: auto;
            border-radius: 8px;
            box-shadow: 0 2px 8px rgba(0,0,0,0.1);
        }

        /* Insights - Markdown rendered */
        .insight-box {
            background: linear-gradient(135deg, #f8f9ff, #f0f4ff);
            border-left: 4px solid var(--primary);
            padding: 1.5rem;
            margin: 1rem 0;
            border-radius: 0 8px 8px 0;
        }

        .insight-box h4 {
            color: var(--primary);
            margin-bottom: 0.75rem;
            display: flex;
            align-items: center;
            gap: 0.5rem;
        }

        .insight-content {
            line-height: 1.8;
        }

        .insight-content p {
            margin-bottom: 0.75rem;
        }

        .insight-content ul, .insight-content ol {
            margin: 0.5rem 0 0.5rem 1.5rem;
        }

        .insight-content li {
            margin-bottom: 0.3rem;
        }

        .insight-content strong {
            color: var(--primary);
        }

        /* Progress bars */
        .progress-bar {
            height: 24px;
            background: #e9ecef;
            border-radius: 12px;
            overflow: hidden;
        }

        .progress-fill {
            height: 100%;
            display: flex;
            align-items: center;
            justify-content: center;
            color: white;
            font-size: 0.8rem;
            font-weight: bold;
            transition: width 0.3s ease;
        }

        /* Tier badges - improved colors */
        .tier-badge {
            display: inline-block;
            padding: 0.25rem 0.75rem;
            border-radius: 12px;
            font-size: 0.8rem;
            font-weight: bold;
        }

        .tier-1 {
            background: linear-gradient(135deg, #6c757d, #8c959d);
            color: white;
        }

        .tier-2 {
            background: linear-gradient(135deg, #4A90D9, #6BA8E8);
            color: white;
        }

        .tier-3 {
            background: linear-gradient(135deg, #2E5090, #4E70B0);
            color: white;
        }

        /* Trend Section */
        .trend-chart-container {
            background: var(--light);
            padding: 1.5rem;
            border-radius: 8px;
            margin: 1rem 0;
        }

        .trend-summary {
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(150px, 1fr));
            gap: 1rem;
            margin-top: 1rem;
        }

        .trend-item {
            text-align: center;
            padding: 1rem;
            background: white;
            border-radius: 8px;
            box-shadow: 0 1px 3px rgba(0,0,0,0.1);
        }

        .trend-item .trend-label {
            font-size: 0.85rem;
            color: #666;
        }

        .trend-item .trend-value {
            font-size: 1.2rem;
            font-weight: bold;
            margin: 0.25rem 0;
        }

        .trend-item .trend-indicator {
            font-size: 0.9rem;
        }

        .trend-improving { color: var(--success); }
        .trend-deteriorating { color: var(--danger); }
        .trend-stable { color: var(--info); }
        .trend-increasing { color: var(--warning); }
        .trend-decreasing { color: var(--info); }

        /* Footer */
        footer {
            text-align: center;
            margin-top: 2rem;
            padding: 1.5rem;
            border-top: 2px solid var(--light);
            color: #666;
            font-size: 0.85rem;
            background: var(--light);
            margin: 2rem -2rem -2rem -2rem;
            border-radius: 0 0 8px 8px;
        }

        footer .footer-meta {
            display: flex;
            justify-content: center;
            gap: 2rem;
            flex-wrap: wrap;
            margin-bottom: 0.5rem;
        }

        footer .disclaimer {
            font-size: 0.75rem;
            color: #999;
            margin-top: 0.5rem;
        }

        /* Print styles */
        @media print {
            body {
                padding: 0;
                background: white;
            }
            .container {
                box-shadow: none;
                max-width: none;
            }
            .toc-nav {
                display: none;
            }
            .chart-container img {
                max-height: 350px;
            }
            section {
                page-break-inside: avoid;
            }
            h2 {
                page-break-after: avoid;
            }
            table {
                page-break-inside: avoid;
            }
            .kpi-grid {
                page-break-inside: avoid;
            }
            header, footer {
                -webkit-print-color-adjust: exact;
                print-color-adjust: exact;
            }
        }
    </style>
</head>
<body>
    <!-- TOC Navigation -->
    <nav class="toc-nav">
        <h4>{{ texts.nav_contents if texts.nav_contents else 'Contents' }}</h4>
        <ul>
            <li><a href="#executive-summary">{{ texts.section_executive_summary }}</a></li>
            {% if trend_analysis and trend_analysis.periods %}
            <li><a href="#trend-analysis">{{ texts.section_trend_analysis if texts.section_trend_analysis else 'Trend Analysis' }}</a></li>
            {% endif %}
            <li><a href="#tier-classification">{{ texts.section_tier_classification }}</a></li>
            <li><a href="#dependency-matrix">{{ texts.section_dependency_matrix }}</a></li>
            <li><a href="#bottleneck-analysis">{{ texts.section_bottleneck_analysis }}</a></li>
            <li><a href="#workload-balance">{{ texts.section_workload_balance }}</a></li>
            <li><a href="#knowledge-silos">{{ texts.section_knowledge_silo }}</a></li>
            <li><a href="#additional-analysis">{{ texts.section_expert_profiles }}</a></li>
        </ul>
    </nav>

    <div class="container">
        <header>
            <h1>{{ texts.report_title }}</h1>
            <p class="subtitle">{{ texts.report_subtitle }}</p>
            <p class="meta">{{ texts.period }}: {{ result.start_date }} {{ texts.to }} {{ result.end_date }}</p>
            <p class="meta">{{ texts.generated_on }}: {{ generated_time }}</p>
        </header>

        <!-- Executive Summary -->
        <section id="executive-summary">
            <h2><span class="section-icon">📊</span> {{ texts.section_executive_summary }}</h2>

            <div class="kpi-grid">
                <div class="kpi-card">
                    <div class="kpi-value">{{ result.total_bo_resolvers }}</div>
                    <div class="kpi-label">{{ texts.exec_total_bo_experts }}</div>
                </div>
                <div class="kpi-card">
                    <div class="kpi-value">
                        {{ result.total_bo_tickets }}
                        {% if trend_analysis and trend_analysis.volume_trend %}
                        <span class="kpi-trend {% if trend_analysis.volume_trend == 'increasing' %}up{% elif trend_analysis.volume_trend == 'decreasing' %}down{% else %}stable{% endif %}">
                            {% if trend_analysis.volume_trend == 'increasing' %}↑{% elif trend_analysis.volume_trend == 'decreasing' %}↓{% else %}→{% endif %}
                        </span>
                        {% endif %}
                    </div>
                    <div class="kpi-label">{{ texts.exec_total_bo_tickets }}</div>
                    {% if trend_analysis and trend_analysis.volume_change_pct %}
                    <div class="kpi-sublabel">{{ "%.1f"|format(trend_analysis.volume_change_pct) }}% vs prev</div>
                    {% endif %}
                </div>
                <div class="kpi-card {% if result.global_avg_mttr > 24 %}warning{% endif %}">
                    <div class="kpi-value">
                        {{ "%.1f"|format(result.global_avg_mttr) }}
                        {% if trend_analysis and trend_analysis.mttr_trend %}
                        <span class="kpi-trend {% if trend_analysis.mttr_trend == 'improving' %}down{% elif trend_analysis.mttr_trend == 'deteriorating' %}up{% else %}stable{% endif %}">
                            {% if trend_analysis.mttr_trend == 'improving' %}↓{% elif trend_analysis.mttr_trend == 'deteriorating' %}↑{% else %}→{% endif %}
                        </span>
                        {% endif %}
                    </div>
                    <div class="kpi-label">{{ texts.exec_avg_mttr }} ({{ texts.exec_hours }})</div>
                    {% if trend_analysis and trend_analysis.mttr_change_pct %}
                    <div class="kpi-sublabel">{{ "%.1f"|format(trend_analysis.mttr_change_pct) }}% vs prev</div>
                    {% endif %}
                </div>
                <div class="kpi-card {% if result.global_sla_rate < 0.9 %}danger{% elif result.global_sla_rate >= 0.95 %}success{% endif %}">
                    <div class="kpi-value">
                        {{ "%.1f"|format(result.global_sla_rate * 100) }}%
                        {% if trend_analysis and trend_analysis.sla_trend %}
                        <span class="kpi-trend {% if trend_analysis.sla_trend == 'improving' %}up{% elif trend_analysis.sla_trend == 'deteriorating' %}down{% else %}stable{% endif %}">
                            {% if trend_analysis.sla_trend == 'improving' %}↑{% elif trend_analysis.sla_trend == 'deteriorating' %}↓{% else %}→{% endif %}
                        </span>
                        {% endif %}
                    </div>
                    <div class="kpi-label">{{ texts.exec_sla_rate }}</div>
                    {% if trend_analysis and trend_analysis.sla_change_pct %}
                    <div class="kpi-sublabel">{{ "%.1f"|format(trend_analysis.sla_change_pct) }}% vs prev</div>
                    {% endif %}
                </div>
                <div class="kpi-card {% if result.high_risk_count > 5 %}danger{% elif result.high_risk_count > 0 %}warning{% else %}success{% endif %}">
                    <div class="kpi-value">{{ result.high_risk_count }}</div>
                    <div class="kpi-label">{{ texts.exec_high_risk_count }}</div>
                </div>
                <div class="kpi-card {% if result.bottleneck_count > 3 %}danger{% elif result.bottleneck_count > 0 %}warning{% else %}success{% endif %}">
                    <div class="kpi-value">{{ result.bottleneck_count }}</div>
                    <div class="kpi-label">{{ texts.exec_bottleneck_count }}</div>
                </div>
            </div>

            {% if insights.executive_summary %}
            <div class="insight-box">
                <h4>🤖 AI {{ texts.insight_summary }}</h4>
                <div class="insight-content">{{ insights.executive_summary | safe }}</div>
            </div>
            {% endif %}
        </section>

        <!-- Trend Analysis Section -->
        {% if trend_analysis and trend_analysis.periods %}
        <section id="trend-analysis">
            <h2><span class="section-icon">📈</span> {{ texts.section_trend_analysis if texts.section_trend_analysis else 'Trend Analysis' }}</h2>

            <div class="trend-chart-container">
                {% if charts.trend_line %}
                <div class="chart-container">
                    <img src="data:image/png;base64,{{ charts.trend_line }}" alt="Trend Analysis">
                </div>
                {% endif %}

                <div class="trend-summary">
                    <div class="trend-item">
                        <div class="trend-label">{{ texts.exec_avg_mttr if texts.exec_avg_mttr else 'MTTR' }} {{ texts.trend if texts.trend else 'Trend' }}</div>
                        <div class="trend-value trend-{{ trend_analysis.mttr_trend }}">
                            {% if trend_analysis.mttr_trend == 'improving' %}↓ {{ texts.trend_improving if texts.trend_improving else 'Improving' }}
                            {% elif trend_analysis.mttr_trend == 'deteriorating' %}↑ {{ texts.trend_deteriorating if texts.trend_deteriorating else 'Deteriorating' }}
                            {% else %}→ {{ texts.trend_stable if texts.trend_stable else 'Stable' }}{% endif %}
                        </div>
                        <div class="trend-indicator">{{ "%.1f"|format(trend_analysis.mttr_change_pct) }}%</div>
                    </div>
                    <div class="trend-item">
                        <div class="trend-label">SLA {{ texts.trend if texts.trend else 'Trend' }}</div>
                        <div class="trend-value trend-{{ trend_analysis.sla_trend }}">
                            {% if trend_analysis.sla_trend == 'improving' %}↑ {{ texts.trend_improving if texts.trend_improving else 'Improving' }}
                            {% elif trend_analysis.sla_trend == 'deteriorating' %}↓ {{ texts.trend_deteriorating if texts.trend_deteriorating else 'Deteriorating' }}
                            {% else %}→ {{ texts.trend_stable if texts.trend_stable else 'Stable' }}{% endif %}
                        </div>
                        <div class="trend-indicator">{{ "%.1f"|format(trend_analysis.sla_change_pct) }}%</div>
                    </div>
                    <div class="trend-item">
                        <div class="trend-label">{{ texts.volume if texts.volume else 'Volume' }} {{ texts.trend if texts.trend else 'Trend' }}</div>
                        <div class="trend-value trend-{{ trend_analysis.volume_trend }}">
                            {% if trend_analysis.volume_trend == 'increasing' %}↑ {{ texts.trend_increasing if texts.trend_increasing else 'Increasing' }}
                            {% elif trend_analysis.volume_trend == 'decreasing' %}↓ {{ texts.trend_decreasing if texts.trend_decreasing else 'Decreasing' }}
                            {% else %}→ {{ texts.trend_stable if texts.trend_stable else 'Stable' }}{% endif %}
                        </div>
                        <div class="trend-indicator">{{ "%.1f"|format(trend_analysis.volume_change_pct) }}%</div>
                    </div>
                    <div class="trend-item">
                        <div class="trend-label">{{ texts.period_type if texts.period_type else 'Period Type' }}</div>
                        <div class="trend-value">{{ trend_analysis.period_type | capitalize }}</div>
                        <div class="trend-indicator">{{ trend_analysis.periods | length }} {{ texts.periods if texts.periods else 'periods' }}</div>
                    </div>
                </div>
            </div>
        </section>
        {% endif %}

        <!-- Tier Classification -->
        <section id="tier-classification">
            <h2><span class="section-icon">👥</span> {{ texts.section_tier_classification }}</h2>
            <p>{{ texts.tier_classification_desc }}</p>

            {% if charts.tier_distribution %}
            <div class="chart-container">
                <img src="data:image/png;base64,{{ charts.tier_distribution }}" alt="Tier Distribution">
            </div>
            {% endif %}

            <div class="table-wrapper">
                <table>
                    <thead>
                        <tr>
                            <th>{{ texts.tier_resolver }}</th>
                            <th>{{ texts.tier_assigned }}</th>
                            <th>{{ texts.tier_complexity_index }}</th>
                            <th>{{ texts.tier_priority_index }}</th>
                            <th>{{ texts.profile_total_tickets }}</th>
                            <th>{{ texts.weighted_workload if texts.weighted_workload else 'Weighted' }}</th>
                            <th>{{ texts.tier_justification }}</th>
                        </tr>
                    </thead>
                    <tbody>
                        {% for profile in bo_profiles %}
                        <tr>
                            <td>{{ profile.name }}</td>
                            <td><span class="tier-badge tier-{{ profile.tier[-1] }}">{{ profile.tier }}</span></td>
                            <td>{{ "%.2f"|format(profile.complexity_index) }}x</td>
                            <td>{{ "%.1f"|format(profile.priority_index * 100) }}%</td>
                            <td>{{ profile.total_tickets }}</td>
                            <td>{{ "%.0f"|format(profile.weighted_workload) }}</td>
                            <td>{{ profile.tier_justification }}</td>
                        </tr>
                        {% endfor %}
                    </tbody>
                </table>
            </div>
        </section>

        <!-- Single-Point Dependency -->
        <section id="dependency-matrix">
            <h2><span class="section-icon">🔗</span> {{ texts.section_dependency_matrix }}</h2>

            {% if charts.dependency_heatmap %}
            <div class="chart-container">
                <img src="data:image/png;base64,{{ charts.dependency_heatmap }}" alt="Dependency Heatmap">
            </div>
            {% endif %}

            {% if charts.category_coverage %}
            <div class="chart-container">
                <img src="data:image/png;base64,{{ charts.category_coverage }}" alt="Category Coverage">
            </div>
            {% endif %}

            <h3>{{ texts.dependency_high_risk }}</h3>
            <div class="table-wrapper">
                <table>
                    <thead>
                        <tr>
                            <th>{{ texts.dependency_category }}</th>
                            <th>{{ texts.dependency_resolver_count }}</th>
                            <th>{{ texts.silo_ticket_count }}</th>
                            <th>{{ texts.dependency_resolvers }}</th>
                            <th>{{ texts.dependency_risk_level }}</th>
                        </tr>
                    </thead>
                    <tbody>
                        {% for risk in high_risk_categories %}
                        <tr>
                            <td>{{ risk.category }}</td>
                            <td>{{ risk.resolver_count }}</td>
                            <td>{{ risk.ticket_count }}</td>
                            <td>{{ risk.resolvers | join(', ') }}</td>
                            <td class="risk-{{ risk.risk_level | lower }}">{{ risk.risk_level }}</td>
                        </tr>
                        {% endfor %}
                    </tbody>
                </table>
            </div>

            {% if insights.dependency %}
            <div class="insight-box">
                <h4>🤖 AI {{ texts.insight_recommendations }}</h4>
                <div class="insight-content">{{ insights.dependency | safe }}</div>
            </div>
            {% endif %}
        </section>

        <!-- Bottleneck Analysis -->
        <section id="bottleneck-analysis">
            <h2><span class="section-icon">⚠️</span> {{ texts.section_bottleneck_analysis }}</h2>

            {% if charts.bottleneck_ranking %}
            <div class="chart-container">
                <img src="data:image/png;base64,{{ charts.bottleneck_ranking }}" alt="Bottleneck Ranking">
            </div>
            {% endif %}

            <div class="table-wrapper">
                <table>
                    <thead>
                        <tr>
                            <th>{{ texts.bottleneck_rank }}</th>
                            <th>{{ texts.tier_resolver }}</th>
                            <th>{{ texts.bottleneck_score }}</th>
                            <th>{{ texts.bottleneck_backlog }}</th>
                            <th>{{ texts.bottleneck_mttr }}</th>
                            <th>{{ texts.exec_sla_rate }}</th>
                            <th>{{ texts.bottleneck_root_cause }}</th>
                        </tr>
                    </thead>
                    <tbody>
                        {% for profile in top_bottlenecks %}
                        <tr>
                            <td>{{ loop.index }}</td>
                            <td>{{ profile.name }}</td>
                            <td>
                                <div class="progress-bar">
                                    <div class="progress-fill" style="width: {{ profile.bottleneck_score * 100 }}%; background: {% if profile.bottleneck_score >= 0.6 %}var(--danger){% elif profile.bottleneck_score >= 0.4 %}var(--warning){% else %}var(--success){% endif %}">
                                        {{ "%.2f"|format(profile.bottleneck_score) }}
                                    </div>
                                </div>
                            </td>
                            <td>{{ profile.current_backlog }}</td>
                            <td>{{ "%.1f"|format(profile.avg_mttr_hours) }}h</td>
                            <td>{{ "%.1f"|format(profile.sla_rate * 100) }}%</td>
                            <td>{{ profile.bottleneck_root_causes | join('; ') }}</td>
                        </tr>
                        {% endfor %}
                    </tbody>
                </table>
            </div>

            {% if insights.bottleneck %}
            <div class="insight-box">
                <h4>🤖 AI {{ texts.insight_recommendations }}</h4>
                <div class="insight-content">{{ insights.bottleneck | safe }}</div>
            </div>
            {% endif %}
        </section>

        <!-- Workload Balance -->
        <section id="workload-balance">
            <h2><span class="section-icon">⚖️</span> {{ texts.section_workload_balance }}</h2>

            {% if charts.workload_distribution %}
            <div class="chart-container">
                <img src="data:image/png;base64,{{ charts.workload_distribution }}" alt="Workload Distribution">
            </div>
            {% endif %}

            {% if workload_balance %}
            <div class="kpi-grid">
                <div class="kpi-card {% if workload_balance.gini_coefficient > 0.5 %}danger{% elif workload_balance.gini_coefficient > 0.3 %}warning{% else %}success{% endif %}">
                    <div class="kpi-value">{{ "%.3f"|format(workload_balance.gini_coefficient) }}</div>
                    <div class="kpi-label">{{ texts.balance_gini }} ({{ texts.raw_workload if texts.raw_workload else 'Raw' }})</div>
                </div>
                {% if workload_balance.weighted_gini_coefficient > 0 %}
                <div class="kpi-card {% if workload_balance.weighted_gini_coefficient > 0.5 %}danger{% elif workload_balance.weighted_gini_coefficient > 0.3 %}warning{% else %}success{% endif %}">
                    <div class="kpi-value">{{ "%.3f"|format(workload_balance.weighted_gini_coefficient) }}</div>
                    <div class="kpi-label">{{ texts.weighted_gini if texts.weighted_gini else 'Weighted Gini' }}</div>
                </div>
                {% endif %}
                <div class="kpi-card">
                    <div class="kpi-value">{{ "%.1f"|format(workload_balance.avg_workload) }}</div>
                    <div class="kpi-label">{{ texts.balance_avg_workload }}</div>
                </div>
                <div class="kpi-card {% if workload_balance.overloaded_resolvers %}warning{% endif %}">
                    <div class="kpi-value">{{ workload_balance.overloaded_resolvers | length }}</div>
                    <div class="kpi-label">{{ texts.balance_overloaded }}</div>
                </div>
            </div>

            <p><strong>{{ texts.balance_interpretation }}:</strong> {{ workload_balance.interpretation }}</p>
            {% if workload_balance.weighted_interpretation %}
            <p><strong>{{ texts.weighted_balance if texts.weighted_balance else 'Priority-Weighted Balance' }}:</strong> {{ workload_balance.weighted_interpretation }}</p>
            {% endif %}

            {% if workload_balance.overloaded_resolvers %}
            <p><strong>{{ texts.balance_overloaded }}:</strong> {{ workload_balance.overloaded_resolvers | join(', ') }}</p>
            {% endif %}
            {% endif %}

            {% if insights.workload_balance %}
            <div class="insight-box">
                <h4>🤖 AI {{ texts.insight_recommendations }}</h4>
                <div class="insight-content">{{ insights.workload_balance | safe }}</div>
            </div>
            {% endif %}
        </section>

        <!-- Priority Distribution -->
        <section id="priority-distribution">
            <h2><span class="section-icon">🎯</span> {{ texts.priority_distribution if texts.priority_distribution else 'Priority Distribution' }}</h2>
            
            {% if charts.priority_heatmap %}
            <div class="chart-container">
                <h3>{{ texts.chart_priority_heatmap if texts.chart_priority_heatmap else 'Priority Distribution Heatmap' }}</h3>
                <img src="data:image/png;base64,{{ charts.priority_heatmap }}" alt="Priority Distribution Heatmap">
            </div>
            {% endif %}
            
            <p>{{ texts.priority_distribution_desc if texts.priority_distribution_desc else 'This heatmap shows how tickets of different priorities are distributed among resolvers. Higher P1/P2 counts indicate specialists handling more critical work.' }}</p>
        </section>

        <!-- Knowledge Silos -->
        <section id="knowledge-silos">
            <h2><span class="section-icon">🏝️</span> {{ texts.section_knowledge_silo }}</h2>

            <div class="table-wrapper">
                <table>
                    <thead>
                        <tr>
                            <th>{{ texts.silo_category }}</th>
                            <th>{{ texts.silo_sole_expert }}</th>
                            <th>{{ texts.silo_ticket_count }}</th>
                            <th>{{ texts.profile_avg_mttr }}</th>
                            <th>{{ texts.silo_priority }}</th>
                        </tr>
                    </thead>
                    <tbody>
                        {% for silo in knowledge_silos %}
                        <tr>
                            <td>{{ silo.category }}</td>
                            <td>{{ silo.sole_experts | join(', ') }}</td>
                            <td>{{ silo.ticket_count }}</td>
                            <td>{{ "%.1f"|format(silo.avg_mttr) }}h</td>
                            <td class="risk-{{ silo.priority | lower }}">{{ silo.priority }}</td>
                        </tr>
                        {% endfor %}
                    </tbody>
                </table>
            </div>
        </section>

        <!-- Additional Charts -->
        <section id="additional-analysis">
            <h2><span class="section-icon">📉</span> {{ texts.section_expert_profiles }}</h2>

            {% if charts.mttr_by_tier %}
            <div class="chart-container">
                <h3>{{ texts.chart_mttr_comparison }}</h3>
                <img src="data:image/png;base64,{{ charts.mttr_by_tier }}" alt="MTTR by Tier">
            </div>
            {% endif %}

            {% if charts.expert_hourly %}
            <div class="chart-container">
                <h3>{{ texts.chart_hourly_pattern if texts.chart_hourly_pattern else 'Expert Hourly Activity Pattern' }}</h3>
                <img src="data:image/png;base64,{{ charts.expert_hourly }}" alt="Expert Hourly Heatmap">
            </div>
            {% endif %}

            {% if charts.sla_comparison %}
            <div class="chart-container">
                <h3>{{ texts.chart_sla_comparison if texts.chart_sla_comparison else 'SLA Compliance Comparison' }}</h3>
                <img src="data:image/png;base64,{{ charts.sla_comparison }}" alt="SLA Comparison">
            </div>
            {% endif %}
        </section>

        <footer>
            <div class="footer-meta">
                <span>{{ texts.report_title }}</span>
                <span>{{ texts.generated_on }}: {{ generated_time }}</span>
                <span>{{ texts.period }}: {{ result.start_date }} {{ texts.to }} {{ result.end_date }}</span>
            </div>
            <div class="disclaimer">
                {{ texts.disclaimer if texts.disclaimer else 'This report is auto-generated for internal use. Data accuracy depends on source system inputs.' }}
            </div>
        </footer>
    </div>
</body>
</html>
"""


class ReportBuilder:
    """
    Report builder for DOCX and HTML output.
    """

    def __init__(
        self,
        result: AnalysisResult,
        charts: Dict[str, str],
        insights: Dict[str, str],
        language: str = "en"
    ):
        """
        Initialize report builder.

        Args:
            result: Analysis result
            charts: Dictionary of chart name to base64 image
            insights: Dictionary of insight name to text
            language: Output language
        """
        self.result = result
        self.charts = charts
        self.insights = insights
        self.language = language
        self.texts = get_all_texts(language)

    def _convert_markdown_to_html(self, text: str) -> str:
        """
        Convert markdown text to HTML.

        Args:
            text: Markdown text

        Returns:
            HTML string
        """
        if not text:
            return ""
        return markdown.markdown(text, extensions=['tables', 'fenced_code', 'nl2br'])

    def _convert_insights_to_html(self, insights: Dict[str, str]) -> Dict[str, str]:
        """
        Convert all insights from markdown to HTML.

        Args:
            insights: Dictionary of insight name to markdown text

        Returns:
            Dictionary of insight name to HTML string
        """
        return {key: self._convert_markdown_to_html(value) for key, value in insights.items()}
    
    def build_html(self) -> str:
        """
        Build HTML report.
        
        Returns:
            HTML string
        """
        template = Template(HTML_TEMPLATE)
        
        # Prepare data
        bo_profiles = [p for p in self.result.resolver_profiles if p.tier in ["Tier 2", "Tier 3"]]
        bo_profiles.sort(key=lambda x: x.tier, reverse=True)
        
        high_risk_categories = [r for r in self.result.category_risks if r.risk_level == "High"]
        
        top_bottlenecks = sorted(
            self.result.resolver_profiles,
            key=lambda x: x.bottleneck_score,
            reverse=True
        )[:10]

        # Convert markdown insights to HTML
        html_insights = self._convert_insights_to_html(self.insights)

        # Render template
        html = template.render(
            language=self.language,
            texts=self.texts,
            colors=COLORS,
            result=self.result,
            charts=self.charts,
            insights=html_insights,
            bo_profiles=bo_profiles,
            high_risk_categories=high_risk_categories,
            top_bottlenecks=top_bottlenecks,
            workload_balance=self.result.workload_balance,
            knowledge_silos=self.result.knowledge_silos[:10],
            trend_analysis=self.result.trend_analysis,
            generated_time=datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        )
        
        return html
    
    def build_docx(self) -> Document:
        """
        Build DOCX report.
        
        Returns:
            python-docx Document object
        """
        doc = Document()
        
        # Set styles
        self._setup_styles(doc)
        
        # Title
        title = doc.add_heading(self.texts["report_title"], 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = doc.add_paragraph(self.texts["report_subtitle"])
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        meta = doc.add_paragraph(
            f"{self.texts['period']}: {self.result.start_date} {self.texts['to']} {self.result.end_date}\n"
            f"{self.texts['generated_on']}: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
        )
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_page_break()
        
        # Table of Contents placeholder
        doc.add_heading("Table of Contents", 1)
        doc.add_paragraph("(Update fields to generate TOC)")
        doc.add_page_break()
        
        # Executive Summary
        doc.add_heading(self.texts["section_executive_summary"], 1)
        self._add_executive_summary_table(doc)
        
        if "executive_summary" in self.insights:
            doc.add_paragraph()
            self._add_insight_box(doc, self.insights["executive_summary"])
        
        doc.add_page_break()
        
        # Tier Classification
        doc.add_heading(self.texts["section_tier_classification"], 1)
        doc.add_paragraph(self.texts["tier_classification_desc"])
        
        if "tier_distribution" in self.charts:
            self._add_chart_image(doc, self.charts["tier_distribution"])
        
        self._add_tier_table(doc)
        
        doc.add_page_break()
        
        # Dependency Matrix
        doc.add_heading(self.texts["section_dependency_matrix"], 1)
        
        if "dependency_heatmap" in self.charts:
            self._add_chart_image(doc, self.charts["dependency_heatmap"])
        
        self._add_risk_table(doc)
        
        if "dependency" in self.insights:
            doc.add_paragraph()
            self._add_insight_box(doc, self.insights["dependency"])
        
        doc.add_page_break()
        
        # Bottleneck Analysis
        doc.add_heading(self.texts["section_bottleneck_analysis"], 1)
        
        if "bottleneck_ranking" in self.charts:
            self._add_chart_image(doc, self.charts["bottleneck_ranking"])
        
        self._add_bottleneck_table(doc)
        
        if "bottleneck" in self.insights:
            doc.add_paragraph()
            self._add_insight_box(doc, self.insights["bottleneck"])
        
        doc.add_page_break()
        
        # Workload Balance
        doc.add_heading(self.texts["section_workload_balance"], 1)
        
        if "workload_distribution" in self.charts:
            self._add_chart_image(doc, self.charts["workload_distribution"])
        
        if self.result.workload_balance:
            self._add_balance_summary(doc)
        
        if "workload_balance" in self.insights:
            doc.add_paragraph()
            self._add_insight_box(doc, self.insights["workload_balance"])
        
        doc.add_page_break()
        
        # Knowledge Silos
        doc.add_heading(self.texts["section_knowledge_silo"], 1)
        self._add_silo_table(doc)
        
        # Additional charts
        if "mttr_by_tier" in self.charts or "sla_comparison" in self.charts:
            doc.add_page_break()
            doc.add_heading(self.texts["section_expert_profiles"], 1)
            
            if "mttr_by_tier" in self.charts:
                self._add_chart_image(doc, self.charts["mttr_by_tier"])
            
            if "sla_comparison" in self.charts:
                self._add_chart_image(doc, self.charts["sla_comparison"])
        
        return doc
    
    def _setup_styles(self, doc: Document) -> None:
        """Setup document styles."""
        # Set default font
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(11)
    
    def _add_executive_summary_table(self, doc: Document) -> None:
        """Add executive summary KPI table."""
        table = doc.add_table(rows=2, cols=6)
        table.style = 'Table Grid'
        
        # Headers
        headers = [
            self.texts["exec_total_bo_experts"],
            self.texts["exec_total_bo_tickets"],
            f"{self.texts['exec_avg_mttr']} ({self.texts['exec_hours']})",
            self.texts["exec_sla_rate"],
            self.texts["exec_high_risk_count"],
            self.texts["exec_bottleneck_count"]
        ]
        
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            self._set_cell_shading(cell, "2E5090")
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.font.bold = True
        
        # Values
        values = [
            str(self.result.total_bo_resolvers),
            str(self.result.total_bo_tickets),
            f"{self.result.global_avg_mttr:.1f}",
            f"{self.result.global_sla_rate:.1%}",
            str(self.result.high_risk_count),
            str(self.result.bottleneck_count)
        ]
        
        for i, value in enumerate(values):
            table.rows[1].cells[i].text = value
    
    def _add_tier_table(self, doc: Document) -> None:
        """Add tier classification table."""
        bo_profiles = [p for p in self.result.resolver_profiles if p.tier in ["Tier 2", "Tier 3"]]
        bo_profiles.sort(key=lambda x: x.tier, reverse=True)
        
        if not bo_profiles:
            doc.add_paragraph(self.texts["no_data"])
            return
        
        table = doc.add_table(rows=len(bo_profiles) + 1, cols=5)
        table.style = 'Table Grid'
        
        # Headers
        headers = [
            self.texts["tier_resolver"],
            self.texts["tier_assigned"],
            self.texts["tier_complexity_index"],
            self.texts["profile_total_tickets"],
            self.texts["tier_justification"]
        ]
        
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            self._set_cell_shading(cell, "2E5090")
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.font.bold = True
        
        # Data rows
        for row_idx, profile in enumerate(bo_profiles, 1):
            table.rows[row_idx].cells[0].text = profile.name
            table.rows[row_idx].cells[1].text = profile.tier
            table.rows[row_idx].cells[2].text = f"{profile.complexity_index:.2f}x"
            table.rows[row_idx].cells[3].text = str(profile.total_tickets)
            table.rows[row_idx].cells[4].text = profile.tier_justification[:50] + "..." if len(profile.tier_justification) > 50 else profile.tier_justification
    
    def _add_risk_table(self, doc: Document) -> None:
        """Add high-risk categories table."""
        high_risk = [r for r in self.result.category_risks if r.risk_level == "High"]
        
        if not high_risk:
            doc.add_paragraph(self.texts["no_data"])
            return
        
        table = doc.add_table(rows=len(high_risk) + 1, cols=4)
        table.style = 'Table Grid'
        
        headers = [
            self.texts["dependency_category"],
            self.texts["dependency_resolver_count"],
            self.texts["silo_ticket_count"],
            self.texts["dependency_resolvers"]
        ]
        
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            self._set_cell_shading(cell, "2E5090")
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.font.bold = True
        
        for row_idx, risk in enumerate(high_risk, 1):
            table.rows[row_idx].cells[0].text = risk.category
            table.rows[row_idx].cells[1].text = str(risk.resolver_count)
            table.rows[row_idx].cells[2].text = str(risk.ticket_count)
            table.rows[row_idx].cells[3].text = ", ".join(risk.resolvers[:3])
    
    def _add_bottleneck_table(self, doc: Document) -> None:
        """Add bottleneck ranking table."""
        top_bottlenecks = sorted(
            self.result.resolver_profiles,
            key=lambda x: x.bottleneck_score,
            reverse=True
        )[:10]
        
        if not top_bottlenecks:
            doc.add_paragraph(self.texts["no_data"])
            return
        
        table = doc.add_table(rows=len(top_bottlenecks) + 1, cols=6)
        table.style = 'Table Grid'
        
        headers = [
            self.texts["bottleneck_rank"],
            self.texts["tier_resolver"],
            self.texts["bottleneck_score"],
            self.texts["bottleneck_backlog"],
            self.texts["exec_sla_rate"],
            self.texts["bottleneck_root_cause"]
        ]
        
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            self._set_cell_shading(cell, "2E5090")
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.font.bold = True
        
        for row_idx, profile in enumerate(top_bottlenecks, 1):
            table.rows[row_idx].cells[0].text = str(row_idx)
            table.rows[row_idx].cells[1].text = profile.name
            table.rows[row_idx].cells[2].text = f"{profile.bottleneck_score:.2f}"
            table.rows[row_idx].cells[3].text = str(profile.current_backlog)
            table.rows[row_idx].cells[4].text = f"{profile.sla_rate:.1%}"
            table.rows[row_idx].cells[5].text = ", ".join(profile.bottleneck_root_causes[:2])
    
    def _add_balance_summary(self, doc: Document) -> None:
        """Add workload balance summary."""
        balance = self.result.workload_balance
        
        p = doc.add_paragraph()
        p.add_run(f"{self.texts['balance_gini']}: ").bold = True
        p.add_run(f"{balance.gini_coefficient:.3f}")
        
        p = doc.add_paragraph()
        p.add_run(f"{self.texts['balance_interpretation']}: ").bold = True
        p.add_run(balance.interpretation)
        
        if balance.overloaded_resolvers:
            p = doc.add_paragraph()
            p.add_run(f"{self.texts['balance_overloaded']}: ").bold = True
            p.add_run(", ".join(balance.overloaded_resolvers))
    
    def _add_silo_table(self, doc: Document) -> None:
        """Add knowledge silo table."""
        silos = self.result.knowledge_silos[:10]
        
        if not silos:
            doc.add_paragraph(self.texts["no_data"])
            return
        
        table = doc.add_table(rows=len(silos) + 1, cols=4)
        table.style = 'Table Grid'
        
        headers = [
            self.texts["silo_category"],
            self.texts["silo_sole_expert"],
            self.texts["silo_ticket_count"],
            self.texts["silo_priority"]
        ]
        
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            self._set_cell_shading(cell, "2E5090")
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.font.bold = True
        
        for row_idx, silo in enumerate(silos, 1):
            table.rows[row_idx].cells[0].text = silo.category
            table.rows[row_idx].cells[1].text = ", ".join(silo.sole_experts)
            table.rows[row_idx].cells[2].text = str(silo.ticket_count)
            table.rows[row_idx].cells[3].text = silo.priority
    
    def _add_insight_box(self, doc: Document, insight: str) -> None:
        """Add AI insight box."""
        p = doc.add_paragraph()
        p.add_run("AI Insights:").bold = True
        
        for line in insight.split('\n'):
            doc.add_paragraph(line)
    
    def _add_chart_image(self, doc: Document, base64_image: str) -> None:
        """Add chart image from base64."""
        if not base64_image:
            return
        
        try:
            image_data = base64.b64decode(base64_image)
            image_stream = BytesIO(image_data)
            doc.add_picture(image_stream, width=Inches(6))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            doc.add_paragraph(f"[Chart could not be rendered: {e}]")
    
    def _set_cell_shading(self, cell, color: str) -> None:
        """Set cell background color."""
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), color)
        cell._tc.get_or_add_tcPr().append(shading)
    
    def save(self, base_filename: str = None) -> Dict[str, str]:
        """
        Save reports to files.
        
        Args:
            base_filename: Optional base filename (without extension)
        
        Returns:
            Dictionary of format to file path
        """
        if base_filename is None:
            lang_suffix = "CN" if self.language == "zh" else "EN"
            base_filename = f"BO_Workload_Performance_Report_{self.result.start_date}_to_{self.result.end_date}_{lang_suffix}"
        
        output_paths = {}
        
        # Save HTML
        html_path = OUTPUT_DIR / f"{base_filename}.html"
        html_content = self.build_html()
        with open(html_path, "w", encoding="utf-8") as f:
            f.write(html_content)
        output_paths["html"] = str(html_path)
        
        # Save DOCX
        docx_path = OUTPUT_DIR / f"{base_filename}.docx"
        doc = self.build_docx()
        doc.save(docx_path)
        output_paths["docx"] = str(docx_path)
        
        return output_paths
