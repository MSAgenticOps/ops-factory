#!/usr/bin/env python3
"""
SLA Violation Attribution Analysis Report - Bilingual Version
Outputs: EN HTML, CN HTML, EN DOCX, CN DOCX
"""

import pandas as pd
import matplotlib.pyplot as plt
import matplotlib.dates as mdates
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os

# ============ Translations ============
LANG = {
    'en': {
        'title': 'SLA Violation Attribution Analysis Report',
        'data_period': 'Data Period',
        'executive_summary': 'Executive Summary',
        'total_tickets': 'Total Tickets',
        'response_sla': 'Response SLA',
        'resolution_sla': 'Resolution SLA',
        'total_violations': 'Total Violations',
        'high_risk_tickets': 'High Risk Tickets',
        'critical_alert': 'Critical Alert',
        'alert_text': 'Resolution SLA ({resolution_rate:.1f}%) is significantly below the 95% target. P1 Resolution SLA is only {p1_rate:.1f}% and P2 is {p2_rate:.1f}%. Immediate action required.',
        'p1p2_analysis': 'High Priority Analysis (P1/P2)',
        'p1p2_desc': 'P1 and P2 incidents require immediate attention. SLA targets: P1 Response 15min/Resolution 2h, P2 Response 30min/Resolution 6h.',
        'p1_performance': 'P1 Performance',
        'p2_performance': 'P2 Performance',
        'monthly_trend': 'Monthly Trend Analysis',
        'sla_by_priority': 'SLA Compliance by Priority',
        'priority': 'Priority',
        'response_target': 'Response Target',
        'resolution_target': 'Resolution Target',
        'total': 'Total',
        'risk_analysis': 'Risk Analysis',
        'high_risk_top10': 'High Risk Tickets (Top 10)',
        'high_risk_desc': 'Tickets approaching or exceeding SLA threshold. "Exceed" shows hours over SLA target.',
        'ticket': 'Ticket #',
        'category': 'Category',
        'resolver': 'Resolver',
        'actual': 'Actual',
        'target': 'Target',
        'exceed': 'Exceed',
        'attribution': 'Attribution',
        'violation_deep_dive': 'Violation Deep Dive',
        'top_violations': 'Top Violations by Severity',
        'resolution': 'Resolution',
        'attribution_analysis': 'Attribution Analysis',
        'attr_footnote': 'Note: Attribution total ({attr_total:,}) may differ from violation count ({violation_count:,}) because a single ticket can have multiple contributing factors.',
        'violations_by_category': 'Violations by Category',
        'recommendations': 'Improvement Recommendations',
        'rec_critical': 'CRITICAL',
        'rec_high': 'HIGH',
        'rec_medium': 'MEDIUM',
        'rec_p1p2_title': 'P1/P2 SLA Performance',
        'rec_p1p2_text': 'P1 Resolution SLA is at {p1_rate:.1f}% (target: 95%). P2 is at {p2_rate:.1f}%.',
        'rec_p1p2_actions': [
            'Establish dedicated P1/P2 response team with <15min initial response',
            'Implement automatic escalation after 50% SLA consumption',
            'Create P1 war room protocol with multi-team coordination',
            'Review P1/P2 classification criteria - ensure proper prioritization'
        ],
        'rec_timewindow_title': 'Time Window Factor ({pct:.0f}% of violations)',
        'rec_timewindow_text': 'Majority of violations occur during off-hours (18:00-09:00) and weekends.',
        'rec_timewindow_actions': [
            'Strengthen on-call coverage during 18:00-09:00 shift',
            'Implement weekend duty rotation with dedicated escalation path',
            'Consider automated triage and routing for off-hours tickets',
            'Evaluate 24/7 NOC coverage for P1/P2 incidents'
        ],
        'rec_external_title': 'External Dependencies ({pct:.0f}% of violations)',
        'rec_external_text': 'Tickets pending on external teams or customers cause significant delays.',
        'rec_external_actions': [
            'Establish SLA agreements with dependent teams (internal OLAs)',
            'Implement automated customer follow-up reminders',
            'Create "pending external" dashboard for proactive management',
            'Define escalation triggers when external wait exceeds threshold'
        ],
        'rec_process_title': 'Process Improvement ({pct:.0f}% of violations)',
        'rec_process_text': 'Reassignment and escalation delays contribute to SLA breaches.',
        'rec_process_actions': [
            'Reduce average reassignment count (target: <2 per ticket)',
            'Implement skill-based routing to minimize transfers',
            'Create clear escalation matrix with defined triggers',
            'Train L1 team on better initial categorization'
        ],
        'rec_category_title': 'Category Focus: {category}',
        'rec_category_text': 'This category accounts for {count:,} violations ({pct:.0f}% of total).',
        'rec_category_actions': [
            'Review monitoring alert thresholds to reduce false positives',
            'Create runbooks for common scenarios',
            'Consider automation for repetitive tasks'
        ],
        'rec_workload_title': 'Top Resolvers Workload',
        'rec_workload_text': 'Top resolver has {count} violations. Consider workload balancing.',
        'rec_workload_actions': [
            'Review workload distribution across team members',
            'Implement workload caps and automatic redistribution',
            'Identify training needs for lower-performing resolvers'
        ],
        'generated_on': 'Generated on',
        'process': 'Process',
        'resource': 'Resource',
        'external': 'External',
        'timewindow': 'TimeWindow',
        'chart_sla_priority': 'SLA Compliance Rate by Priority',
        'chart_risk_dist': 'Risk Level Distribution',
        'chart_attribution': 'Violation Attribution Analysis',
        'chart_violations_cat': 'Top 10 Violations by Category',
        'chart_severity': 'Violation Severity Distribution',
        'chart_monthly': 'Monthly SLA Trend',
        'chart_p1p2': 'High Priority (P1/P2) SLA Performance',
        'compliance_rate': 'Compliance Rate (%)',
        'violation_count': 'Violation Count',
        'count': 'Count',
        'month': 'Month',
        'met': 'Met',
        'violated': 'Violated',
        'minor': 'Minor (<2h)',
        'moderate': 'Moderate (2-8h)',
        'severe': 'Severe (8-24h)',
        'critical': 'Critical (>24h)',
        'high': 'High',
        'medium': 'Medium',
        'low': 'Low',
    },
    'cn': {
        'title': 'SLA违约归因分析报告',
        'data_period': '数据周期',
        'executive_summary': '执行摘要',
        'total_tickets': '总工单数',
        'response_sla': '响应SLA',
        'resolution_sla': '解决SLA',
        'total_violations': '违约总数',
        'high_risk_tickets': '高风险工单',
        'critical_alert': '严重告警',
        'alert_text': '解决SLA ({resolution_rate:.1f}%) 显著低于95%目标。P1解决SLA仅{p1_rate:.1f}%，P2为{p2_rate:.1f}%。需立即采取行动。',
        'p1p2_analysis': '高优先级分析 (P1/P2)',
        'p1p2_desc': 'P1和P2事件需要立即关注。SLA目标：P1响应15分钟/解决2小时，P2响应30分钟/解决6小时。',
        'p1_performance': 'P1表现',
        'p2_performance': 'P2表现',
        'monthly_trend': '月度趋势分析',
        'sla_by_priority': '按优先级SLA达成率',
        'priority': '优先级',
        'response_target': '响应目标',
        'resolution_target': '解决目标',
        'total': '总计',
        'risk_analysis': '风险分析',
        'high_risk_top10': '高风险工单 (Top 10)',
        'high_risk_desc': '接近或超过SLA阈值的工单。"超出"显示超过SLA目标的小时数。',
        'ticket': '工单号',
        'category': '类别',
        'resolver': '处理人',
        'actual': '实际',
        'target': '目标',
        'exceed': '超出',
        'attribution': '归因',
        'violation_deep_dive': '违约深度分析',
        'top_violations': '按严重程度排序的违约',
        'resolution': '解决时长',
        'attribution_analysis': '归因分析',
        'attr_footnote': '注：归因总数 ({attr_total:,}) 可能与违约数 ({violation_count:,}) 不同，因为单个工单可能有多个归因因素。',
        'violations_by_category': '按类别违约分布',
        'recommendations': '改进建议',
        'rec_critical': '紧急',
        'rec_high': '高',
        'rec_medium': '中',
        'rec_p1p2_title': 'P1/P2 SLA表现',
        'rec_p1p2_text': 'P1解决SLA为{p1_rate:.1f}%（目标：95%）。P2为{p2_rate:.1f}%。',
        'rec_p1p2_actions': [
            '建立专门的P1/P2响应团队，确保<15分钟初始响应',
            '实施50% SLA消耗后自动升级机制',
            '创建P1战时协议，实现多团队协调',
            '审查P1/P2分类标准，确保正确优先级排序'
        ],
        'rec_timewindow_title': '时间窗口因素 (占违约{pct:.0f}%)',
        'rec_timewindow_text': '大多数违约发生在非工作时间（18:00-09:00）和周末。',
        'rec_timewindow_actions': [
            '加强18:00-09:00班次的值班覆盖',
            '实施周末轮值制度，配备专门升级路径',
            '考虑非工作时间工单的自动分诊和路由',
            '评估P1/P2事件的24/7 NOC覆盖'
        ],
        'rec_external_title': '外部依赖 (占违约{pct:.0f}%)',
        'rec_external_text': '等待外部团队或客户导致显著延迟。',
        'rec_external_actions': [
            '与依赖团队建立SLA协议（内部OLA）',
            '实施自动客户跟进提醒',
            '创建"等待外部"仪表板，主动管理',
            '定义外部等待超过阈值时的升级触发器'
        ],
        'rec_process_title': '流程改进 (占违约{pct:.0f}%)',
        'rec_process_text': '转派和升级延迟导致SLA违约。',
        'rec_process_actions': [
            '减少平均转派次数（目标：<2次/工单）',
            '实施基于技能的路由，减少转派',
            '创建清晰的升级矩阵，定义触发条件',
            '培训L1团队更好地进行初始分类'
        ],
        'rec_category_title': '类别聚焦：{category}',
        'rec_category_text': '该类别占{count:,}个违约（总数的{pct:.0f}%）。',
        'rec_category_actions': [
            '审查监控告警阈值，减少误报',
            '为常见场景创建操作手册',
            '考虑自动化重复性任务'
        ],
        'rec_workload_title': '处理人工作负荷',
        'rec_workload_text': '排名第一的处理人有{count}个违约。考虑工作负荷均衡。',
        'rec_workload_actions': [
            '审查团队成员间的工作负荷分配',
            '实施工作负荷上限和自动重分配',
            '识别表现较差处理人的培训需求'
        ],
        'generated_on': '生成时间',
        'process': '流程',
        'resource': '资源',
        'external': '外部',
        'timewindow': '时间窗',
        'chart_sla_priority': '按优先级SLA达成率',
        'chart_risk_dist': '风险等级分布',
        'chart_attribution': '违约归因分析',
        'chart_violations_cat': 'Top 10 类别违约',
        'chart_severity': '违约严重程度分布',
        'chart_monthly': '月度SLA趋势',
        'chart_p1p2': '高优先级 (P1/P2) SLA表现',
        'compliance_rate': '达成率 (%)',
        'violation_count': '违约数',
        'count': '数量',
        'month': '月份',
        'met': '达标',
        'violated': '违约',
        'minor': '轻微 (<2h)',
        'moderate': '中等 (2-8h)',
        'severe': '严重 (8-24h)',
        'critical': '紧急 (>24h)',
        'high': '高',
        'medium': '中',
        'low': '低',
    }
}

# Theme colors
COLORS = {
    'primary': '#1e40af',
    'primary_light': '#3b82f6',
    'danger': '#ef4444',
    'warning': '#f59e0b',
    'success': '#10b981',
    'gray': '#6b7280',
    'process': '#3b82f6',
    'resource': '#8b5cf6',
    'external': '#f59e0b',
    'timewindow': '#ef4444',
}

# Set matplotlib style
import matplotlib
matplotlib.rcParams['font.family'] = ['Heiti TC', 'STHeiti', 'PingFang HK', 'Arial Unicode MS', 'Arial', 'sans-serif']
matplotlib.rcParams['axes.unicode_minus'] = False
plt.rcParams['font.size'] = 11
plt.rcParams['axes.titlesize'] = 14
plt.rcParams['axes.titleweight'] = 'bold'
plt.rcParams['axes.labelsize'] = 11
plt.rcParams['axes.spines.top'] = False
plt.rcParams['axes.spines.right'] = False
plt.rcParams['figure.facecolor'] = 'white'
plt.rcParams['axes.facecolor'] = 'white'
plt.rcParams['axes.edgecolor'] = '#e5e7eb'

# ============ Load Data ============
xl = pd.ExcelFile('data/Incidents-exported.xlsx')
sla_criteria = pd.read_excel(xl, sheet_name=0)
df = pd.read_excel(xl, sheet_name=1)

# SLA lookup
SLA = {}
for _, row in sla_criteria.iterrows():
    SLA[row['Priority']] = {
        'response_min': row['Response （minutes）'],
        'resolution_hours': row['Resolution （hours）']
    }

# Calculate metrics
df['Resolution_Hours'] = df['Resolution Time(m)'] / 60
df['Response_Minutes'] = df['Response Time(m)']
df['SLA_Response_Target'] = df['Priority'].map(lambda x: SLA.get(x, {}).get('response_min', 60))
df['SLA_Resolution_Target'] = df['Priority'].map(lambda x: SLA.get(x, {}).get('resolution_hours', 48))
df['Response_SLA_Met'] = df['Response_Minutes'] <= df['SLA_Response_Target']
df['Resolution_SLA_Met'] = df['Resolution_Hours'] <= df['SLA_Resolution_Target']
df['Response_Exceed_Min'] = (df['Response_Minutes'] - df['SLA_Response_Target']).clip(lower=0)
df['Resolution_Exceed_Hours'] = (df['Resolution_Hours'] - df['SLA_Resolution_Target']).clip(lower=0)
df['Resolution_SLA_Pct'] = (df['Resolution_Hours'] / df['SLA_Resolution_Target'] * 100).clip(upper=999)
df['Risk_Level'] = df['Resolution_SLA_Pct'].apply(
    lambda x: 'High' if x >= 80 else ('Medium' if x >= 60 else 'Low')
)

# Attribution
def get_attribution(row):
    attrs = []
    if row['Resolution Time(m)'] > 0 and row['Total Time(m)'] > row['Resolution Time(m)'] * 1.2:
        attrs.append('Process')
    if row['Suspend Time(m)'] > row['Resolution Time(m)'] * 0.3:
        attrs.append('External')
    begin = pd.to_datetime(row['Begin Date'])
    if begin.hour >= 18 or begin.hour < 9 or begin.weekday() >= 5:
        attrs.append('TimeWindow')
    return attrs if attrs else ['Resource']

df['Attributions'] = df.apply(get_attribution, axis=1)
df['Primary_Attribution'] = df['Attributions'].apply(lambda x: x[0] if x else 'Unknown')
df['Begin_Date'] = pd.to_datetime(df['Begin Date'])
df['Month'] = df['Begin_Date'].dt.to_period('M')

violations = df[~df['Resolution_SLA_Met']].copy()
violations['Severity'] = violations['Resolution_Exceed_Hours'].apply(
    lambda h: 'Minor (<2h)' if h < 2 else ('Moderate (2-8h)' if h < 8 else ('Severe (8-24h)' if h < 24 else 'Critical (>24h)'))
)

# Statistics
stats = {
    'total_tickets': len(df),
    'response_rate': df['Response_SLA_Met'].mean() * 100,
    'resolution_rate': df['Resolution_SLA_Met'].mean() * 100,
    'total_violations': len(violations),
    'high_risk_count': len(df[df['Risk_Level'] == 'High']),
    'p1_resolution_rate': df[df['Priority'] == 'P1']['Resolution_SLA_Met'].mean() * 100,
    'p2_resolution_rate': df[df['Priority'] == 'P2']['Resolution_SLA_Met'].mean() * 100,
    'p1_response_rate': df[df['Priority'] == 'P1']['Response_SLA_Met'].mean() * 100,
    'p2_response_rate': df[df['Priority'] == 'P2']['Response_SLA_Met'].mean() * 100,
    'p1_count': len(df[df['Priority'] == 'P1']),
    'p2_count': len(df[df['Priority'] == 'P2']),
    'date_min': df['Begin_Date'].min().strftime('%Y-%m-%d'),
    'date_max': df['Begin_Date'].max().strftime('%Y-%m-%d'),
}

# Attribution counts
attr_counts = {}
for attrs in violations['Attributions']:
    for a in attrs:
        attr_counts[a] = attr_counts.get(a, 0) + 1
attr_total = sum(attr_counts.values())
stats['attr_counts'] = attr_counts
stats['attr_total'] = attr_total

# Top categories and resolvers
top_categories = violations.groupby('Category').size().nlargest(5)
top_resolvers = violations.groupby('Resolver').size().nlargest(5)
stats['top_category'] = top_categories.index[0] if len(top_categories) > 0 else 'Unknown'
stats['top_category_count'] = top_categories.iloc[0] if len(top_categories) > 0 else 0
stats['top_resolver_count'] = top_resolvers.iloc[0] if len(top_resolvers) > 0 else 0

# Output directory
os.makedirs('output/images', exist_ok=True)

# ============ Generate Charts (Bilingual) ============

def generate_charts(lang):
    t = LANG[lang]
    suffix = f'_{lang.upper()}'

    # 1. SLA by Priority
    fig, ax = plt.subplots(figsize=(8, 5))
    priority_stats = df.groupby('Priority').agg({
        'Resolution_SLA_Met': 'mean', 'Response_SLA_Met': 'mean'
    }).reindex(['P1', 'P2', 'P3', 'P4'])
    x = range(len(priority_stats))
    width = 0.35
    bars1 = ax.bar([i - width/2 for i in x], priority_stats['Response_SLA_Met'] * 100, width,
                   label=t['response_sla'], color=COLORS['primary_light'])
    bars2 = ax.bar([i + width/2 for i in x], priority_stats['Resolution_SLA_Met'] * 100, width,
                   label=t['resolution_sla'], color=COLORS['primary'])
    ax.set_ylabel(t['compliance_rate'])
    ax.set_xlabel(t['priority'])
    ax.set_title(t['chart_sla_priority'])
    ax.set_xticks(x)
    ax.set_xticklabels(['P1', 'P2', 'P3', 'P4'])
    ax.legend()
    ax.set_ylim(0, 105)
    for bar in bars1 + bars2:
        ax.annotate(f'{bar.get_height():.1f}%', xy=(bar.get_x() + bar.get_width()/2, bar.get_height()),
                    ha='center', va='bottom', fontsize=9)
    plt.tight_layout()
    plt.savefig(f'output/images/sla_by_priority{suffix}.png', dpi=150)
    plt.close()

    # 2. Risk Distribution
    fig, ax = plt.subplots(figsize=(7, 5))
    risk_labels = [t['high'], t['medium'], t['low']]
    risk_counts = df['Risk_Level'].value_counts().reindex(['High', 'Medium', 'Low'])
    colors_risk = [COLORS['danger'], COLORS['warning'], COLORS['success']]
    bars = ax.bar(risk_labels, risk_counts.values, color=colors_risk)
    ax.set_title(t['chart_risk_dist'])
    ax.set_ylabel(t['count'])
    for bar in bars:
        ax.annotate(f'{int(bar.get_height())}', xy=(bar.get_x() + bar.get_width()/2, bar.get_height()),
                    ha='center', va='bottom', fontweight='bold')
    plt.tight_layout()
    plt.savefig(f'output/images/risk_distribution{suffix}.png', dpi=150)
    plt.close()

    # 3. Attribution
    fig, ax = plt.subplots(figsize=(8, 5))
    attr_order = ['Process', 'Resource', 'External', 'TimeWindow']
    attr_labels_loc = [t['process'], t['resource'], t['external'], t['timewindow']]
    attr_values = [attr_counts.get(a, 0) for a in attr_order]
    attr_colors = [COLORS['process'], COLORS['resource'], COLORS['external'], COLORS['timewindow']]
    bars = ax.bar(attr_labels_loc, attr_values, color=attr_colors)
    ax.set_title(t['chart_attribution'])
    ax.set_ylabel(t['count'])
    for bar in bars:
        ax.annotate(f'{int(bar.get_height())}', xy=(bar.get_x() + bar.get_width()/2, bar.get_height()),
                    ha='center', va='bottom', fontweight='bold')
    plt.tight_layout()
    plt.savefig(f'output/images/attribution{suffix}.png', dpi=150)
    plt.close()

    # 4. Violations by Category
    fig, ax = plt.subplots(figsize=(10, 6))
    cat_violations = violations['Category'].value_counts().head(10)
    colors_cat = [COLORS['danger'] if i == 0 else COLORS['primary_light'] for i in range(len(cat_violations))]
    bars = ax.barh(cat_violations.index[::-1], cat_violations.values[::-1], color=colors_cat[::-1])
    ax.set_title(t['chart_violations_cat'])
    ax.set_xlabel(t['violation_count'])
    for bar in bars:
        ax.annotate(f'{int(bar.get_width())}', xy=(bar.get_width(), bar.get_y() + bar.get_height()/2),
                    ha='left', va='center', fontsize=9)
    plt.tight_layout()
    plt.savefig(f'output/images/violations_by_category{suffix}.png', dpi=150)
    plt.close()

    # 5. Severity Distribution
    fig, ax = plt.subplots(figsize=(7, 6))
    severity_order = ['Minor (<2h)', 'Moderate (2-8h)', 'Severe (8-24h)', 'Critical (>24h)']
    severity_labels_loc = [t['minor'], t['moderate'], t['severe'], t['critical']]
    severity_counts = violations['Severity'].value_counts().reindex(severity_order).fillna(0)
    colors_sev = [COLORS['success'], COLORS['warning'], COLORS['primary'], COLORS['danger']]
    ax.pie(severity_counts.values, labels=severity_labels_loc, autopct='%1.1f%%', colors=colors_sev, startangle=90)
    ax.set_title(t['chart_severity'])
    plt.tight_layout()
    plt.savefig(f'output/images/severity_distribution{suffix}.png', dpi=150)
    plt.close()

    # 6. Monthly Trend
    fig, ax = plt.subplots(figsize=(12, 5))
    monthly_stats = df.groupby('Month').agg({
        'Resolution_SLA_Met': ['sum', 'count'], 'Response_SLA_Met': 'sum'
    })
    monthly_stats.columns = ['Resolution_Met', 'Total', 'Response_Met']
    monthly_stats['Resolution_Rate'] = monthly_stats['Resolution_Met'] / monthly_stats['Total'] * 100
    monthly_stats['Response_Rate'] = monthly_stats['Response_Met'] / monthly_stats['Total'] * 100
    monthly_stats['Violations'] = monthly_stats['Total'] - monthly_stats['Resolution_Met']
    months = [p.to_timestamp() for p in monthly_stats.index]
    ax2 = ax.twinx()
    bars = ax2.bar(months, monthly_stats['Violations'], width=20, alpha=0.3, color=COLORS['danger'], label=t['total_violations'])
    line1, = ax.plot(months, monthly_stats['Resolution_Rate'], 'o-', color=COLORS['primary'], linewidth=2, label=t['resolution_sla'])
    line2, = ax.plot(months, monthly_stats['Response_Rate'], 's--', color=COLORS['success'], linewidth=2, label=t['response_sla'])
    ax.set_ylabel(t['compliance_rate'])
    ax2.set_ylabel(t['violation_count'])
    ax.set_xlabel(t['month'])
    ax.set_title(t['chart_monthly'])
    ax.set_ylim(0, 105)
    ax.axhline(y=95, color=COLORS['gray'], linestyle=':', alpha=0.7)
    ax.xaxis.set_major_formatter(mdates.DateFormatter('%Y-%m'))
    plt.xticks(rotation=45)
    ax.legend([line1, line2, bars], [t['resolution_sla'], t['response_sla'], t['total_violations']], loc='upper left')
    plt.tight_layout()
    plt.savefig(f'output/images/monthly_trend{suffix}.png', dpi=150)
    plt.close()

    # 7. P1/P2 Analysis
    fig, axes = plt.subplots(1, 2, figsize=(12, 5))
    p1_data = df[df['Priority'] == 'P1']
    p1_violations_count = len(p1_data[~p1_data['Resolution_SLA_Met']])
    axes[0].pie([p1_data['Resolution_SLA_Met'].sum(), p1_violations_count],
                labels=[t['met'], t['violated']], autopct='%1.1f%%',
                colors=[COLORS['success'], COLORS['danger']], startangle=90)
    axes[0].set_title(f'{t["p1_performance"]}\n(Target: 2h, Total: {len(p1_data)})')
    p2_data = df[df['Priority'] == 'P2']
    p2_violations_count = len(p2_data[~p2_data['Resolution_SLA_Met']])
    axes[1].pie([p2_data['Resolution_SLA_Met'].sum(), p2_violations_count],
                labels=[t['met'], t['violated']], autopct='%1.1f%%',
                colors=[COLORS['success'], COLORS['danger']], startangle=90)
    axes[1].set_title(f'{t["p2_performance"]}\n(Target: 6h, Total: {len(p2_data)})')
    plt.suptitle(t['chart_p1p2'], fontsize=14, fontweight='bold')
    plt.tight_layout()
    plt.savefig(f'output/images/p1_p2_analysis{suffix}.png', dpi=150)
    plt.close()

# ============ Generate HTML ============

def generate_html(lang):
    t = LANG[lang]
    suffix = f'_{lang.upper()}'

    # High risk tickets
    high_risk_tickets = df[df['Risk_Level'] == 'High'].nsmallest(10, 'Resolution_Exceed_Hours')[
        ['Order Number', 'Priority', 'Category', 'Resolver', 'Resolution_Hours', 'SLA_Resolution_Target', 'Resolution_Exceed_Hours', 'Primary_Attribution']
    ]

    # Top violations
    top_violations = violations.nlargest(15, 'Resolution_Exceed_Hours')[
        ['Order Number', 'Priority', 'Category', 'Resolver', 'Resolution_Hours', 'Resolution_Exceed_Hours', 'Primary_Attribution']
    ]

    html = f'''<!DOCTYPE html>
<html lang="{lang}">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{t['title']}</title>
    <style>
        * {{ margin: 0; padding: 0; box-sizing: border-box; }}
        body {{ font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, 'PingFang SC', 'Microsoft YaHei', sans-serif; background: #f3f4f6; color: #374151; line-height: 1.6; }}
        .container {{ max-width: 1200px; margin: 0 auto; padding: 20px; }}
        .header {{ background: linear-gradient(135deg, #1e40af, #3b82f6); color: white; padding: 30px; border-radius: 12px; margin-bottom: 24px; }}
        .header h1 {{ font-size: 28px; margin-bottom: 8px; }}
        .header .date {{ opacity: 0.9; }}
        .card {{ background: white; border-radius: 12px; padding: 24px; margin-bottom: 20px; box-shadow: 0 1px 3px rgba(0,0,0,0.1); }}
        .card h2 {{ font-size: 20px; color: #1e40af; margin-bottom: 16px; padding-bottom: 8px; border-bottom: 2px solid #e5e7eb; }}
        .card h3 {{ font-size: 16px; color: #374151; margin: 20px 0 12px; }}
        .metrics {{ display: grid; grid-template-columns: repeat(auto-fit, minmax(180px, 1fr)); gap: 16px; margin-bottom: 20px; }}
        .metric {{ background: #f9fafb; padding: 16px; border-radius: 8px; text-align: center; }}
        .metric .value {{ font-size: 28px; font-weight: bold; color: #1e40af; }}
        .metric .label {{ color: #6b7280; font-size: 13px; margin-top: 4px; }}
        .metric.danger .value {{ color: #ef4444; }}
        .metric.success .value {{ color: #10b981; }}
        table {{ width: 100%; border-collapse: collapse; margin: 16px 0; font-size: 14px; }}
        th, td {{ padding: 10px 12px; text-align: left; border-bottom: 1px solid #e5e7eb; }}
        th {{ background: #f9fafb; font-weight: 600; color: #374151; }}
        tr:hover {{ background: #f9fafb; }}
        .badge {{ display: inline-block; padding: 4px 10px; border-radius: 20px; font-size: 11px; font-weight: 500; }}
        .badge.high, .badge.critical {{ background: #fee2e2; color: #991b1b; }}
        .badge.medium, .badge.severe {{ background: #fef3c7; color: #92400e; }}
        .badge.low, .badge.minor {{ background: #d1fae5; color: #065f46; }}
        .badge.process {{ background: #dbeafe; color: #1e40af; }}
        .badge.resource {{ background: #ede9fe; color: #6b21a8; }}
        .badge.external {{ background: #fef3c7; color: #92400e; }}
        .badge.timewindow {{ background: #fee2e2; color: #991b1b; }}
        .chart-container {{ text-align: center; margin: 20px 0; }}
        .chart-container img {{ max-width: 100%; height: auto; border-radius: 8px; }}
        .recommendation {{ padding: 16px; border-left: 4px solid #3b82f6; background: #f9fafb; margin: 12px 0; border-radius: 0 8px 8px 0; }}
        .recommendation.critical {{ border-left-color: #ef4444; background: #fef2f2; }}
        .recommendation.high {{ border-left-color: #f59e0b; background: #fffbeb; }}
        .recommendation strong {{ color: #374151; }}
        .recommendation ul {{ margin: 8px 0 0 20px; }}
        .recommendation li {{ margin: 4px 0; }}
        .grid-2 {{ display: grid; grid-template-columns: 1fr 1fr; gap: 20px; }}
        @media (max-width: 768px) {{ .grid-2 {{ grid-template-columns: 1fr; }} }}
        .footnote {{ font-size: 12px; color: #6b7280; margin-top: 12px; padding-top: 12px; border-top: 1px dashed #e5e7eb; }}
        .alert {{ padding: 12px 16px; border-radius: 8px; margin: 16px 0; }}
        .alert.danger {{ background: #fef2f2; border: 1px solid #fecaca; color: #991b1b; }}
        .footer {{ text-align: center; padding: 20px; color: #6b7280; font-size: 14px; }}
    </style>
</head>
<body>
    <div class="container">
        <div class="header">
            <h1>{t['title']}</h1>
            <div class="date">{t['data_period']}: {stats['date_min']} ~ {stats['date_max']}</div>
        </div>

        <div class="card">
            <h2>{t['executive_summary']}</h2>
            <div class="metrics">
                <div class="metric"><div class="value">{stats['total_tickets']:,}</div><div class="label">{t['total_tickets']}</div></div>
                <div class="metric {'success' if stats['response_rate'] >= 95 else 'danger'}"><div class="value">{stats['response_rate']:.1f}%</div><div class="label">{t['response_sla']}</div></div>
                <div class="metric {'success' if stats['resolution_rate'] >= 95 else 'danger'}"><div class="value">{stats['resolution_rate']:.1f}%</div><div class="label">{t['resolution_sla']}</div></div>
                <div class="metric danger"><div class="value">{stats['total_violations']:,}</div><div class="label">{t['total_violations']}</div></div>
                <div class="metric danger"><div class="value">{stats['high_risk_count']:,}</div><div class="label">{t['high_risk_tickets']}</div></div>
            </div>
            <div class="alert danger"><strong>{t['critical_alert']}:</strong> {t['alert_text'].format(resolution_rate=stats['resolution_rate'], p1_rate=stats['p1_resolution_rate'], p2_rate=stats['p2_resolution_rate'])}</div>
        </div>

        <div class="card">
            <h2>{t['p1p2_analysis']}</h2>
            <p style="color: #6b7280; margin-bottom: 16px;">{t['p1p2_desc']}</p>
            <div class="grid-2">
                <div>
                    <h3>{t['p1_performance']} (Total: {stats['p1_count']})</h3>
                    <div class="metrics">
                        <div class="metric {'success' if stats['p1_response_rate'] >= 95 else 'danger'}"><div class="value">{stats['p1_response_rate']:.1f}%</div><div class="label">{t['response_sla']}</div></div>
                        <div class="metric {'success' if stats['p1_resolution_rate'] >= 95 else 'danger'}"><div class="value">{stats['p1_resolution_rate']:.1f}%</div><div class="label">{t['resolution_sla']}</div></div>
                    </div>
                </div>
                <div>
                    <h3>{t['p2_performance']} (Total: {stats['p2_count']})</h3>
                    <div class="metrics">
                        <div class="metric {'success' if stats['p2_response_rate'] >= 95 else 'danger'}"><div class="value">{stats['p2_response_rate']:.1f}%</div><div class="label">{t['response_sla']}</div></div>
                        <div class="metric {'success' if stats['p2_resolution_rate'] >= 95 else 'danger'}"><div class="value">{stats['p2_resolution_rate']:.1f}%</div><div class="label">{t['resolution_sla']}</div></div>
                    </div>
                </div>
            </div>
            <div class="chart-container"><img src="images/p1_p2_analysis{suffix}.png" alt="P1/P2 Analysis"></div>
        </div>

        <div class="card">
            <h2>{t['monthly_trend']}</h2>
            <div class="chart-container"><img src="images/monthly_trend{suffix}.png" alt="Monthly Trend"></div>
        </div>

        <div class="card">
            <h2>{t['sla_by_priority']}</h2>
            <div class="chart-container"><img src="images/sla_by_priority{suffix}.png" alt="SLA by Priority"></div>
            <table>
                <thead><tr><th>{t['priority']}</th><th>{t['response_target']}</th><th>{t['response_sla']}</th><th>{t['resolution_target']}</th><th>{t['resolution_sla']}</th><th>{t['total']}</th></tr></thead>
                <tbody>'''

    for priority in ['P1', 'P2', 'P3', 'P4']:
        p_data = df[df['Priority'] == priority]
        resp_rate = p_data['Response_SLA_Met'].mean() * 100
        res_rate = p_data['Resolution_SLA_Met'].mean() * 100
        html += f'''<tr><td><strong>{priority}</strong></td><td>{SLA[priority]['response_min']} min</td><td><span class="badge {'low' if resp_rate >= 95 else 'high'}">{resp_rate:.1f}%</span></td><td>{SLA[priority]['resolution_hours']} h</td><td><span class="badge {'low' if res_rate >= 95 else 'high'}">{res_rate:.1f}%</span></td><td>{len(p_data):,}</td></tr>'''

    html += f'''</tbody></table></div>

        <div class="card">
            <h2>{t['risk_analysis']}</h2>
            <div class="chart-container"><img src="images/risk_distribution{suffix}.png" alt="Risk Distribution"></div>
            <h3>{t['high_risk_top10']}</h3>
            <p style="color: #6b7280; font-size: 13px; margin-bottom: 12px;">{t['high_risk_desc']}</p>
            <table>
                <thead><tr><th>{t['ticket']}</th><th>{t['priority']}</th><th>{t['category']}</th><th>{t['resolver']}</th><th>{t['actual']}</th><th>{t['target']}</th><th>{t['exceed']}</th><th>{t['attribution']}</th></tr></thead>
                <tbody>'''

    for _, row in high_risk_tickets.iterrows():
        exceed = row['Resolution_Exceed_Hours']
        attr = row['Primary_Attribution'].lower()
        html += f'''<tr><td>{row['Order Number']}</td><td>{row['Priority']}</td><td>{row['Category']}</td><td>{row['Resolver']}</td><td>{row['Resolution_Hours']:.1f}h</td><td>{row['SLA_Resolution_Target']:.0f}h</td><td><span class="badge {'high' if exceed > 0 else 'low'}">{'+' if exceed > 0 else ''}{exceed:.1f}h</span></td><td><span class="badge {attr}">{row['Primary_Attribution']}</span></td></tr>'''

    html += f'''</tbody></table></div>

        <div class="card">
            <h2>{t['violation_deep_dive']}</h2>
            <div class="chart-container"><img src="images/severity_distribution{suffix}.png" alt="Severity"></div>
            <h3>{t['top_violations']} (Total: {stats['total_violations']:,})</h3>
            <table>
                <thead><tr><th>{t['ticket']}</th><th>{t['priority']}</th><th>{t['category']}</th><th>{t['resolver']}</th><th>{t['resolution']}</th><th>{t['exceed']}</th><th>{t['attribution']}</th></tr></thead>
                <tbody>'''

    for _, row in top_violations.iterrows():
        exceed = row['Resolution_Exceed_Hours']
        severity = 'critical' if exceed > 24 else ('severe' if exceed > 8 else ('medium' if exceed > 2 else 'minor'))
        attr = row['Primary_Attribution'].lower()
        html += f'''<tr><td>{row['Order Number']}</td><td>{row['Priority']}</td><td>{row['Category']}</td><td>{row['Resolver']}</td><td>{row['Resolution_Hours']:.1f}h</td><td><span class="badge {severity}">+{exceed:.1f}h</span></td><td><span class="badge {attr}">{row['Primary_Attribution']}</span></td></tr>'''

    attr_total = stats['attr_total']
    html += f'''</tbody></table></div>

        <div class="card">
            <h2>{t['attribution_analysis']}</h2>
            <div class="chart-container"><img src="images/attribution{suffix}.png" alt="Attribution"></div>
            <div class="metrics">'''

    for attr_name, attr_key in [('Process', 'process'), ('Resource', 'resource'), ('External', 'external'), ('TimeWindow', 'timewindow')]:
        count = attr_counts.get(attr_name, 0)
        pct = count / attr_total * 100 if attr_total > 0 else 0
        html += f'''<div class="metric"><div class="value" style="color: {COLORS[attr_key]};">{count:,}</div><div class="label">{t[attr_key]} ({pct:.0f}%)</div></div>'''

    html += f'''</div>
            <div class="footnote">{t['attr_footnote'].format(attr_total=attr_total, violation_count=stats['total_violations'])}</div>
        </div>

        <div class="card">
            <h2>{t['violations_by_category']}</h2>
            <div class="chart-container"><img src="images/violations_by_category{suffix}.png" alt="Category"></div>
        </div>

        <div class="card">
            <h2>{t['recommendations']}</h2>'''

    # Recommendations
    tw_pct = attr_counts.get('TimeWindow', 0) / attr_total * 100 if attr_total > 0 else 0
    ext_pct = attr_counts.get('External', 0) / attr_total * 100 if attr_total > 0 else 0
    proc_pct = attr_counts.get('Process', 0) / attr_total * 100 if attr_total > 0 else 0
    cat_pct = stats['top_category_count'] / stats['total_violations'] * 100 if stats['total_violations'] > 0 else 0

    recs = [
        ('critical', t['rec_p1p2_title'], t['rec_p1p2_text'].format(p1_rate=stats['p1_resolution_rate'], p2_rate=stats['p2_resolution_rate']), t['rec_p1p2_actions']),
        ('critical', t['rec_timewindow_title'].format(pct=tw_pct), t['rec_timewindow_text'], t['rec_timewindow_actions']),
        ('high', t['rec_external_title'].format(pct=ext_pct), t['rec_external_text'], t['rec_external_actions']),
        ('high', t['rec_process_title'].format(pct=proc_pct), t['rec_process_text'], t['rec_process_actions']),
        ('recommendation', t['rec_category_title'].format(category=stats['top_category']), t['rec_category_text'].format(count=stats['top_category_count'], pct=cat_pct), t['rec_category_actions']),
        ('recommendation', t['rec_workload_title'], t['rec_workload_text'].format(count=stats['top_resolver_count']), t['rec_workload_actions']),
    ]

    for level, title, text, actions in recs:
        html += f'''<div class="recommendation {level}"><strong>[{t['rec_critical'] if level == 'critical' else (t['rec_high'] if level == 'high' else t['rec_medium'])}] {title}</strong><p>{text}</p><ul>'''
        for action in actions:
            html += f'<li>{action}</li>'
        html += '</ul></div>'

    html += f'''</div>

        <div class="footer">{t['generated_on']} {datetime.now().strftime('%Y-%m-%d %H:%M')} | {t['title']}</div>
    </div>
</body>
</html>'''

    with open(f'output/SLA_Violation_Analysis_Report{suffix}.html', 'w', encoding='utf-8') as f:
        f.write(html)

# ============ Generate DOCX ============

def generate_docx(lang):
    t = LANG[lang]
    suffix = f'_{lang.upper()}'

    doc = Document()

    # Title
    title = doc.add_heading(t['title'], 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Date
    date_para = doc.add_paragraph(f"{t['data_period']}: {stats['date_min']} ~ {stats['date_max']}")
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Executive Summary
    doc.add_heading(t['executive_summary'], 1)

    summary_table = doc.add_table(rows=2, cols=5)
    summary_table.style = 'Table Grid'
    headers = [t['total_tickets'], t['response_sla'], t['resolution_sla'], t['total_violations'], t['high_risk_tickets']]
    values = [f"{stats['total_tickets']:,}", f"{stats['response_rate']:.1f}%", f"{stats['resolution_rate']:.1f}%",
              f"{stats['total_violations']:,}", f"{stats['high_risk_count']:,}"]

    for i, header in enumerate(headers):
        summary_table.cell(0, i).text = header
        summary_table.cell(1, i).text = values[i]

    doc.add_paragraph()
    alert = doc.add_paragraph()
    alert.add_run(f"{t['critical_alert']}: ").bold = True
    alert.add_run(t['alert_text'].format(resolution_rate=stats['resolution_rate'], p1_rate=stats['p1_resolution_rate'], p2_rate=stats['p2_resolution_rate']))

    # P1/P2 Analysis
    doc.add_heading(t['p1p2_analysis'], 1)
    doc.add_paragraph(t['p1p2_desc'])

    p1p2_table = doc.add_table(rows=3, cols=3)
    p1p2_table.style = 'Table Grid'
    p1p2_table.cell(0, 0).text = ''
    p1p2_table.cell(0, 1).text = t['response_sla']
    p1p2_table.cell(0, 2).text = t['resolution_sla']
    p1p2_table.cell(1, 0).text = f"P1 (n={stats['p1_count']})"
    p1p2_table.cell(1, 1).text = f"{stats['p1_response_rate']:.1f}%"
    p1p2_table.cell(1, 2).text = f"{stats['p1_resolution_rate']:.1f}%"
    p1p2_table.cell(2, 0).text = f"P2 (n={stats['p2_count']})"
    p1p2_table.cell(2, 1).text = f"{stats['p2_response_rate']:.1f}%"
    p1p2_table.cell(2, 2).text = f"{stats['p2_resolution_rate']:.1f}%"

    doc.add_paragraph()
    doc.add_picture(f'output/images/p1_p2_analysis{suffix}.png', width=Inches(6))

    # Monthly Trend
    doc.add_heading(t['monthly_trend'], 1)
    doc.add_picture(f'output/images/monthly_trend{suffix}.png', width=Inches(6))

    # SLA by Priority
    doc.add_heading(t['sla_by_priority'], 1)
    doc.add_picture(f'output/images/sla_by_priority{suffix}.png', width=Inches(5.5))

    sla_table = doc.add_table(rows=5, cols=6)
    sla_table.style = 'Table Grid'
    sla_headers = [t['priority'], t['response_target'], t['response_sla'], t['resolution_target'], t['resolution_sla'], t['total']]
    for i, h in enumerate(sla_headers):
        sla_table.cell(0, i).text = h

    for idx, priority in enumerate(['P1', 'P2', 'P3', 'P4'], 1):
        p_data = df[df['Priority'] == priority]
        resp_rate = p_data['Response_SLA_Met'].mean() * 100
        res_rate = p_data['Resolution_SLA_Met'].mean() * 100
        sla_table.cell(idx, 0).text = priority
        sla_table.cell(idx, 1).text = f"{SLA[priority]['response_min']} min"
        sla_table.cell(idx, 2).text = f"{resp_rate:.1f}%"
        sla_table.cell(idx, 3).text = f"{SLA[priority]['resolution_hours']} h"
        sla_table.cell(idx, 4).text = f"{res_rate:.1f}%"
        sla_table.cell(idx, 5).text = f"{len(p_data):,}"

    # Risk Analysis
    doc.add_heading(t['risk_analysis'], 1)
    doc.add_picture(f'output/images/risk_distribution{suffix}.png', width=Inches(5))

    # Attribution Analysis
    doc.add_heading(t['attribution_analysis'], 1)
    doc.add_picture(f'output/images/attribution{suffix}.png', width=Inches(5.5))

    attr_table = doc.add_table(rows=2, cols=4)
    attr_table.style = 'Table Grid'
    attr_names = [t['process'], t['resource'], t['external'], t['timewindow']]
    attr_keys = ['Process', 'Resource', 'External', 'TimeWindow']
    for i, (name, key) in enumerate(zip(attr_names, attr_keys)):
        count = attr_counts.get(key, 0)
        pct = count / attr_total * 100 if attr_total > 0 else 0
        attr_table.cell(0, i).text = name
        attr_table.cell(1, i).text = f"{count:,} ({pct:.0f}%)"

    doc.add_paragraph(t['attr_footnote'].format(attr_total=attr_total, violation_count=stats['total_violations']))

    # Severity Distribution
    doc.add_heading(t['violation_deep_dive'], 1)
    doc.add_picture(f'output/images/severity_distribution{suffix}.png', width=Inches(5))

    # Violations by Category
    doc.add_heading(t['violations_by_category'], 1)
    doc.add_picture(f'output/images/violations_by_category{suffix}.png', width=Inches(6))

    # Recommendations
    doc.add_heading(t['recommendations'], 1)

    tw_pct = attr_counts.get('TimeWindow', 0) / attr_total * 100 if attr_total > 0 else 0
    ext_pct = attr_counts.get('External', 0) / attr_total * 100 if attr_total > 0 else 0
    proc_pct = attr_counts.get('Process', 0) / attr_total * 100 if attr_total > 0 else 0
    cat_pct = stats['top_category_count'] / stats['total_violations'] * 100 if stats['total_violations'] > 0 else 0

    recs = [
        (t['rec_critical'], t['rec_p1p2_title'], t['rec_p1p2_text'].format(p1_rate=stats['p1_resolution_rate'], p2_rate=stats['p2_resolution_rate']), t['rec_p1p2_actions']),
        (t['rec_critical'], t['rec_timewindow_title'].format(pct=tw_pct), t['rec_timewindow_text'], t['rec_timewindow_actions']),
        (t['rec_high'], t['rec_external_title'].format(pct=ext_pct), t['rec_external_text'], t['rec_external_actions']),
        (t['rec_high'], t['rec_process_title'].format(pct=proc_pct), t['rec_process_text'], t['rec_process_actions']),
        (t['rec_medium'], t['rec_category_title'].format(category=stats['top_category']), t['rec_category_text'].format(count=stats['top_category_count'], pct=cat_pct), t['rec_category_actions']),
        (t['rec_medium'], t['rec_workload_title'], t['rec_workload_text'].format(count=stats['top_resolver_count']), t['rec_workload_actions']),
    ]

    for level, title, text, actions in recs:
        para = doc.add_paragraph()
        para.add_run(f"[{level}] {title}").bold = True
        doc.add_paragraph(text)
        for action in actions:
            doc.add_paragraph(f"• {action}")
        doc.add_paragraph()

    # Footer
    footer = doc.add_paragraph(f"{t['generated_on']} {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(f'output/SLA_Violation_Analysis_Report{suffix}.docx')

# ============ Main ============

if __name__ == '__main__':
    print("Generating bilingual charts...")
    generate_charts('en')
    generate_charts('cn')

    print("Generating HTML reports...")
    generate_html('en')
    generate_html('cn')

    print("Generating DOCX reports...")
    generate_docx('en')
    generate_docx('cn')

    print("\nOutput files generated:")
    print("  - output/SLA_Violation_Analysis_Report_EN.html")
    print("  - output/SLA_Violation_Analysis_Report_CN.html")
    print("  - output/SLA_Violation_Analysis_Report_EN.docx")
    print("  - output/SLA_Violation_Analysis_Report_CN.docx")
    print(f"\nKey Statistics:")
    print(f"  Total Tickets: {stats['total_tickets']:,}")
    print(f"  Response SLA: {stats['response_rate']:.1f}%")
    print(f"  Resolution SLA: {stats['resolution_rate']:.1f}%")
    print(f"  P1 Resolution SLA: {stats['p1_resolution_rate']:.1f}%")
    print(f"  P2 Resolution SLA: {stats['p2_resolution_rate']:.1f}%")
