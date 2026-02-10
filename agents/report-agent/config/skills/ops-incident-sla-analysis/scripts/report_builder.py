"""
Report Builder for SLA Violation Analysis.
Generates HTML and DOCX reports.
"""
import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from utils import OUTPUT_DIR, IMG_DIR, COLORS
from i18n import get_text


class ReportBuilder:
    def __init__(self, results, charts, language='en'):
        self.results = results
        self.charts = charts
        self.lang = language
        self.t = lambda key: get_text(key, language)

    def build_html(self):
        """Build HTML report."""
        overview = self.results['overview']
        risk = self.results['risk']
        violations = self.results['violations']
        attribution = self.results['attribution']
        recommendations = self.results['recommendations']

        # Generate date suffix
        date_suffix = overview['date_range'].replace(' ~ ', '_to_').replace('-', '')
        lang_suffix = 'EN' if self.lang == 'en' else 'CN'
        filename = f"SLA_Violation_Analysis_{date_suffix}_{lang_suffix}.html"
        filepath = os.path.join(OUTPUT_DIR, filename)

        html = f"""<!DOCTYPE html>
<html lang="{self.lang}">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{self.t('report_title')}</title>
    <style>
        * {{ margin: 0; padding: 0; box-sizing: border-box; }}
        body {{ font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif; background: #f3f4f6; color: #374151; line-height: 1.6; }}
        .container {{ max-width: 1200px; margin: 0 auto; padding: 20px; }}
        .header {{ background: linear-gradient(135deg, #1e40af, #3b82f6); color: white; padding: 30px; border-radius: 12px; margin-bottom: 24px; }}
        .header h1 {{ font-size: 28px; margin-bottom: 8px; }}
        .header .date {{ opacity: 0.9; }}
        .card {{ background: white; border-radius: 12px; padding: 24px; margin-bottom: 20px; box-shadow: 0 1px 3px rgba(0,0,0,0.1); }}
        .card h2 {{ font-size: 20px; color: #1e40af; margin-bottom: 16px; padding-bottom: 8px; border-bottom: 2px solid #e5e7eb; }}
        .metrics {{ display: grid; grid-template-columns: repeat(auto-fit, minmax(200px, 1fr)); gap: 16px; margin-bottom: 20px; }}
        .metric {{ background: #f9fafb; padding: 16px; border-radius: 8px; text-align: center; }}
        .metric .value {{ font-size: 32px; font-weight: bold; color: #1e40af; }}
        .metric .label {{ color: #6b7280; font-size: 14px; }}
        .metric.danger .value {{ color: #ef4444; }}
        .metric.warning .value {{ color: #f59e0b; }}
        .metric.success .value {{ color: #10b981; }}
        table {{ width: 100%; border-collapse: collapse; margin: 16px 0; }}
        th, td {{ padding: 12px; text-align: left; border-bottom: 1px solid #e5e7eb; }}
        th {{ background: #f9fafb; font-weight: 600; color: #374151; }}
        tr:hover {{ background: #f9fafb; }}
        .badge {{ display: inline-block; padding: 4px 12px; border-radius: 20px; font-size: 12px; font-weight: 500; }}
        .badge.high {{ background: #fee2e2; color: #991b1b; }}
        .badge.medium {{ background: #fef3c7; color: #92400e; }}
        .badge.low {{ background: #d1fae5; color: #065f46; }}
        .chart-container {{ text-align: center; margin: 20px 0; }}
        .chart-container img {{ max-width: 100%; height: auto; border-radius: 8px; }}
        .recommendation {{ padding: 16px; border-left: 4px solid #3b82f6; background: #f9fafb; margin: 12px 0; border-radius: 0 8px 8px 0; }}
        .recommendation.high {{ border-left-color: #ef4444; }}
        .recommendation.medium {{ border-left-color: #f59e0b; }}
        .footer {{ text-align: center; padding: 20px; color: #6b7280; font-size: 14px; }}
    </style>
</head>
<body>
    <div class="container">
        <div class="header">
            <h1>{self.t('report_title')}</h1>
            <div class="date">{self.t('data_period')}: {overview['date_range']}</div>
        </div>

        <!-- SLA Overview -->
        <div class="card">
            <h2>{self.t('sla_overview')}</h2>
            <div class="metrics">
                <div class="metric">
                    <div class="value">{overview['total']:,}</div>
                    <div class="label">{self.t('total_tickets')}</div>
                </div>
                <div class="metric {'danger' if overview['res_rate'] < 90 else 'warning' if overview['res_rate'] < 95 else 'success'}">
                    <div class="value">{overview['res_rate']:.1f}%</div>
                    <div class="label">{self.t('resolution_sla')}</div>
                </div>
                <div class="metric danger">
                    <div class="value">{overview['res_violations']}</div>
                    <div class="label">{self.t('total_violations')}</div>
                </div>
                <div class="metric {'danger' if risk['high_risk_count'] > 10 else 'warning' if risk['high_risk_count'] > 5 else ''}">
                    <div class="value">{risk['high_risk_count']}</div>
                    <div class="label">{self.t('high_risk_count')}</div>
                </div>
            </div>
            {self._render_chart('sla_by_priority')}
        </div>

        <!-- Risk Analysis -->
        <div class="card">
            <h2>{self.t('risk_analysis')}</h2>
            {self._render_chart('risk_distribution')}
            <h3 style="margin: 20px 0 12px; color: #374151;">{self.t('high_risk')} ({risk['high_risk_count']})</h3>
            {self._render_risk_table(risk['high_risk_list'][:10])}
        </div>

        <!-- Violation Deep Dive -->
        <div class="card">
            <h2>{self.t('violation_deep_dive')}</h2>
            {self._render_chart('severity_distribution')}
            <h3 style="margin: 20px 0 12px; color: #374151;">{self.t('total_violations')}: {violations['total']}</h3>
            {self._render_violation_table(violations['violation_list'][:15])}
        </div>

        <!-- Attribution Analysis -->
        <div class="card">
            <h2>{self.t('attribution_analysis')}</h2>
            {self._render_chart('attribution')}
            {self._render_attribution_summary(attribution)}
        </div>

        <!-- Recommendations -->
        <div class="card">
            <h2>{self.t('recommendations')}</h2>
            {self._render_recommendations(recommendations)}
        </div>

        <div class="footer">
            Generated on {datetime.now().strftime('%Y-%m-%d %H:%M')} | SLA Violation Analysis Report
        </div>
    </div>
</body>
</html>"""

        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(html)

        return filepath

    def _render_chart(self, chart_name):
        """Render chart if exists."""
        if chart_name in self.charts and self.charts[chart_name]:
            img_path = os.path.basename(self.charts[chart_name])
            return f'<div class="chart-container"><img src="images/{img_path}" alt="{chart_name}"></div>'
        return ''

    def _render_risk_table(self, items):
        """Render high risk tickets table."""
        if not items:
            return '<p>No high risk tickets.</p>'

        rows = ''
        for item in items:
            rows += f"""<tr>
                <td>{item['Order Number']}</td>
                <td>{item['Priority']}</td>
                <td>{item['Category']}</td>
                <td>{item['Resolver']}</td>
                <td>{item['SLA_Res_Pct']:.0f}%</td>
            </tr>"""

        return f"""<table>
            <thead><tr>
                <th>{self.t('order_number')}</th>
                <th>{self.t('priority')}</th>
                <th>{self.t('category')}</th>
                <th>{self.t('resolver')}</th>
                <th>SLA %</th>
            </tr></thead>
            <tbody>{rows}</tbody>
        </table>"""

    def _render_violation_table(self, items):
        """Render violation details table."""
        if not items:
            return '<p>No violations.</p>'

        rows = ''
        for item in items:
            severity_class = 'high' if 'Critical' in item['severity'] or 'Severe' in item['severity'] else 'medium' if 'Moderate' in item['severity'] else 'low'
            rows += f"""<tr>
                <td>{item['order_number']}</td>
                <td>{item['priority']}</td>
                <td>{item['category']}</td>
                <td>{item['resolver']}</td>
                <td>{item['resolution_hours']:.1f}h</td>
                <td>{item['overage_hours']:.1f}h</td>
                <td><span class="badge {severity_class}">{item['severity']}</span></td>
            </tr>"""

        return f"""<table>
            <thead><tr>
                <th>{self.t('order_number')}</th>
                <th>{self.t('priority')}</th>
                <th>{self.t('category')}</th>
                <th>{self.t('resolver')}</th>
                <th>{self.t('resolution_time')}</th>
                <th>{self.t('overage')}</th>
                <th>Severity</th>
            </tr></thead>
            <tbody>{rows}</tbody>
        </table>"""

    def _render_attribution_summary(self, attribution):
        """Render attribution summary."""
        total = sum(attribution.values())
        if total == 0:
            return '<p>No attribution data.</p>'

        items = ''
        for key, val in attribution.items():
            pct = val / total * 100 if total > 0 else 0
            label = self.t(f'attribution_{key.lower()}') if self.lang == 'zh' else key
            items += f'<div class="metric"><div class="value">{val}</div><div class="label">{label} ({pct:.0f}%)</div></div>'

        return f'<div class="metrics">{items}</div>'

    def _render_recommendations(self, recommendations):
        """Render recommendations."""
        if not recommendations:
            return '<p>No recommendations.</p>'

        html = ''
        for rec in recommendations:
            priority_class = rec['priority'].lower()
            text = rec.get('text_zh' if self.lang == 'zh' else 'text', rec['text'])
            html += f'<div class="recommendation {priority_class}"><strong>[{rec["priority"]}]</strong> {text}</div>'

        return html

    def build_docx(self):
        """Build DOCX report."""
        overview = self.results['overview']
        violations = self.results['violations']
        recommendations = self.results['recommendations']

        # Generate filename
        date_suffix = overview['date_range'].replace(' ~ ', '_to_').replace('-', '')
        lang_suffix = 'EN' if self.lang == 'en' else 'CN'
        filename = f"SLA_Violation_Analysis_{date_suffix}_{lang_suffix}.docx"
        filepath = os.path.join(OUTPUT_DIR, filename)

        doc = Document()

        # Title
        title = doc.add_heading(self.t('report_title'), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Date
        date_para = doc.add_paragraph(f"{self.t('data_period')}: {overview['date_range']}")
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()

        # Overview section
        doc.add_heading(self.t('sla_overview'), level=1)

        table = doc.add_table(rows=2, cols=4)
        table.style = 'Table Grid'

        headers = [self.t('total_tickets'), self.t('resolution_sla'), self.t('total_violations'), self.t('high_risk_count')]
        values = [str(overview['total']), f"{overview['res_rate']:.1f}%", str(overview['res_violations']), str(self.results['risk']['high_risk_count'])]

        for i, header in enumerate(headers):
            table.rows[0].cells[i].text = header
        for i, value in enumerate(values):
            table.rows[1].cells[i].text = value

        doc.add_paragraph()

        # Add chart if exists
        if 'sla_by_priority' in self.charts and self.charts['sla_by_priority']:
            doc.add_picture(self.charts['sla_by_priority'], width=Inches(6))

        # Violations section
        doc.add_heading(self.t('violation_deep_dive'), level=1)
        doc.add_paragraph(f"{self.t('total_violations')}: {violations['total']}")

        if violations['violation_list']:
            table = doc.add_table(rows=min(len(violations['violation_list']), 10) + 1, cols=5)
            table.style = 'Table Grid'

            headers = [self.t('order_number'), self.t('priority'), self.t('category'), self.t('resolution_time'), self.t('overage')]
            for i, h in enumerate(headers):
                table.rows[0].cells[i].text = h

            for i, v in enumerate(violations['violation_list'][:10]):
                table.rows[i+1].cells[0].text = str(v['order_number'])
                table.rows[i+1].cells[1].text = str(v['priority'])
                table.rows[i+1].cells[2].text = str(v['category'])
                table.rows[i+1].cells[3].text = f"{v['resolution_hours']:.1f}h"
                table.rows[i+1].cells[4].text = f"{v['overage_hours']:.1f}h"

        doc.add_paragraph()

        # Recommendations
        doc.add_heading(self.t('recommendations'), level=1)
        for rec in recommendations:
            text = rec.get('text_zh' if self.lang == 'zh' else 'text', rec['text'])
            doc.add_paragraph(f"[{rec['priority']}] {text}")

        doc.save(filepath)
        return filepath
