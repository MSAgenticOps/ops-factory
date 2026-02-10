#!/usr/bin/env python3
"""
Contract Risk Review Tool
合同风险审核工具

Features:
- Parse DOCX contracts and extract sections
- Evaluate sections against baseline using LLM
- Generate color-coded annotated DOCX (red=high, orange=medium, blue=low)
- Generate executive summary PDF with risk statistics
- Generate side-by-side comparison HTML view
- Support bilingual output (zh/en)
- Confidence scoring with review markers
"""
from __future__ import annotations

import argparse
import hashlib
import json
import logging
import os
import re
from concurrent.futures import ThreadPoolExecutor, as_completed
from dataclasses import dataclass, field
from datetime import datetime, timezone
from difflib import SequenceMatcher
from pathlib import Path
from typing import Any, Iterable, List, Mapping, Optional, Sequence

import fitz  # PyMuPDF
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor
from openai import OpenAI

# ============================================================================
# Configuration
# ============================================================================

SCRIPT_DIR = Path(__file__).parent
DEFAULT_DATA_DIR = SCRIPT_DIR.parent / "data"
DEFAULT_OUTPUT_DIR = SCRIPT_DIR.parent / "output"
DEFAULT_BASELINE = DEFAULT_DATA_DIR / "baseline.md"
DEFAULT_CONTRACT = DEFAULT_DATA_DIR / "contract.docx"

LOGGER = logging.getLogger(__name__)

# Bilingual labels
LABELS = {
    "zh": {
        "contract": "合同",
        "baseline": "基线",
        "no_type": "无类型",
        "verdict_conflict": "冲突",
        "verdict_violation": "明显违背",
        "verdict_inconsistent": "可能不一致",
        "verdict_ok": "无冲突/等价",
        "verdict_na": "不适用/无对应",
        "severity_high": "高",
        "severity_medium": "中",
        "severity_low": "低",
        "confidence_high": "高",
        "confidence_medium": "中",
        "confidence_low": "低",
        "needs_review": "需人工复核",
        "report_title": "合同审校结果汇总",
        "summary_title": "合同风险审核报告",
        "total_sections": "总条款数",
        "high_risk": "高风险",
        "medium_risk": "中风险",
        "low_risk": "低风险",
        "no_risk": "无风险",
        "review_needed": "需复核",
        "comparison_title": "合同条款对比分析",
    },
    "en": {
        "contract": "Contract",
        "baseline": "Baseline",
        "no_type": "N/A",
        "verdict_conflict": "Conflict",
        "verdict_violation": "Violation",
        "verdict_inconsistent": "Possibly Inconsistent",
        "verdict_ok": "No Conflict",
        "verdict_na": "Not Applicable",
        "severity_high": "High",
        "severity_medium": "Medium",
        "severity_low": "Low",
        "confidence_high": "High",
        "confidence_medium": "Medium",
        "confidence_low": "Low",
        "needs_review": "Needs Review",
        "report_title": "Contract Review Summary",
        "summary_title": "Contract Risk Review Report",
        "total_sections": "Total Sections",
        "high_risk": "High Risk",
        "medium_risk": "Medium Risk",
        "low_risk": "Low Risk",
        "no_risk": "No Risk",
        "review_needed": "Review Needed",
        "comparison_title": "Contract Clause Comparison Analysis",
    },
}

# Risk colors for highlighting
RISK_COLORS = {
    "高": "red",
    "中": "yellow",
    "低": "cyan",
    "High": "red",
    "Medium": "yellow",
    "Low": "cyan",
}

# ============================================================================
# Data Models
# ============================================================================


@dataclass(slots=True)
class Section:
    """A parsed section from the contract document."""
    section_id: str
    heading_path: List[str]
    page: Optional[int]
    text: str
    paragraph_indexes: List[int] = field(default_factory=list)


@dataclass(slots=True)
class Judgment:
    """Result returned by the LLM for a given section."""
    section_id: str
    heading_path: List[str]
    page: Optional[int]
    verdict: str
    type: Optional[str]
    severity: Optional[str]
    rationale: str
    evidence_contract: str
    evidence_baseline: str
    confidence: str = "中"  # 高/中/低
    needs_review: bool = False

    def is_actionable(self) -> bool:
        """Whether this judgment requires annotating the document."""
        return self.verdict in {"冲突", "明显违背", "可能不一致", "Conflict", "Violation", "Possibly Inconsistent"}

    def to_dict(self) -> dict:
        """Convert the judgment to a serializable dictionary."""
        return {
            "section_id": self.section_id,
            "heading_path": self.heading_path,
            "page": self.page,
            "verdict": self.verdict,
            "type": self.type,
            "severity": self.severity,
            "rationale": self.rationale,
            "evidence_contract": self.evidence_contract,
            "evidence_baseline": self.evidence_baseline,
            "confidence": self.confidence,
            "needs_review": self.needs_review,
        }


# ============================================================================
# Document Parser
# ============================================================================

HEADING_PREFIX = "Heading"

# Section splitting configuration
MAX_SECTION_CHARS = 8000  # ~4000 tokens, suited for 64K+ context models
MIN_SECTION_CHARS = 200   # Don't create tiny fragments


def _paragraph_heading_level(paragraph) -> int | None:
    try:
        style_name = paragraph.style.name or ""
    except AttributeError:
        style_name = ""
    if not style_name.startswith(HEADING_PREFIX):
        return None
    try:
        level = int(style_name.split()[1])
    except (IndexError, ValueError):
        return None
    return level


def _split_long_section(section: Section, max_chars: int = MAX_SECTION_CHARS) -> List[Section]:
    """Split a long section into smaller chunks at sentence boundaries."""
    if len(section.text) <= max_chars:
        return [section]

    # Split by sentence endings (Chinese and English)
    import re
    sentences = re.split(r'(?<=[。.!?！？;；])\s*', section.text)

    chunks: List[Section] = []
    current_chunk: List[str] = []
    current_length = 0
    chunk_index = 1

    for sentence in sentences:
        sentence = sentence.strip()
        if not sentence:
            continue

        # If adding this sentence would exceed limit, flush current chunk
        if current_length + len(sentence) > max_chars and current_chunk:
            chunk_text = " ".join(current_chunk)
            if len(chunk_text) >= MIN_SECTION_CHARS:
                chunks.append(Section(
                    section_id=f"{section.section_id}-{chunk_index}",
                    heading_path=section.heading_path.copy(),
                    page=section.page,
                    text=chunk_text,
                    paragraph_indexes=section.paragraph_indexes.copy(),
                ))
                chunk_index += 1
            current_chunk = []
            current_length = 0

        current_chunk.append(sentence)
        current_length += len(sentence) + 1  # +1 for space

    # Flush remaining
    if current_chunk:
        chunk_text = " ".join(current_chunk)
        if len(chunk_text) >= MIN_SECTION_CHARS or not chunks:
            chunks.append(Section(
                section_id=f"{section.section_id}-{chunk_index}" if chunk_index > 1 else section.section_id,
                heading_path=section.heading_path.copy(),
                page=section.page,
                text=chunk_text,
                paragraph_indexes=section.paragraph_indexes.copy(),
            ))
        elif chunks:
            # Merge tiny remainder with last chunk
            last = chunks[-1]
            chunks[-1] = Section(
                section_id=last.section_id,
                heading_path=last.heading_path,
                page=last.page,
                text=last.text + " " + chunk_text,
                paragraph_indexes=last.paragraph_indexes,
            )

    return chunks if chunks else [section]


def parse_contract(path: str, max_section_chars: int = MAX_SECTION_CHARS) -> List[Section]:
    """Parse a contract (docx or pdf) into a list of sections with heading context."""
    ext = Path(path).suffix.lower()
    if ext == ".pdf":
        return parse_contract_pdf(path, max_section_chars)
    document = Document(path)
    sections: List[Section] = []
    heading_stack: list[str] = []
    buffer: list[str] = []
    buffer_paragraph_indexes: list[int] = []
    section_counter = 1

    def flush() -> None:
        nonlocal section_counter, buffer, buffer_paragraph_indexes
        if not buffer:
            return
        section = Section(
            section_id=f"S{section_counter:03d}",
            heading_path=heading_stack.copy(),
            page=None,
            text="\n".join(buffer).strip(),
            paragraph_indexes=buffer_paragraph_indexes.copy(),
        )
        sections.append(section)
        section_counter += 1
        buffer.clear()
        buffer_paragraph_indexes.clear()

    for index, paragraph in enumerate(document.paragraphs):
        level = _paragraph_heading_level(paragraph)
        if level is not None:
            flush()
            while len(heading_stack) >= level:
                heading_stack.pop()
            heading_stack.append(paragraph.text.strip())
            continue
        text = paragraph.text.strip()
        if not text:
            continue
        buffer.append(text)
        buffer_paragraph_indexes.append(index)

    flush()

    # Split long sections into smaller chunks
    if max_section_chars > 0:
        split_sections: List[Section] = []
        for section in sections:
            split_sections.extend(_split_long_section(section, max_section_chars))
        LOGGER.info("Section splitting: %d -> %d sections (max %d chars)",
                    len(sections), len(split_sections), max_section_chars)
        return split_sections

    return sections


def _is_heading_font(span: dict, median_size: float) -> bool:
    """Heuristic: a span is a heading if font is significantly larger or bold."""
    size = span.get("size", 0)
    flags = span.get("flags", 0)  # bit 4 = bold in PyMuPDF
    is_bold = bool(flags & (1 << 4))
    is_large = size >= median_size * 1.2 and size > median_size
    return is_large or (is_bold and size >= median_size)


def _guess_heading_level(size: float, median_size: float) -> int:
    """Map font size ratio to a heading level (1-3)."""
    ratio = size / median_size if median_size > 0 else 1.0
    if ratio >= 1.6:
        return 1
    if ratio >= 1.3:
        return 2
    return 3


def parse_contract_pdf(path: str, max_section_chars: int = MAX_SECTION_CHARS) -> List[Section]:
    """Parse a contract PDF into a list of sections with heading context."""
    doc = fitz.open(path)

    # First pass: collect all font sizes to compute median body size
    all_sizes: list[float] = []
    for page in doc:
        blocks = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)["blocks"]
        for block in blocks:
            if block["type"] != 0:  # text blocks only
                continue
            for line in block["lines"]:
                for span in line["spans"]:
                    text = span["text"].strip()
                    if text:
                        all_sizes.append(span["size"])

    if not all_sizes:
        doc.close()
        return []

    median_size = sorted(all_sizes)[len(all_sizes) // 2]

    # Second pass: extract sections
    sections: List[Section] = []
    heading_stack: list[str] = []
    buffer: list[str] = []
    buffer_indexes: list[int] = []  # encoded as page * 10000 + block_idx
    section_counter = 1

    def flush() -> None:
        nonlocal section_counter, buffer, buffer_indexes
        if not buffer:
            return
        text = "\n".join(buffer).strip()
        if not text:
            buffer.clear()
            buffer_indexes.clear()
            return
        sections.append(Section(
            section_id=f"S{section_counter:03d}",
            heading_path=heading_stack.copy(),
            page=buffer_indexes[0] // 10000 + 1 if buffer_indexes else None,
            text=text,
            paragraph_indexes=buffer_indexes.copy(),
        ))
        section_counter += 1
        buffer.clear()
        buffer_indexes.clear()

    for page_num, page in enumerate(doc):
        blocks = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)["blocks"]
        for block_idx, block in enumerate(blocks):
            if block["type"] != 0:
                continue
            # Combine all spans in the block
            block_text_parts: list[str] = []
            block_is_heading = False
            heading_size = 0.0
            for line in block["lines"]:
                for span in line["spans"]:
                    text = span["text"].strip()
                    if text:
                        block_text_parts.append(text)
                        if _is_heading_font(span, median_size):
                            block_is_heading = True
                            heading_size = max(heading_size, span["size"])

            block_text = " ".join(block_text_parts).strip()
            if not block_text:
                continue

            encoded_idx = page_num * 10000 + block_idx

            if block_is_heading and len(block_text) < 200:
                flush()
                level = _guess_heading_level(heading_size, median_size)
                while len(heading_stack) >= level:
                    heading_stack.pop()
                heading_stack.append(block_text)
            else:
                buffer.append(block_text)
                buffer_indexes.append(encoded_idx)

    flush()
    doc.close()

    # Split long sections
    if max_section_chars > 0:
        split_sections: List[Section] = []
        for section in sections:
            split_sections.extend(_split_long_section(section, max_section_chars))
        LOGGER.info("PDF section splitting: %d -> %d sections (max %d chars)",
                    len(sections), len(split_sections), max_section_chars)
        return split_sections

    return sections


# ============================================================================
# LLM Client
# ============================================================================

ALLOWED_VERDICTS = {"冲突", "明显违背", "可能不一致", "无冲突/等价", "不适用/无对应"}
ALLOWED_TYPES = {"直接矛盾", "条件削弱", "单方权利扩张", "缺失关键限制", "术语定义冲突", "措辞模糊", "责任/范围无上限", "违约责任失衡", None, "", "null"}

PROMPT_TEMPLATE = '''
系统指令：
你是资深法律顾问，负责审核合同条款是否与公司 baseline（审核基线）存在风险。
Baseline 分为三个风险等级：
- 🔴 红线条款（severity: 高）：绝不允许，触及公司核心利益底线，必须拒绝
- 🟠 橘红条款（severity: 中）：会导致吃亏但战略上可商量，需高层审批
- 🔵 蓝色条款（severity: 低）：措辞有歧义，建议修改以消除风险

请输出严格 JSON：verdict, type, severity, rationale, evidence_contract, evidence_baseline, confidence。
必须引用合同与 baseline 的原文片段作为证据。

标签说明：
- verdict:
  - 冲突：与红线条款直接冲突，必须拒绝
  - 明显违背：严重偏离 baseline 要求
  - 可能不一致：存在潜在风险，需进一步确认
  - 无冲突/等价：符合 baseline 或在可接受范围内
  - 不适用/无对应：baseline 中无相关条款
- type（必须严格使用以下标签，禁止自行创造）:
  - 直接矛盾：合同条款与 baseline 要求方向完全相反
  - 条件削弱：合同放宽了 baseline 要求的标准或条件
  - 单方权利扩张：一方获得 baseline 未授予的单方面权利
  - 缺失关键限制：baseline 要求的限制条款在合同中完全缺失
  - 责任/范围无上限：合同条款设定了无上限/无边界的义务或赔偿责任
  - 违约责任失衡：双方违约后果不对等，惩罚比例严重失衡
  - 术语定义冲突：合同与 baseline 对同一术语定义不同
  - 措辞模糊：使用主观、含糊表述，可能引发争议
  - null：verdict 为"无冲突/等价"或"不适用/无对应"时使用
- severity:
  - 高：触犯红线，绝不接受
  - 中：橘红区域，战略可商量
  - 低：蓝色区域，措辞风险
- confidence: 高 | 中 | 低（你对此判断的把握程度，低置信度表示需要人工复核）

合同条款（含标题链）：
标题链：{heading}
正文：
"""{text}"""

baseline（全文）：
"""{baseline}"""

仅输出 JSON：
'''


# ============================================================================
# LLM Cache
# ============================================================================

class LLMCache:
    """Persistent cache for LLM responses to avoid redundant API calls.

    Batches disk writes to reduce I/O overhead on large runs. Writes every
    *flush_interval* new entries instead of on every set().
    """

    def __init__(self, cache_dir: Path | str | None = None, flush_interval: int = 20):
        if cache_dir is None:
            cache_dir = DEFAULT_OUTPUT_DIR / ".cache"
        self.cache_dir = Path(cache_dir)
        self.cache_dir.mkdir(parents=True, exist_ok=True)
        self.cache_file = self.cache_dir / "llm_cache.json"
        self._cache: dict[str, dict] = self._load_cache()
        self._dirty = 0
        self._flush_interval = flush_interval
        self._lock = __import__("threading").Lock()

    def _load_cache(self) -> dict[str, dict]:
        if self.cache_file.exists():
            try:
                return json.loads(self.cache_file.read_text(encoding="utf-8"))
            except (json.JSONDecodeError, IOError):
                return {}
        return {}

    def _save_cache(self) -> None:
        self.cache_file.write_text(
            json.dumps(self._cache, ensure_ascii=False, indent=2),
            encoding="utf-8"
        )

    def _make_key(self, section_text: str, baseline_text: str, model: str) -> str:
        """Create a unique cache key based on content hash."""
        content = f"{model}::{section_text}::{baseline_text}"
        return hashlib.sha256(content.encode("utf-8")).hexdigest()[:16]

    def get(self, section_text: str, baseline_text: str, model: str) -> dict | None:
        """Get cached result if exists."""
        key = self._make_key(section_text, baseline_text, model)
        return self._cache.get(key)

    def set(self, section_text: str, baseline_text: str, model: str, data: dict) -> None:
        """Store result in cache. Flushes to disk every flush_interval entries."""
        key = self._make_key(section_text, baseline_text, model)
        with self._lock:
            self._cache[key] = data
            self._dirty += 1
            if self._dirty >= self._flush_interval:
                self._save_cache()
                self._dirty = 0

    def flush(self) -> None:
        """Force write any pending cache entries to disk."""
        with self._lock:
            if self._dirty > 0:
                self._save_cache()
                self._dirty = 0

    def stats(self) -> dict:
        """Return cache statistics."""
        return {
            "entries": len(self._cache),
            "cache_file": str(self.cache_file),
        }


class LLMClient:
    """Wrapper around the OpenAI SDK with retry and validation logic."""

    def __init__(
        self,
        model: str | None = None,
        base_url: str | None = None,
        api_key: str | None = None,
        max_retries: int = 2,
        cache: LLMCache | None = None,
    ) -> None:
        self.model = model or os.environ.get("OPENAI_MODEL")
        if not self.model:
            raise ValueError("OPENAI_MODEL is required")
        base_url = base_url or os.environ.get("OPENAI_BASE_URL")
        api_key = api_key or os.environ.get("OPENAI_API_KEY")
        if not api_key:
            raise ValueError("OPENAI_API_KEY is required")
        self.client = OpenAI(api_key=api_key, base_url=base_url)
        self.max_retries = max_retries
        self.cache = cache
        self._cache_hits = 0
        self._cache_misses = 0

    def _build_prompt(self, section: Section, baseline_text: str) -> str:
        heading = " > ".join(section.heading_path) if section.heading_path else "(未命名条款)"
        escaped_text = section.text.replace("{", "{{").replace("}", "}}")
        escaped_baseline = baseline_text.replace("{", "{{").replace("}", "}}")
        return PROMPT_TEMPLATE.format(
            heading=heading,
            text=escaped_text,
            baseline=escaped_baseline,
        )

    def evaluate_section(self, section: Section, baseline_text: str) -> Judgment:
        # Check cache first
        if self.cache:
            cached = self.cache.get(section.text, baseline_text, self.model)
            if cached:
                self._cache_hits += 1
                LOGGER.debug("Cache hit for %s", section.section_id)
                return self._build_judgment(section, cached)
            self._cache_misses += 1

        prompt = self._build_prompt(section, baseline_text)
        last_error: Exception | None = None
        for attempt in range(1, self.max_retries + 1):
            try:
                response = self.client.chat.completions.create(
                    model=self.model,
                    messages=[{"role": "user", "content": prompt}],
                    temperature=0.0,
                )
                text = response.choices[0].message.content.strip()
                data = self._safe_load_json(text)

                # Store in cache
                if self.cache:
                    self.cache.set(section.text, baseline_text, self.model, data)

                return self._build_judgment(section, data)
            except Exception as exc:
                last_error = exc
                LOGGER.warning(
                    "LLM evaluation failed on attempt %s for %s: %s", attempt, section.section_id, exc
                )
        LOGGER.error("LLM evaluation ultimately failed for %s: %s", section.section_id, last_error)
        raise RuntimeError(f"LLM服务不可用，无法完成合同审核。请检查API配置和网络连接。") from last_error

    def cache_stats(self) -> dict:
        """Return cache hit/miss statistics."""
        total = self._cache_hits + self._cache_misses
        return {
            "hits": self._cache_hits,
            "misses": self._cache_misses,
            "hit_rate": f"{self._cache_hits / total * 100:.1f}%" if total > 0 else "N/A",
        }

    def _safe_load_json(self, text: str) -> dict[str, Any]:
        candidate = text.strip()
        if candidate.startswith("```"):
            candidate = candidate.strip("`")
            parts = candidate.split("\n", 1)
            candidate = parts[1] if len(parts) > 1 else ""
            if candidate.endswith("```"):
                candidate = candidate[:-3]
        try:
            return json.loads(candidate)
        except json.JSONDecodeError as exc:
            raise ValueError(f"模型输出非 JSON: {candidate!r}") from exc

    def _build_judgment(self, section: Section, data: dict[str, Any]) -> Judgment:
        verdict = self._normalize_verdict(data.get("verdict"))
        conflict_type = self._normalize_type(data.get("type"))
        severity = self._normalize_severity(data.get("severity"))
        confidence = self._normalize_confidence(data.get("confidence"))
        rationale = (data.get("rationale") or "").strip()
        evidence_contract = (data.get("evidence_contract") or "").strip()
        evidence_baseline = (data.get("evidence_baseline") or "").strip()
        needs_review = confidence == "低"
        return Judgment(
            section_id=section.section_id,
            heading_path=section.heading_path,
            page=section.page,
            verdict=verdict,
            type=conflict_type,
            severity=severity,
            rationale=rationale,
            evidence_contract=evidence_contract,
            evidence_baseline=evidence_baseline,
            confidence=confidence,
            needs_review=needs_review,
        )

    def _normalize_verdict(self, value: Any) -> str:
        text = (value or "").strip()
        if text not in ALLOWED_VERDICTS:
            LOGGER.warning("Unexpected verdict %s, defaulting to '可能不一致'", text)
            return "可能不一致"
        return text

    # Mapping for common LLM-invented type labels → canonical types
    _TYPE_ALIASES: dict[str, str] = {
        "范围无边界": "责任/范围无上限",
        "服务范围无边界": "责任/范围无上限",
        "责任无上限": "责任/范围无上限",
        "无限责任": "责任/范围无上限",
        "无上限": "责任/范围无上限",
        "责任失衡": "违约责任失衡",
        "权利义务失衡": "违约责任失衡",
        "惩罚失衡": "违约责任失衡",
        "定义冲突": "术语定义冲突",
        "术语冲突": "术语定义冲突",
        "权利扩张": "单方权利扩张",
        "单方扩张": "单方权利扩张",
        "缺失限制": "缺失关键限制",
        "关键限制缺失": "缺失关键限制",
    }

    def _normalize_type(self, value: Any) -> str | None:
        if value is None:
            return None
        text = str(value).strip()
        if text in {"", "null", "None"}:
            return None
        if text in ALLOWED_TYPES:
            return text
        # Try alias mapping
        mapped = self._TYPE_ALIASES.get(text)
        if mapped:
            LOGGER.debug("Mapped type '%s' -> '%s'", text, mapped)
            return mapped
        LOGGER.warning("Unexpected type %s, defaulting to None", text)
        return None

    def _normalize_severity(self, value: Any) -> str | None:
        if value is None:
            return None
        text = str(value).strip()
        if text in {"高", "高风险", "HIGH", "High"}:
            return "高"
        if text in {"中", "中风险", "MEDIUM", "Medium"}:
            return "中"
        if text in {"低", "低风险", "LOW", "Low"}:
            return "低"
        LOGGER.warning("Unexpected severity %s, defaulting to '中'", text)
        return "中"

    def _normalize_confidence(self, value: Any) -> str:
        if value is None:
            return "中"
        text = str(value).strip()
        if text in {"高", "HIGH", "High"}:
            return "高"
        if text in {"中", "MEDIUM", "Medium"}:
            return "中"
        if text in {"低", "LOW", "Low"}:
            return "低"
        return "中"


# ============================================================================
# Document Annotation (with color highlighting)
# ============================================================================

HIGHLIGHT_COLORS = {
    "高": "red",
    "中": "yellow",
    "低": "cyan",
}


def _add_highlight_to_paragraph(paragraph, color: str) -> None:
    """Add highlight color to all runs in a paragraph."""
    for run in paragraph.runs:
        rpr = run._r.get_or_add_rPr()
        highlight = OxmlElement("w:highlight")
        highlight.set(qn("w:val"), color)
        rpr.append(highlight)


def _add_comment(paragraph, comments_part, text: str, author: str = "DocReview", initials: str = "DR") -> None:
    comment_id = comments_part.next_id
    comment = OxmlElement("w:comment")
    comment.set(qn("w:id"), str(comment_id))
    comment.set(qn("w:author"), author)
    comment.set(qn("w:initials"), initials)
    comment.set(qn("w:date"), datetime.now(timezone.utc).replace(microsecond=0).isoformat().replace("+00:00", "Z"))

    body_p = OxmlElement("w:p")
    body_r = OxmlElement("w:r")
    body_t = OxmlElement("w:t")
    body_t.text = text
    body_r.append(body_t)
    body_p.append(body_r)
    comment.append(body_p)
    comments_part.element.append(comment)

    start = OxmlElement("w:commentRangeStart")
    start.set(qn("w:id"), str(comment_id))
    end = OxmlElement("w:commentRangeEnd")
    end.set(qn("w:id"), str(comment_id))

    paragraph._p.insert(0, start)
    paragraph._p.append(end)

    reference_run = OxmlElement("w:r")
    reference = OxmlElement("w:commentReference")
    reference.set(qn("w:id"), str(comment_id))
    reference_run.append(reference)
    paragraph._p.append(reference_run)


def _clean_excerpt(text: str, limit: int = 200) -> str:
    """Collapse whitespace and truncate excerpts to keep comments concise."""
    stripped = (text or "").strip()
    if not stripped:
        return ""
    collapsed = " ".join(stripped.split())
    if len(collapsed) <= limit:
        return collapsed
    return collapsed[: limit - 1] + "…"


def build_comment_text(judgment: Judgment, lang: str = "zh") -> str:
    L = LABELS[lang]
    rationale = _clean_excerpt(judgment.rationale)
    contract_excerpt = _clean_excerpt(judgment.evidence_contract)
    baseline_excerpt = _clean_excerpt(judgment.evidence_baseline)

    parts = [f"[{judgment.verdict}/{judgment.type or L['no_type']}/{judgment.severity or L['severity_medium']}]"]
    if judgment.needs_review:
        parts[0] += f" ⚠️ {L['needs_review']}"
    if rationale:
        parts.append(rationale)
    if contract_excerpt:
        parts.append(f"{L['contract']}: {contract_excerpt}")
    if baseline_excerpt:
        parts.append(f"{L['baseline']}: {baseline_excerpt}")
    return "\n".join(parts)


def _find_best_matching_paragraph(document, section: Section, evidence_text: str):
    """Find the paragraph that best matches the evidence text."""
    if not evidence_text or not section.paragraph_indexes:
        return document.paragraphs[section.paragraph_indexes[0]] if section.paragraph_indexes else None

    best_ratio = 0
    best_para = None
    for idx in section.paragraph_indexes:
        para = document.paragraphs[idx]
        ratio = SequenceMatcher(None, para.text, evidence_text).ratio()
        if ratio > best_ratio:
            best_ratio = ratio
            best_para = para

    # Fallback to first paragraph if no good match
    if best_ratio < 0.3 and section.paragraph_indexes:
        return document.paragraphs[section.paragraph_indexes[0]]
    return best_para


def annotate_document(
    source_path: str,
    destination_path: str,
    judgments: Iterable[Judgment],
    sections: Mapping[str, Section],
    lang: str = "zh",
) -> None:
    document = Document(source_path)
    comments_part = document.part._comments_part

    for judgment in judgments:
        if not judgment.is_actionable():
            continue
        section = sections.get(judgment.section_id)
        if not section or not section.paragraph_indexes:
            continue

        # Find best matching paragraph based on evidence
        paragraph = _find_best_matching_paragraph(document, section, judgment.evidence_contract)
        if not paragraph:
            continue

        # Add color highlight based on severity
        if judgment.severity in HIGHLIGHT_COLORS:
            _add_highlight_to_paragraph(paragraph, HIGHLIGHT_COLORS[judgment.severity])

        # Add comment
        comment_text = build_comment_text(judgment, lang)
        author = f"AI-{judgment.severity or '中'}"
        _add_comment(paragraph, comments_part, comment_text, author=author)

    document.save(destination_path)


# PDF annotation colors (RGB tuples for PyMuPDF, range 0-1)
PDF_HIGHLIGHT_COLORS = {
    "高": (1.0, 0.6, 0.6),       # light red
    "中": (1.0, 0.9, 0.4),       # light yellow/orange
    "低": (0.6, 0.85, 1.0),      # light blue
    "High": (1.0, 0.6, 0.6),
    "Medium": (1.0, 0.9, 0.4),
    "Low": (0.6, 0.85, 1.0),
}


def _find_text_on_page(page, search_text: str, max_instances: int = 1):
    """Search for text on a PDF page, return list of Rect quads."""
    if not search_text:
        return []
    # Try progressively shorter prefixes if full text not found
    for length in (len(search_text), 100, 50):
        snippet = search_text[:length].strip()
        if not snippet:
            continue
        quads = page.search_for(snippet, quads=True)
        if quads:
            return quads[:max_instances]
    return []


def annotate_pdf(
    source_path: str,
    destination_path: str,
    judgments: Iterable[Judgment],
    sections: Mapping[str, Section],
    lang: str = "zh",
) -> None:
    """Annotate a PDF with highlights and sticky-note comments."""
    doc = fitz.open(source_path)

    for judgment in judgments:
        if not judgment.is_actionable():
            continue
        section = sections.get(judgment.section_id)
        if not section:
            continue

        color = PDF_HIGHLIGHT_COLORS.get(judgment.severity, (1.0, 0.9, 0.4))
        comment_text = build_comment_text(judgment, lang)

        # Determine which page(s) to search
        target_pages: list[int] = []
        if section.paragraph_indexes:
            target_pages = sorted(set(idx // 10000 for idx in section.paragraph_indexes))
        elif section.page is not None:
            target_pages = [section.page - 1]  # 0-indexed

        evidence = judgment.evidence_contract or section.text[:100]
        annotated = False

        for page_num in target_pages:
            if page_num < 0 or page_num >= len(doc):
                continue
            page = doc[page_num]
            quads = _find_text_on_page(page, evidence)
            if quads:
                # Add highlight annotation
                highlight = page.add_highlight_annot(quads)
                highlight.set_colors(stroke=color)
                highlight.set_info(title=f"AI-{judgment.severity or '中'}", content=comment_text)
                highlight.update()
                annotated = True
                break

        # Fallback: add a sticky note at top of the first target page
        if not annotated and target_pages:
            page_num = target_pages[0]
            if 0 <= page_num < len(doc):
                page = doc[page_num]
                note = page.add_text_annot((72, 72), comment_text)
                note.set_info(title=f"AI-{judgment.severity or '中'}")
                note.update()

    doc.save(destination_path)
    doc.close()


# ============================================================================
# Report Generation
# ============================================================================

def export_json(path: str, judgments: Iterable[Judgment]) -> None:
    data = [judgment.to_dict() for judgment in judgments]
    Path(path).write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")


def _format_heading(heading_path: list[str]) -> str:
    return " > ".join(heading_path) if heading_path else "(未命名条款)"



# ============================================================================
# HTML Report (Executive Summary + Side-by-Side Comparison)
# ============================================================================

def _render_markdown_to_html(md_text: str) -> str:
    """Convert markdown text to HTML with proper styling."""
    import html as html_module

    lines = md_text.split('\n')
    html_output = []
    in_list = False
    in_table = False
    table_rows = []
    current_risk_class = ""

    def close_list():
        nonlocal in_list
        if in_list:
            html_output.append('</ul>')
            in_list = False

    def close_table():
        nonlocal in_table, table_rows
        if in_table and table_rows:
            html_output.append('<table class="baseline-table">')
            html_output.append('<thead><tr>')
            headers = table_rows[0]
            for h in headers:
                html_output.append(f'<th>{h}</th>')
            html_output.append('</tr></thead>')
            html_output.append('<tbody>')
            for row in table_rows[2:]:  # Skip header and separator
                html_output.append('<tr>')
                for cell in row:
                    cell_class = ""
                    if "🔴" in cell or "红线" in cell:
                        cell_class = ' class="cell-red"'
                    elif "🟠" in cell or "橘红" in cell:
                        cell_class = ' class="cell-orange"'
                    elif "🔵" in cell or "蓝色" in cell:
                        cell_class = ' class="cell-blue"'
                    html_output.append(f'<td{cell_class}>{cell}</td>')
                html_output.append('</tr>')
            html_output.append('</tbody></table>')
            table_rows = []
            in_table = False

    def process_inline(text: str) -> str:
        """Process inline markdown: bold, italic, code."""
        # Bold **text**
        text = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', text)
        # Italic *text*
        text = re.sub(r'\*(.+?)\*', r'<em>\1</em>', text)
        # Code `text`
        text = re.sub(r'`(.+?)`', r'<code>\1</code>', text)
        return text

    for line in lines:
        stripped = line.strip()

        # Skip empty lines
        if not stripped:
            close_list()
            close_table()
            continue

        # Horizontal rule
        if stripped == '---':
            close_list()
            close_table()
            html_output.append('<hr class="baseline-hr">')
            continue

        # Table row
        if stripped.startswith('|') and stripped.endswith('|'):
            close_list()
            if not in_table:
                in_table = True
            cells = [c.strip() for c in stripped[1:-1].split('|')]
            # Check if it's separator row
            if all(c.replace('-', '').replace(':', '') == '' for c in cells):
                table_rows.append(cells)  # Keep separator to mark header end
            else:
                table_rows.append(cells)
            continue
        else:
            close_table()

        # Headers
        if stripped.startswith('# '):
            close_list()
            text = process_inline(stripped[2:])
            html_output.append(f'<h1 class="baseline-h1">{text}</h1>')
            continue
        if stripped.startswith('## '):
            close_list()
            text = process_inline(stripped[3:])
            risk_class = ""
            if "🔴" in text or "红线" in text:
                risk_class = " red-section"
                current_risk_class = "red"
            elif "🟠" in text or "橘红" in text:
                risk_class = " orange-section"
                current_risk_class = "orange"
            elif "🔵" in text or "蓝色" in text:
                risk_class = " blue-section"
                current_risk_class = "blue"
            else:
                current_risk_class = ""
            html_output.append(f'<h2 class="baseline-h2{risk_class}">{text}</h2>')
            continue
        if stripped.startswith('### '):
            close_list()
            text = process_inline(stripped[4:])
            html_output.append(f'<h3 class="baseline-h3">{text}</h3>')
            continue

        # Blockquote
        if stripped.startswith('> '):
            close_list()
            text = process_inline(stripped[2:])
            html_output.append(f'<blockquote class="baseline-quote">{text}</blockquote>')
            continue

        # List items
        if stripped.startswith('- '):
            if not in_list:
                html_output.append(f'<ul class="baseline-list {current_risk_class}">')
                in_list = True
            text = process_inline(stripped[2:])
            # Handle nested content with indentation
            item_class = ""
            if stripped.startswith('- **'):
                item_class = ' class="list-bold"'
            html_output.append(f'<li{item_class}>{text}</li>')
            continue

        # Indented list continuation or example
        if line.startswith('  ') and (stripped.startswith('- ') or stripped.startswith('❌') or stripped.startswith('✅')):
            text = process_inline(stripped)
            if stripped.startswith('❌'):
                html_output.append(f'<div class="example bad">{text}</div>')
            elif stripped.startswith('✅'):
                html_output.append(f'<div class="example good">{text}</div>')
            else:
                html_output.append(f'<div class="sub-item">{text[2:]}</div>')
            continue

        # Regular paragraph
        close_list()
        text = process_inline(stripped)
        html_output.append(f'<p class="baseline-p">{text}</p>')

    close_list()
    close_table()
    return '\n'.join(html_output)


def generate_comparison_html(
    path: str,
    judgments: List[Judgment],
    sections: Mapping[str, Section],
    baseline_text: str,
    lang: str = "zh",
) -> None:
    """Generate an HTML report with executive summary and side-by-side comparison.

    Combines risk statistics, distribution chart, high-risk details table,
    search, print support, and contract vs baseline comparison.
    """
    import html as _html

    L = LABELS[lang]

    # Render baseline markdown to HTML
    baseline_html = _render_markdown_to_html(baseline_text)

    # Statistics
    total = len(judgments)
    high = sum(1 for j in judgments if j.severity == "高")
    medium = sum(1 for j in judgments if j.severity == "中" and j.is_actionable())
    low = sum(1 for j in judgments if j.severity == "低" and j.is_actionable())
    ok = total - high - medium - low
    review = sum(1 for j in judgments if j.needs_review)
    high_pct = f"{high / total * 100:.1f}" if total else "0"
    med_pct = f"{medium / total * 100:.1f}" if total else "0"
    low_pct = f"{low / total * 100:.1f}" if total else "0"
    ok_pct = f"{ok / total * 100:.1f}" if total else "0"

    # Build high-risk summary table rows
    high_items = [j for j in judgments if j.severity == "高"]
    medium_items = [j for j in judgments if j.severity == "中" and j.is_actionable()]

    summary_rows = ""
    for idx, item in enumerate(high_items):
        heading = _html.escape(_format_heading(item.heading_path))
        review_mark = ' <span class="tag review" style="font-size:10px;padding:2px 6px;">⚠</span>' if item.needs_review else ""
        summary_rows += f'''<tr class="summary-row-high" onclick="scrollToClause('clause-{item.section_id}')" style="cursor:pointer">
            <td><span class="severity-dot high-dot"></span> {heading}{review_mark}</td>
            <td>{_html.escape(item.verdict or "")}</td>
            <td>{_html.escape(item.type or "")}</td>
            <td>{_html.escape(item.rationale[:120] + "..." if len(item.rationale) > 120 else item.rationale)}</td>
        </tr>'''
    for idx, item in enumerate(medium_items):
        heading = _html.escape(_format_heading(item.heading_path))
        review_mark = ' <span class="tag review" style="font-size:10px;padding:2px 6px;">⚠</span>' if item.needs_review else ""
        summary_rows += f'''<tr class="summary-row-medium" onclick="scrollToClause('clause-{item.section_id}')" style="cursor:pointer">
            <td><span class="severity-dot medium-dot"></span> {heading}{review_mark}</td>
            <td>{_html.escape(item.verdict or "")}</td>
            <td>{_html.escape(item.type or "")}</td>
            <td>{_html.escape(item.rationale[:120] + "..." if len(item.rationale) > 120 else item.rationale)}</td>
        </tr>'''

    actionable_count = len([j for j in judgments if j.is_actionable()])

    html_content = f'''<!DOCTYPE html>
<html lang="{lang}">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{L["summary_title"]}</title>
    <link rel="preconnect" href="https://fonts.googleapis.com">
    <link rel="preconnect" href="https://fonts.gstatic.com" crossorigin>
    <link href="https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&family=Noto+Sans+SC:wght@400;500;600;700&display=swap" rel="stylesheet">
    <style>
        :root {{
            --primary: #0f172a;
            --primary-light: #1e293b;
            --accent: #3b82f6;
            --accent-light: #60a5fa;
            --success: #10b981;
            --success-bg: #ecfdf5;
            --warning: #f59e0b;
            --warning-bg: #fffbeb;
            --danger: #ef4444;
            --danger-bg: #fef2f2;
            --info: #3b82f6;
            --info-bg: #eff6ff;
            --purple: #8b5cf6;
            --purple-bg: #f5f3ff;
            --bg: #f8fafc;
            --bg-card: #ffffff;
            --border: #e2e8f0;
            --border-light: #f1f5f9;
            --text: #0f172a;
            --text-secondary: #475569;
            --text-muted: #94a3b8;
            --radius: 12px;
            --radius-sm: 8px;
            --shadow: 0 1px 3px rgba(0,0,0,0.08), 0 1px 2px rgba(0,0,0,0.06);
            --shadow-md: 0 4px 6px -1px rgba(0,0,0,0.08), 0 2px 4px -1px rgba(0,0,0,0.04);
        }}

        * {{ box-sizing: border-box; margin: 0; padding: 0; }}

        body {{
            font-family: 'Inter', 'Noto Sans SC', -apple-system, BlinkMacSystemFont, sans-serif;
            background: var(--bg);
            color: var(--text);
            line-height: 1.6;
            font-size: 14px;
            -webkit-font-smoothing: antialiased;
        }}

        /* Header */
        .header {{
            background: var(--primary);
            color: white;
            padding: 24px 32px;
        }}
        .header-inner {{
            max-width: 1800px;
            margin: 0 auto;
            display: flex;
            justify-content: space-between;
            align-items: center;
        }}
        .header h1 {{ font-size: 20px; font-weight: 600; letter-spacing: -0.3px; }}
        .header-actions {{ display: flex; align-items: center; gap: 16px; }}
        .header-meta {{ font-size: 13px; color: var(--text-muted); }}
        .btn-print {{
            padding: 7px 16px;
            border: 1px solid rgba(255,255,255,0.3);
            border-radius: 6px;
            background: transparent;
            color: white;
            font-size: 13px;
            cursor: pointer;
            transition: all 0.15s;
        }}
        .btn-print:hover {{ background: rgba(255,255,255,0.1); border-color: rgba(255,255,255,0.5); }}

        .container {{ max-width: 1800px; margin: 0 auto; padding: 24px 32px; }}

        /* Stats Row */
        .stats-row {{ display: flex; gap: 12px; margin-bottom: 20px; flex-wrap: wrap; }}
        .stat-item {{
            background: var(--bg-card);
            border: 1px solid var(--border);
            border-radius: var(--radius-sm);
            padding: 16px 24px;
            display: flex;
            align-items: center;
            gap: 16px;
            flex: 1;
            min-width: 150px;
            transition: all 0.15s ease;
        }}
        .stat-item:hover {{ border-color: var(--accent); box-shadow: var(--shadow); }}
        .stat-icon {{
            width: 44px; height: 44px;
            border-radius: var(--radius-sm);
            display: flex; align-items: center; justify-content: center;
            font-size: 20px; flex-shrink: 0;
        }}
        .stat-item.high .stat-icon {{ background: var(--danger-bg); }}
        .stat-item.medium .stat-icon {{ background: var(--warning-bg); }}
        .stat-item.low .stat-icon {{ background: var(--info-bg); }}
        .stat-item.ok .stat-icon {{ background: var(--success-bg); }}
        .stat-item.review .stat-icon {{ background: var(--purple-bg); }}
        .stat-content {{ display: flex; flex-direction: column; }}
        .stat-number {{ font-size: 28px; font-weight: 700; line-height: 1.1; }}
        .stat-item.high .stat-number {{ color: var(--danger); }}
        .stat-item.medium .stat-number {{ color: var(--warning); }}
        .stat-item.low .stat-number {{ color: var(--info); }}
        .stat-item.ok .stat-number {{ color: var(--success); }}
        .stat-item.review .stat-number {{ color: var(--purple); }}
        .stat-label {{ font-size: 12px; color: var(--text-muted); font-weight: 500; text-transform: uppercase; letter-spacing: 0.5px; }}

        /* Executive Summary */
        .summary-panel {{
            background: var(--bg-card);
            border: 1px solid var(--border);
            border-radius: var(--radius);
            margin-bottom: 20px;
            overflow: hidden;
        }}
        .summary-header {{
            padding: 16px 20px;
            border-bottom: 1px solid var(--border);
            background: var(--border-light);
            display: flex; justify-content: space-between; align-items: center;
        }}
        .summary-header h2 {{ font-size: 14px; font-weight: 600; }}
        .summary-body {{ padding: 20px; }}

        /* Distribution Bar */
        .dist-bar-container {{ margin-bottom: 20px; }}
        .dist-bar {{
            display: flex; height: 28px; border-radius: 6px; overflow: hidden;
            background: var(--border-light); margin-bottom: 10px;
        }}
        .dist-segment {{ display: flex; align-items: center; justify-content: center; color: white; font-size: 11px; font-weight: 600; min-width: 2px; transition: width 0.5s ease; }}
        .dist-segment.high {{ background: var(--danger); }}
        .dist-segment.medium {{ background: var(--warning); }}
        .dist-segment.low {{ background: var(--info); }}
        .dist-segment.ok {{ background: var(--success); }}
        .dist-legend {{
            display: flex; gap: 20px; font-size: 12px; color: var(--text-secondary); flex-wrap: wrap;
        }}
        .dist-legend-item {{ display: flex; align-items: center; gap: 6px; }}
        .dist-legend-dot {{ width: 10px; height: 10px; border-radius: 3px; }}

        /* Summary Table */
        .summary-table {{
            width: 100%; border-collapse: collapse; font-size: 13px; margin-top: 16px;
        }}
        .summary-table th {{
            text-align: left; padding: 10px 12px; font-weight: 600; color: var(--text-secondary);
            border-bottom: 2px solid var(--border); font-size: 12px; text-transform: uppercase; letter-spacing: 0.3px;
        }}
        .summary-table td {{ padding: 10px 12px; border-bottom: 1px solid var(--border-light); }}
        .summary-table tr:hover td {{ background: var(--bg); }}
        .summary-row-high td:first-child {{ border-left: 3px solid var(--danger); }}
        .summary-row-medium td:first-child {{ border-left: 3px solid var(--warning); }}
        .severity-dot {{ display: inline-block; width: 8px; height: 8px; border-radius: 50%; margin-right: 6px; }}
        .high-dot {{ background: var(--danger); }}
        .medium-dot {{ background: var(--warning); }}

        /* Toolbar */
        .toolbar {{
            display: flex; gap: 8px; margin-bottom: 20px; flex-wrap: wrap; align-items: center;
        }}
        .filter-pill {{
            padding: 8px 16px;
            border: 1px solid var(--border);
            border-radius: 20px;
            background: var(--bg-card);
            color: var(--text-secondary);
            font-size: 13px; font-weight: 500;
            cursor: pointer; transition: all 0.15s ease;
            display: flex; align-items: center; gap: 6px;
        }}
        .filter-pill:hover {{ border-color: var(--accent); color: var(--accent); }}
        .filter-pill.active {{ background: var(--primary); border-color: var(--primary); color: white; }}
        .search-box {{
            margin-left: auto;
            padding: 8px 14px;
            border: 1px solid var(--border);
            border-radius: 20px;
            font-size: 13px;
            width: 240px;
            outline: none;
            transition: border-color 0.15s;
            font-family: inherit;
        }}
        .search-box:focus {{ border-color: var(--accent); }}

        /* Two Column Layout */
        .layout {{ display: grid; grid-template-columns: 1fr 1fr; gap: 24px; align-items: start; }}

        /* Panel */
        .panel {{
            background: var(--bg-card);
            border: 1px solid var(--border);
            border-radius: var(--radius);
            overflow: hidden;
        }}
        .panel-header {{
            padding: 16px 20px;
            border-bottom: 1px solid var(--border);
            display: flex; align-items: center; justify-content: space-between;
            background: var(--border-light);
        }}
        .panel-title {{ font-size: 14px; font-weight: 600; display: flex; align-items: center; gap: 8px; }}
        .panel-badge {{
            background: var(--accent); color: white;
            font-size: 11px; padding: 2px 8px; border-radius: 10px; font-weight: 600;
        }}
        .panel-body {{ padding: 16px; max-height: calc(100vh - 280px); overflow-y: auto; }}

        /* Clause Cards */
        .clause {{
            background: white;
            border: 1px solid var(--border);
            border-radius: var(--radius-sm);
            margin-bottom: 10px;
            transition: all 0.15s ease;
            scroll-margin-top: 80px;
        }}
        .clause:hover {{ border-color: var(--accent-light); box-shadow: var(--shadow); }}
        .clause.highlight {{ box-shadow: 0 0 0 2px var(--accent), var(--shadow-md); }}

        .clause-header {{
            padding: 14px 16px;
            display: flex; justify-content: space-between; align-items: center;
            gap: 12px; cursor: pointer; background: white;
        }}
        .clause.expanded .clause-header {{ border-bottom: 1px solid var(--border-light); }}
        .clause-title {{ font-weight: 500; font-size: 13px; color: var(--text); flex: 1; line-height: 1.4; }}
        .clause-tags {{ display: flex; gap: 6px; flex-shrink: 0; align-items: center; }}
        .tag {{ padding: 4px 10px; border-radius: 6px; font-size: 11px; font-weight: 600; }}
        .tag.high {{ background: var(--danger); color: white; }}
        .tag.medium {{ background: var(--warning); color: white; }}
        .tag.low {{ background: var(--info); color: white; }}
        .tag.ok {{ background: var(--success); color: white; }}
        .tag.review {{ background: var(--purple); color: white; }}
        .tag.conf {{ background: transparent; color: var(--text-muted); font-weight: 500; }}

        .clause-body {{ padding: 16px; background: var(--bg); display: none; }}
        .clause.expanded .clause-body {{ display: block; }}
        .clause-text {{
            font-size: 13px; color: var(--text-secondary); line-height: 1.7;
            margin-bottom: 14px; padding: 12px 14px; background: white; border-radius: var(--radius-sm);
        }}
        .clause-analysis {{
            background: white; border-radius: var(--radius-sm);
            padding: 14px; margin-bottom: 10px;
        }}
        .clause-analysis-label {{
            font-size: 11px; font-weight: 600; color: var(--text-muted);
            margin-bottom: 6px; text-transform: uppercase; letter-spacing: 0.3px;
        }}
        .clause-analysis-text {{ font-size: 13px; color: var(--text); line-height: 1.6; }}

        /* Evidence Columns */
        .evidence-grid {{ display: grid; grid-template-columns: 1fr 1fr; gap: 10px; }}
        .evidence-block {{
            font-size: 12px; color: var(--text-secondary); font-style: italic;
            padding: 10px 14px; background: white; border-radius: var(--radius-sm);
            border-left: 3px solid var(--border);
        }}
        .evidence-block.contract {{ border-left-color: var(--accent); }}
        .evidence-block.baseline {{ border-left-color: var(--purple); }}
        .evidence-block-label {{
            font-size: 10px; font-weight: 600; color: var(--text-muted);
            text-transform: uppercase; letter-spacing: 0.3px; margin-bottom: 4px; font-style: normal;
        }}

        /* Baseline Styles */
        .baseline-content {{ padding: 24px; }}
        .baseline-h1 {{ font-size: 18px; font-weight: 700; color: var(--text); margin-bottom: 16px; }}
        .baseline-h2 {{
            font-size: 15px; font-weight: 600; color: var(--text);
            margin: 28px 0 14px 0; padding: 14px 18px;
            border-radius: var(--radius-sm); display: flex; align-items: center; gap: 10px;
        }}
        .baseline-h2.red-section {{ background: var(--danger-bg); color: var(--danger); }}
        .baseline-h2.orange-section {{ background: var(--warning-bg); color: #92400e; }}
        .baseline-h2.blue-section {{ background: var(--info-bg); color: var(--info); }}
        .baseline-h3 {{ font-size: 14px; font-weight: 600; color: var(--text); margin: 18px 0 10px 0; }}
        .baseline-quote {{
            background: var(--border-light); padding: 14px 18px; margin: 14px 0;
            font-size: 13px; color: var(--text-secondary); border-radius: var(--radius-sm); font-style: italic;
        }}
        .baseline-p {{ font-size: 13px; color: var(--text-secondary); margin: 10px 0; line-height: 1.7; }}
        .baseline-list {{ margin: 10px 0 10px 24px; font-size: 13px; color: var(--text-secondary); }}
        .baseline-list li {{ margin: 8px 0; line-height: 1.6; }}
        .baseline-list li strong {{ color: var(--text); }}
        .baseline-list.red li strong {{ color: var(--danger); }}
        .baseline-list.orange li strong {{ color: #92400e; }}
        .baseline-list.blue li strong {{ color: var(--info); }}
        .baseline-hr {{ border: none; height: 1px; background: var(--border); margin: 28px 0; }}
        .baseline-table {{ width: 100%; border-collapse: collapse; margin: 16px 0; font-size: 13px; }}
        .baseline-table th {{ background: var(--primary); color: white; padding: 10px 12px; text-align: left; font-weight: 600; }}
        .baseline-table td {{ padding: 10px 12px; border-bottom: 1px solid var(--border); }}
        .baseline-table tr:hover td {{ background: var(--bg); }}
        .baseline-table .cell-red {{ color: var(--danger); font-weight: 500; }}
        .baseline-table .cell-orange {{ color: var(--warning); font-weight: 500; }}
        .baseline-table .cell-blue {{ color: var(--info); font-weight: 500; }}
        .example {{ font-size: 12px; padding: 10px 14px; margin: 6px 0 6px 20px; border-radius: var(--radius-sm); font-family: 'SF Mono', Monaco, monospace; }}
        .example.bad {{ background: var(--danger-bg); color: #991b1b; }}
        .example.good {{ background: var(--success-bg); color: #166534; }}
        .sub-item {{ margin-left: 24px; font-size: 12px; color: var(--text-muted); }}

        .footer {{
            text-align: center; padding: 24px; color: var(--text-muted);
            font-size: 12px; border-top: 1px solid var(--border); margin-top: 24px;
        }}

        /* Responsive */
        @media (max-width: 1200px) {{
            .layout {{ grid-template-columns: 1fr; }}
            .panel-body {{ max-height: none; }}
            .evidence-grid {{ grid-template-columns: 1fr; }}
        }}
        @media (max-width: 768px) {{
            .container {{ padding: 16px; }}
            .stats-row {{ gap: 8px; }}
            .stat-item {{ padding: 12px 16px; min-width: 140px; }}
            .search-box {{ width: 100%; margin-left: 0; margin-top: 8px; }}
        }}

        /* Print */
        @media print {{
            body {{ background: white; font-size: 12px; }}
            .header {{ background: #333; padding: 16px; -webkit-print-color-adjust: exact; print-color-adjust: exact; }}
            .btn-print {{ display: none; }}
            .container {{ padding: 16px; }}
            .layout {{ display: block; }}
            .panel {{ break-inside: avoid; margin-bottom: 16px; }}
            .panel-body {{ max-height: none; overflow: visible; }}
            .clause-body {{ display: block !important; }}
            .clause {{ break-inside: avoid; }}
            .stats-row {{ flex-wrap: wrap; }}
            .stat-item {{ -webkit-print-color-adjust: exact; print-color-adjust: exact; }}
            .tag {{ -webkit-print-color-adjust: exact; print-color-adjust: exact; }}
            .dist-segment {{ -webkit-print-color-adjust: exact; print-color-adjust: exact; }}
            .search-box {{ display: none; }}
            .summary-table tr {{ break-inside: avoid; }}
        }}
    </style>
</head>
<body>
    <header class="header">
        <div class="header-inner">
            <h1>{L["summary_title"]}</h1>
            <div class="header-actions">
                <span class="header-meta">{datetime.now().strftime('%Y-%m-%d %H:%M')}</span>
                <button class="btn-print" onclick="window.print()">🖨 Print</button>
            </div>
        </div>
    </header>

    <div class="container">
        <!-- Stats Cards -->
        <div class="stats-row">
            <div class="stat-item high">
                <div class="stat-icon">🔴</div>
                <div class="stat-content">
                    <span class="stat-number">{high}</span>
                    <span class="stat-label">{L["high_risk"]}</span>
                </div>
            </div>
            <div class="stat-item medium">
                <div class="stat-icon">🟠</div>
                <div class="stat-content">
                    <span class="stat-number">{medium}</span>
                    <span class="stat-label">{L["medium_risk"]}</span>
                </div>
            </div>
            <div class="stat-item low">
                <div class="stat-icon">🔵</div>
                <div class="stat-content">
                    <span class="stat-number">{low}</span>
                    <span class="stat-label">{L["low_risk"]}</span>
                </div>
            </div>
            <div class="stat-item ok">
                <div class="stat-icon">✓</div>
                <div class="stat-content">
                    <span class="stat-number">{ok}</span>
                    <span class="stat-label">{L["no_risk"]}</span>
                </div>
            </div>
            <div class="stat-item review">
                <div class="stat-icon">⚠</div>
                <div class="stat-content">
                    <span class="stat-number">{review}</span>
                    <span class="stat-label">{L["review_needed"]}</span>
                </div>
            </div>
        </div>

        <!-- Executive Summary -->
        <div class="summary-panel">
            <div class="summary-header">
                <h2>📊 Executive Summary</h2>
                <span style="font-size:12px;color:var(--text-muted)">{total} clauses analyzed</span>
            </div>
            <div class="summary-body">
                <div class="dist-bar-container">
                    <div class="dist-bar">
                        <div class="dist-segment high" style="width:{high_pct}%">{high if high else ""}</div>
                        <div class="dist-segment medium" style="width:{med_pct}%">{medium if medium else ""}</div>
                        <div class="dist-segment low" style="width:{low_pct}%">{low if low else ""}</div>
                        <div class="dist-segment ok" style="width:{ok_pct}%">{ok if ok else ""}</div>
                    </div>
                    <div class="dist-legend">
                        <div class="dist-legend-item"><div class="dist-legend-dot" style="background:var(--danger)"></div>{L["high_risk"]} {high_pct}%</div>
                        <div class="dist-legend-item"><div class="dist-legend-dot" style="background:var(--warning)"></div>{L["medium_risk"]} {med_pct}%</div>
                        <div class="dist-legend-item"><div class="dist-legend-dot" style="background:var(--info)"></div>{L["low_risk"]} {low_pct}%</div>
                        <div class="dist-legend-item"><div class="dist-legend-dot" style="background:var(--success)"></div>{L["no_risk"]} {ok_pct}%</div>
                    </div>
                </div>
'''

    if summary_rows:
        html_content += f'''
                <table class="summary-table">
                    <thead>
                        <tr>
                            <th style="width:30%">Clause</th>
                            <th style="width:15%">Verdict</th>
                            <th style="width:18%">Type</th>
                            <th>Rationale</th>
                        </tr>
                    </thead>
                    <tbody>{summary_rows}</tbody>
                </table>
'''

    html_content += '''
            </div>
        </div>
'''

    # Toolbar
    html_content += f'''
        <!-- Toolbar: Filters + Search -->
        <div class="toolbar">
            <button class="filter-pill active" onclick="filterClauses('all')">全部</button>
            <button class="filter-pill" onclick="filterClauses('high')">🔴 {L["high_risk"]}</button>
            <button class="filter-pill" onclick="filterClauses('medium')">🟠 {L["medium_risk"]}</button>
            <button class="filter-pill" onclick="filterClauses('low')">🔵 {L["low_risk"]}</button>
            <button class="filter-pill" onclick="filterClauses('review')">⚠ {L["review_needed"]}</button>
            <input class="search-box" type="text" placeholder="🔍 Search clauses..." oninput="searchClauses(this.value)">
        </div>

        <!-- Two-Column Layout -->
        <div class="layout">
            <div class="panel">
                <div class="panel-header">
                    <span class="panel-title">📄 {L["contract"]}</span>
                    <span class="panel-badge">{actionable_count} items</span>
                </div>
                <div class="panel-body">
'''

    # Contract clause cards
    judgment_map = {j.section_id: j for j in judgments}
    for section in sections.values():
        j = judgment_map.get(section.section_id)
        if not j:
            continue

        severity_class = "ok"
        if j.severity == "高":
            severity_class = "high"
        elif j.severity == "中" and j.is_actionable():
            severity_class = "medium"
        elif j.severity == "低" and j.is_actionable():
            severity_class = "low"

        review_tag = '<span class="tag review">⚠</span>' if j.needs_review else ''
        heading = _html.escape(_format_heading(section.heading_path))
        section_text = _html.escape(section.text)
        rationale_text = _html.escape(j.rationale)
        evidence_c = _html.escape(j.evidence_contract) if j.evidence_contract else ""
        evidence_b = _html.escape(j.evidence_baseline) if j.evidence_baseline else ""

        html_content += f'''
                    <div class="clause {severity_class}" id="clause-{j.section_id}" data-severity="{severity_class}" data-review="{str(j.needs_review).lower()}" onclick="this.classList.toggle('expanded')">
                        <div class="clause-header">
                            <span class="clause-title">{heading}</span>
                            <div class="clause-tags">
                                <span class="tag {severity_class}">{_html.escape(j.verdict or "")}</span>
                                <span class="tag conf">{_html.escape(j.confidence or "")}</span>
                                {review_tag}
                            </div>
                        </div>
                        <div class="clause-body">
                            <div class="clause-text">{section_text}</div>
                            <div class="clause-analysis">
                                <div class="clause-analysis-label">Risk Analysis</div>
                                <div class="clause-analysis-text">{rationale_text}</div>
                            </div>
                            <div class="evidence-grid">
                                <div class="evidence-block contract">
                                    <div class="evidence-block-label">{L["contract"]} Evidence</div>
                                    {evidence_c}
                                </div>
                                <div class="evidence-block baseline">
                                    <div class="evidence-block-label">{L["baseline"]} Evidence</div>
                                    {evidence_b}
                                </div>
                            </div>
                        </div>
                    </div>
'''

    html_content += f'''
                </div>
            </div>
            <div class="panel">
                <div class="panel-header">
                    <span class="panel-title">📋 {L["baseline"]}</span>
                    <span class="panel-badge">{L["baseline"]}</span>
                </div>
                <div class="panel-body">
                    <div class="baseline-content">
                        {baseline_html}
                    </div>
                </div>
            </div>
        </div>

        <div class="footer">
            Contract Risk Review Tool · Generated automatically for reference only
        </div>
    </div>

    <script>
        function filterClauses(filter) {{
            document.querySelectorAll('.filter-pill').forEach(btn => btn.classList.remove('active'));
            event.target.classList.add('active');
            document.querySelectorAll('.clause').forEach(clause => {{
                const severity = clause.dataset.severity;
                const needsReview = clause.dataset.review === 'true';
                if (filter === 'all') clause.style.display = '';
                else if (filter === 'review') clause.style.display = needsReview ? '' : 'none';
                else clause.style.display = severity === filter ? '' : 'none';
            }});
        }}

        function searchClauses(query) {{
            const q = query.toLowerCase();
            document.querySelectorAll('.clause').forEach(clause => {{
                const text = clause.textContent.toLowerCase();
                clause.style.display = q === '' || text.includes(q) ? '' : 'none';
            }});
            // Reset filter pills
            if (q) document.querySelectorAll('.filter-pill').forEach(btn => btn.classList.remove('active'));
            else document.querySelector('.filter-pill').classList.add('active');
        }}

        function scrollToClause(id) {{
            const el = document.getElementById(id);
            if (!el) return;
            el.scrollIntoView({{ behavior: 'smooth', block: 'center' }});
            el.classList.add('expanded', 'highlight');
            setTimeout(() => el.classList.remove('highlight'), 2000);
        }}

        // Expand first high-risk item by default
        const firstHigh = document.querySelector('.clause.high');
        if (firstHigh) firstHigh.classList.add('expanded');
    </script>
</body>
</html>
'''

    Path(path).write_text(html_content, encoding="utf-8")
    LOGGER.info("HTML report generated: %s", path)


# ============================================================================
# Pipeline Runner
# ============================================================================

DEFAULT_WORKERS = 8


def _load_baseline(path: str) -> str:
    return Path(path).read_text(encoding="utf-8")


def _load_env_file(path: str | Path = ".env") -> None:
    """Load key/value pairs from a .env file if it exists."""
    env_path = Path(path)
    if not env_path.exists():
        return
    for line in env_path.read_text(encoding="utf-8").splitlines():
        stripped = line.strip()
        if not stripped or stripped.startswith("#"):
            continue
        if "=" not in stripped:
            continue
        key, value = stripped.split("=", 1)
        key = key.strip()
        value = value.strip().strip('"').strip("'")
        os.environ.setdefault(key, value)


# Pattern for sections that are purely structural / non-substantive.
# Uses fullmatch — the *entire* text must match to be skipped.
_SKIP_PATTERN = re.compile(
    r'^[\d\s.\-、，,（）()\[\]【】:：]+$'      # only numbers / punctuation
    r'|^第[一二三四五六七八九十百\d]+[条章节部分款项]$'  # bare heading like "第一条"
    r'|^\s*$',                                  # blank
)


def _is_trivial_section(section: Section) -> bool:
    """Return True if a section is purely structural and has no substantive text.

    Conservative: only skips sections where the entire text is numbers/punctuation
    or a bare heading label. Never skips text containing Chinese characters or
    alphabetic words, regardless of length.
    """
    text = section.text.strip()
    if not text:
        return True
    return bool(_SKIP_PATTERN.fullmatch(text))


def _make_skip_judgment(section: Section) -> Judgment:
    """Create a 'Not Applicable' judgment for skipped sections."""
    return Judgment(
        section_id=section.section_id,
        heading_path=section.heading_path,
        page=section.page,
        verdict="不适用/无对应",
        type=None,
        severity=None,
        rationale="Section too short or purely structural; skipped.",
        evidence_contract="",
        evidence_baseline="",
        confidence="高",
        needs_review=False,
    )


def _evaluate_sections(
    sections: Sequence[Section],
    baseline_text: str,
    llm_client: LLMClient,
    max_workers: int,
) -> List[Judgment]:
    if not sections:
        return []

    # Pre-filter trivial sections
    to_eval: list[Section] = []
    results: list[Judgment] = []
    for s in sections:
        if _is_trivial_section(s):
            results.append(_make_skip_judgment(s))
        else:
            to_eval.append(s)

    skipped = len(sections) - len(to_eval)
    if skipped:
        LOGGER.info("Skipped %d trivial sections, %d sections to evaluate", skipped, len(to_eval))

    if not to_eval:
        results.sort(key=lambda j: j.section_id)
        return results

    # Connectivity test with first section
    LOGGER.info("Testing LLM connectivity...")
    first_section = to_eval[0]
    try:
        first_result = llm_client.evaluate_section(first_section, baseline_text)
        LOGGER.info("LLM connected. Processing %d sections...", len(to_eval))
    except Exception as exc:
        LOGGER.error("LLM service unavailable")
        raise RuntimeError("LLM服务不可用，无法完成合同审核。请检查API配置和网络连接。") from exc

    results.append(first_result)
    remaining = to_eval[1:]
    total = len(to_eval)
    completed = 1

    if remaining:
        import threading
        counter_lock = threading.Lock()

        with ThreadPoolExecutor(max_workers=max_workers) as executor:
            future_map = {
                executor.submit(llm_client.evaluate_section, section, baseline_text): section
                for section in remaining
            }
            for future in as_completed(future_map):
                try:
                    result = future.result()
                    results.append(result)
                except Exception as exc:
                    section = future_map[future]
                    LOGGER.error("Failed to process %s: %s", section.section_id, exc)

                with counter_lock:
                    completed += 1
                    if completed % 10 == 0 or completed == total:
                        LOGGER.info("[%d/%d] %.1f%% completed", completed, total, completed / total * 100)

    # Flush cache to disk
    if llm_client.cache:
        llm_client.cache.flush()

    results.sort(key=lambda j: j.section_id)
    return results


def run_review(
    contract_path: str,
    baseline_path: str,
    output_dir: str,
    workers: int = DEFAULT_WORKERS,
    lang: str = "zh",
    llm_client: LLMClient | None = None,
    use_cache: bool = True,
    max_section_chars: int = MAX_SECTION_CHARS,
) -> List[Judgment]:
    """Execute the full review pipeline and return judgments."""

    # Load environment
    _load_env_file(SCRIPT_DIR.parent / ".env")

    output_path = Path(output_dir)
    output_path.mkdir(parents=True, exist_ok=True)

    # Initialize cache
    cache = LLMCache(output_path / ".cache") if use_cache else None
    if cache:
        LOGGER.info("Cache enabled: %s", cache.cache_file)

    # Output paths: {original_name}_{type}_{timestamp}.{ext}
    contract_stem = Path(contract_path).stem
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    prefix = f"{contract_stem}_{timestamp}"

    out_json = output_path / f"{prefix}_judgments.json"
    out_html = output_path / f"{prefix}_report.html"

    # Load inputs
    baseline_text = _load_baseline(baseline_path)
    sections = parse_contract(contract_path, max_section_chars=max_section_chars)
    sections_by_id: Mapping[str, Section] = {section.section_id: section for section in sections}

    LOGGER.info("Parsed %d sections from contract", len(sections))

    # Evaluate with LLM
    if llm_client is None:
        llm_client = LLMClient(cache=cache)
    judgments = _evaluate_sections(sections, baseline_text, llm_client, max_workers=workers)

    # Generate outputs
    export_json(str(out_json), judgments)
    LOGGER.info("JSON output: %s", out_json)

    contract_ext = Path(contract_path).suffix.lower()
    if contract_ext == ".pdf":
        out_annotated = output_path / f"{prefix}_annotated.pdf"
        annotate_pdf(contract_path, str(out_annotated), judgments, sections_by_id, lang)
        LOGGER.info("Annotated PDF: %s", out_annotated)
    else:
        out_annotated = output_path / f"{prefix}_annotated.docx"
        annotate_document(contract_path, str(out_annotated), judgments, sections_by_id, lang)
        LOGGER.info("Annotated document: %s", out_annotated)

    generate_comparison_html(str(out_html), judgments, sections_by_id, baseline_text, lang)

    # Summary
    actionable = [j for j in judgments if j.is_actionable()]
    high_risk = sum(1 for j in actionable if j.severity == "高")
    medium_risk = sum(1 for j in actionable if j.severity == "中")
    low_risk = sum(1 for j in actionable if j.severity == "低")
    review_needed = sum(1 for j in judgments if j.needs_review)

    LOGGER.info("=" * 50)
    LOGGER.info("Review complete!")
    LOGGER.info("High risk: %d, Medium risk: %d, Low risk: %d", high_risk, medium_risk, low_risk)
    LOGGER.info("Needs manual review: %d", review_needed)

    # Cache statistics
    if cache:
        stats = llm_client.cache_stats()
        LOGGER.info("Cache stats: %d hits, %d misses (hit rate: %s)",
                    stats["hits"], stats["misses"], stats["hit_rate"])

    LOGGER.info("=" * 50)

    return judgments


# ============================================================================
# CLI
# ============================================================================

def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        description="Contract Risk Review Tool - 合同风险审核工具",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  # Specify contract and baseline files
  python run_review.py my_contract.docx my_baseline.md

  # Only specify contract, use default baseline
  python run_review.py my_contract.pdf

  # Use all defaults (data/contract.docx, data/baseline.md -> output/)
  python run_review.py

  # With optional flags
  python run_review.py contract.docx baseline.md -o results/ -l en -w 8
        """,
    )
    parser.add_argument(
        "contract",
        nargs="?",
        default=str(DEFAULT_CONTRACT),
        help="Contract file path, .docx or .pdf (default: data/contract.docx)"
    )
    parser.add_argument(
        "baseline",
        nargs="?",
        default=str(DEFAULT_BASELINE),
        help="Baseline markdown file path (default: data/baseline.md)"
    )
    parser.add_argument(
        "-o", "--output",
        default=str(DEFAULT_OUTPUT_DIR),
        help="Output directory (default: output/)"
    )
    parser.add_argument(
        "-l", "--lang",
        choices=["zh", "en"],
        default="zh",
        help="Output language (default: zh)"
    )
    parser.add_argument(
        "-w", "--workers",
        type=int,
        default=DEFAULT_WORKERS,
        help="Number of parallel workers (default: 8)"
    )
    parser.add_argument(
        "--no-cache",
        action="store_true",
        help="Disable LLM response caching"
    )
    parser.add_argument(
        "--max-section-chars",
        type=int,
        default=MAX_SECTION_CHARS,
        help=f"Maximum characters per section before splitting (default: {MAX_SECTION_CHARS})"
    )
    return parser


def _validate_paths(contract: str, baseline: str) -> None:
    """Validate that input files exist and have supported formats."""
    contract_path = Path(contract)
    if not contract_path.exists():
        raise SystemExit(f"Error: Contract file not found: {contract}")
    if contract_path.suffix.lower() not in (".docx", ".pdf"):
        raise SystemExit(f"Error: Unsupported contract format '{contract_path.suffix}', expected .docx or .pdf")
    baseline_path = Path(baseline)
    if not baseline_path.exists():
        raise SystemExit(f"Error: Baseline file not found: {baseline}")


def main(args: list[str] | None = None) -> None:
    logging.basicConfig(
        level=logging.INFO,
        format="[%(levelname)s] %(message)s"
    )
    parser = build_parser()
    namespace = parser.parse_args(args=args)

    _validate_paths(namespace.contract, namespace.baseline)

    run_review(
        contract_path=namespace.contract,
        baseline_path=namespace.baseline,
        output_dir=namespace.output,
        workers=namespace.workers,
        lang=namespace.lang,
        use_cache=not namespace.no_cache,
        max_section_chars=namespace.max_section_chars,
    )


if __name__ == "__main__":
    main()
